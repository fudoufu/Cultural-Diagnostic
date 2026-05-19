import io
import re
import time as _time
from contextlib import contextmanager as _ctxmgr
from typing import Optional
import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import warnings
from factor_analyzer import FactorAnalyzer
import semopy

warnings.filterwarnings("ignore")

# ── Performance timing helpers ────────────────────────────────────────────────────
_timings: list = []

@_ctxmgr
def _timed(label: str):
    t = _time.perf_counter()
    yield
    elapsed = _time.perf_counter() - t
    if st.session_state.get("_show_timing"):
        _timings.append((label, elapsed))

# ── Page config ───────────────────────────────────────────────────────────────────
st.set_page_config(page_title="Cultural Diagnostic", layout="wide", initial_sidebar_state="expanded")

# ── CSS / font injection ──────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html, body, [class*="css"], .stMarkdown, .stDataFrame { font-family: 'Inter', sans-serif !important; }
.stApp, [data-testid="stAppViewContainer"], .main, .block-container { background-color: #F7F9FC; }
.block-container { padding-top: 4.25rem; padding-bottom: 2rem; }
.stTabs [data-baseweb="tab-list"] { gap: 6px; background: transparent; }
.stTabs [data-baseweb="tab"] {
    background-color: #FFFFFF;
    border: 1px solid #D6E0EA;
    border-radius: 6px;
    padding: 6px 18px;
    color: #1A2B3C;
    font-weight: 500;
    font-size: 13px;
}
.stTabs [aria-selected="true"] {
    background-color: #0F4C6B !important;
    color: #FFFFFF !important;
    border-color: #0F4C6B !important;
}
/* Level 2 tabs (Correlation / Descriptive) */
.stTabs .stTabs [aria-selected="true"] {
    background-color: #2E7096 !important;
    border-color: #2E7096 !important;
    color: #FFFFFF !important;
}
/* Level 3 tabs (A1/A2/A3 etc.) */
.stTabs .stTabs .stTabs [aria-selected="true"] {
    background-color: #5A9BB5 !important;
    border-color: #5A9BB5 !important;
    color: #FFFFFF !important;
}
/* Correlational Heatmaps tab (3rd at level 3) — grey to signal supporting analysis */
.stTabs .stTabs .stTabs [data-baseweb="tab"]:nth-child(3) {
    background-color: #DDE3EA !important;
    color: #5A7080 !important;
    border-color: #C8D0DA !important;
}
.stTabs .stTabs .stTabs [data-baseweb="tab"]:nth-child(3)[aria-selected="true"] {
    background-color: #A8B4BF !important;
    border-color: #8C9AAA !important;
    color: #1A2B3C !important;
}
/* Reset: undo grey for 3rd tab at level 4+ (e.g. Outcomes × Outcomes inside heatmaps) */
.stTabs .stTabs .stTabs .stTabs [data-baseweb="tab"]:nth-child(3) {
    background-color: #FFFFFF !important;
    color: #1A2B3C !important;
    border-color: #D6E0EA !important;
}
.stTabs .stTabs .stTabs .stTabs [data-baseweb="tab"]:nth-child(3)[aria-selected="true"] {
    background-color: #8DC0D4 !important;
    border-color: #8DC0D4 !important;
    color: #1A2B3C !important;
}
/* Level 4 tabs (Place / Individual) */
.stTabs .stTabs .stTabs .stTabs [aria-selected="true"] {
    background-color: #8DC0D4 !important;
    border-color: #8DC0D4 !important;
    color: #1A2B3C !important;
}
.metric-card {
    background: #FFFFFF;
    border: 1px solid #D6E0EA;
    border-radius: 6px;
    padding: 8px 12px;
    margin-bottom: 8px;
}
.metric-card .card-label { color: #5A7080; font-size: 11px; font-weight: 500; margin: 0; }
.metric-card .card-value { color: #0F4C6B; font-size: 15px; font-weight: 700; margin: 2px 0 0; }
.metric-card .card-sub   { color: #1A2B3C; font-size: 12px; font-weight: 500; margin: 1px 0 0; }
.metric-card-lg {
    background: #FFFFFF;
    border: 1px solid #D6E0EA;
    border-radius: 8px;
    padding: 14px 18px;
    margin-bottom: 12px;
}
.metric-card-lg .card-label { color: #5A7080; font-size: 12px; font-weight: 500; margin: 0; }
.metric-card-lg .card-value { color: #0F4C6B; font-size: 20px; font-weight: 700; margin: 4px 0 0; }
.metric-card-lg .card-sub   { color: #1A2B3C; font-size: 13px; font-weight: 500; margin: 2px 0 0; }
.section-pill {
    display: inline-block;
    background: #0F4C6B;
    color: white;
    border-radius: 6px;
    padding: 6px 16px;
    font-size: 15px;
    font-weight: 600;
    margin-bottom: 12px;
}
.section-pill-alt {
    display: inline-block;
    background: #00A8A8;
    color: white;
    border-radius: 6px;
    padding: 6px 16px;
    font-size: 15px;
    font-weight: 600;
    margin-bottom: 12px;
    margin-top: 18px;
}
div[data-testid="stRadio"] label,
div[data-testid="stRadio"] label p {
    color: #1A2B3C !important;
    font-weight: 500;
}
h4 {
    color: #0F4C6B !important;
    font-size: 17px !important;
    font-weight: 600 !important;
    margin: 4px 0 12px 0 !important;
}
[data-testid="stCaptionContainer"] p, .stCaption {
    color: #3A4D5C !important;
    font-size: 13px !important;
}
[data-testid="stSlider"] label, [data-testid="stSlider"] p,
[data-testid="stSlider"] [data-testid="stMarkdownContainer"] p {
    color: #1A2B3C !important;
    font-weight: 500 !important;
}
[data-testid="stSlider"] [data-testid="stTickBarMin"],
[data-testid="stSlider"] [data-testid="stTickBarMax"] {
    color: #1A2B3C !important;
}
details summary, [data-testid="stExpander"] summary p {
    color: #1A2B3C !important;
    font-weight: 500 !important;
}
details summary:hover, details summary:hover p,
[data-testid="stExpander"] summary:hover p {
    color: #3A4D5C !important;
}
details[open] summary, details[open] summary p,
details[open] [data-testid="stExpander"] summary p {
    color: #FFFFFF !important;
}
details[open] summary:hover, details[open] summary:hover p,
details[open] [data-testid="stExpander"] summary:hover p {
    color: #1A2B3C !important;
}
.landing-hero {
    min-height: 62vh;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    text-align: center;
    color: #1A2B3C;
}
</style>
""", unsafe_allow_html=True)

# ── Constants ─────────────────────────────────────────────────────────────────────
WOW_THEMES = [
    "Sense of Collective Responsibility", "Individual Accountability",
    "Top-down Decision-making", "Autonomy over Decisions",
    "Prioritise People's Well-Being", "Prioritise Results",
    "Challenge Decisions", "Preserve Cohesion",
    "Follow Procedures", "Adapt to Situation",
    "Stick to Current Ways", "Experiment & Innovate",
    "Prioritise Immediate Results", "Consider the Long Term",
    "Proactive Learning", "Reactive Learning",
    "Target-Driven Interactions", "Invest in Relationships",
    "Plan-Based Working", "Agile Working",
    "Recognise Contributions", "Status Awareness",
]

OUTCOME_LABELS = [
    "Intent to stay", "Good place to work", "Feeling valued", "Sense of pride",
    "Sense of impact", "Empowered to be Effective", "Workload manageability", "Role clarity",
    "Voice heard", "Psychological safety", "Breaking silos",
    "Opportunity for contribution", "Meaningful LM time", "LM effectiveness", "Employer rating",
]

LIKERT_MAP = {
    "Strongly agree": 4, "Agree": 3, "Neither agree nor disagree": 2,
    "Disagree": 1, "Strongly disagree": 0,
}

LOS_ORDER = [
    "Less than 1 year", "1-3 years", "4-10 years", "11-20 years", "More than 20 years",
]

EDI_FILTERS = {
    "Q8":  "Length of service",
    "Q72": "Age",
    "Q73": "Gender",
    "Q74": "Ethnic group",
    "Q75": "Sexual orientation",
    "Q76": "Disabled",
    "Q77": "Carer",
}

# Column ranges
WOW_PLACE_COLS = [f"Q{i}" for i in range(11, 33)]   # Q11–Q32
WOW_IND_COLS   = [f"Q{i}" for i in range(33, 55)]   # Q33–Q54
WOW_ALL_COLS   = WOW_PLACE_COLS + WOW_IND_COLS
OUTCOME_COLS   = [f"Q{i}" for i in range(55, 70)]   # Q55–Q69

WOW_ALL_LABELS = [f"P · {t}" for t in WOW_THEMES] + [f"I · {t}" for t in WOW_THEMES]

# Survey statement text for hover tooltips
WOW_PLACE_STATEMENTS = {
    "Sense of Collective Responsibility":    "Teamwork comes first here – it's about what we achieve together, not individually.",
    "Individual Accountability":    "People are accountable for their own work and credited for their results.",
    "Top-down Decision-making":     "Decisions are mostly made by senior people and passed down.",
    "Autonomy over Decisions":      "People are trusted and encouraged to make decisions at their level.",
    "Prioritise People's Well-Being": "People's wellbeing comes first in everyday work.",
    "Prioritise Results":           "People make decisions based on what will get the best results.",
    "Challenge Decisions":          "It's okay to question decisions and suggest different approaches.",
    "Preserve Cohesion":            "People tend to avoid conflict, even when something needs to be said.",
    "Follow Procedures":            "There are processes we must follow in how we work.",
    "Adapt to Situation":           "People are free to step outside the process based on what's needed.",
    "Stick to Current Ways":        "People tend to keep doing things the way they've always been done.",
    "Experiment & Innovate":        "People are encouraged to try new things and new ways of working.",
    "Prioritise Immediate Results": "People prioritise what will deliver results soon rather than further down the line.",
    "Consider the Long Term":       "People consider what something will mean months or years down the line, not just right now.",
    "Proactive Learning":           "People are always looking for ways to do things better around here.",
    "Reactive Learning":            "We tend to fix things after they go wrong, not before.",
    "Target-Driven Interactions":   "People mostly talk to each other when there's a specific task to deal with.",
    "Invest in Relationships":     "Building good relationships is seen as an important part of the job here.",
    "Plan-Based Working":           "What we do and when we do it is mostly planned out in advance.",
    "Agile Working":                "The work shapes the day – we change what we do based on what's needed at the time.",
    "Recognise Contributions":      "People get credit for what they actually do and contribute.",
    "Status Awareness":      "How senior you are tends to affect how people work with you and treat you.",
}

WOW_IND_STATEMENTS = {
    "Sense of Collective Responsibility":    "I'm more focused on what we achieve as a team than what I achieve personally.",
    "Individual Accountability":    "I like to be responsible for my own work and get on with it independently.",
    "Top-down Decision-making":     "I prefer to know that the right people have approved a decision before moving forward.",
    "Autonomy over Decisions":      "When something needs doing, I'm comfortable making a call and getting on with it.",
    "Prioritise People's Well-Being": "When I have to make a tough decision, I think about the impact on my colleagues first.",
    "Prioritise Results":           "Getting the job done is my main priority.",
    "Challenge Decisions":          "When I disagree with something, I like to raise it.",
    "Preserve Cohesion":            "I prefer to keep the peace and avoid disagreements.",
    "Follow Procedures":            "I work best when there's a clear process to follow.",
    "Adapt to Situation":           "I'm comfortable changing how I work depending on what the situation needs.",
    "Stick to Current Ways":        "I prefer to stick with ways of doing things that I know work.",
    "Experiment & Innovate":        "I often think about how things could be done differently.",
    "Prioritise Immediate Results": "I focus on what needs to get done today rather than further ahead.",
    "Consider the Long Term":       "I think about where things are heading in the long run whenever I make decisions about today.",
    "Proactive Learning":           "I seek out opportunities to learn and improve regularly.",
    "Reactive Learning":            "I mostly learn when something goes wrong and I have to figure it out.",
    "Target-Driven Interactions":   "Most of my conversations with colleagues are about work.",
    "Invest in Relationships":     "I like to have close connections with my colleagues.",
    "Plan-Based Working":           "I like to know the plan upfront and work to it.",
    "Agile Working":                "I prefer to keep plans flexible and have my work shape my day.",
    "Recognise Contributions":      "I care more about what someone can contribute than how long they've been here or what their job title is.",
    "Status Awareness":      "I'm conscious of people's seniority and title when working with them.",
}

OUTCOME_STATEMENTS = {
    "Intent to stay":            "I plan to stay with Somerset Council for the next 12 months.",
    "Good place to work":        "I would recommend Somerset Council as a good place to work.",
    "Feeling valued":            "I am valued as an employee of Somerset Council.",
    "Sense of pride":            "I am proud of the culture and behaviours I see in my day-to-day work.",
    "Sense of impact":           "I can see a link between what I do and the impact on the people who live in Somerset.",
    "Empowered to be Effective": "I am empowered to do a good job.",
    "Workload manageability":    "I find my workload manageable.",
    "Role clarity":              "I know what is expected of me at work.",
    "Voice heard":               "I think my views are heard by the organisation.",
    "Psychological safety":      "I am confident to share any work issues with my colleagues/teams.",
    "Breaking silos":            "I have a good sense of what's happening in other parts of the council.",
    "Opportunity for contribution": "I can contribute to improvements in my area of work.",
    "Meaningful LM time":        "The time I have with my line manager is meaningful to me.",
    "LM effectiveness":          "I believe my line manager manages me effectively.",
    "Employer rating":           "On a rating of 1–10, how do you feel as an employee of Somerset Council right now?",
}

Q9_ORDER = ["My directorate", "My service", "My immediate team", "Other"]

HEADCOUNT = {
    "Adult Services and Housing": 589,
    "Chief Exec's Office (including Executive and Service Directors)": 53,
    "Children, Families and Education": 1475,
    "Community, Place and Economy": 1916,
    "Resources, Strategy and Transformation": 533,
    "Finance and Procurement": 280,
}
TOTAL_HEADCOUNT = 4846

SERVICE_AREA_HEADCOUNT = {
    # Adult Services & Housing
    "Adults Commissioning":         58,
    "Adults Housing General Fund":  120,   # = Housing
    # CEO Office & Elections
    "CEO and Electoral Services":   53,    # Electoral Services (33) + Chief Executives Office & Directors (20)
    # Children, Families & Education
    "Children and Families":        923,
    "Commissioning & Performance":  173,
    "Education":                    379,
    # Community, Place & Economy
    "Economic Development, Skills & Climate": 139,
    "HRA & Property":               178,
    "Infrastructure and Transport": 499,
    "Partnerships, Localities and Culture": 219,
    "Planning":                     154,
    "Regulatory and Operations":    660,
    # Resources, Strategy & Transformation
    "Democratic & Governance":      22,
    "HR and OD":                    118,   # = HR & Organisational Development
    "ICT Transformation":           151,   # = ICT Services
    "Legal":                        34,
    "Public Health":                83,
    "Strategy, Performance and Communications": 125,
}

# Colours
PRIMARY  = "#0F4C6B"
RED      = "#C0392B"
AMBER    = "#F39C12"
GREEN    = "#27AE60"

TREND_COLOURS = [
    "#1f77b4","#ff7f0e","#2ca02c","#d62728","#9467bd",
    "#8c564b","#e377c2","#7f7f7f","#bcbd22","#17becf",
    "#aec7e8","#ffbb78","#98df8a","#ff9896","#c5b0d5",
    "#c49c94","#f7b6d2","#c7c7c7","#dbdb8d","#9edae5",
    "#393b79","#637939",
]

HEATMAP_COLORSCALE = [[0.0, "#0F4C6B"], [0.5, "#FFFFFF"], [1.0, "#C0392B"]]

# ── Data loading ──────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner="Loading survey data…")
def load_data(file_bytes: bytes) -> pd.DataFrame:
    df = pd.read_excel(io.BytesIO(file_bytes), header=0)
    df = df.iloc[:, :77].copy()
    df.columns = [f"Q{i+1}" for i in range(77)]

    # Coalesce Q2–Q7 into a single service_area column
    sa_cols = [f"Q{i}" for i in range(2, 8)]
    _sa_col_to_dir = {
        "Q2": "Adult Services and Housing",
        "Q3": "Chief Exec's Office (including Executive and Service Directors)",
        "Q4": "Children, Families and Education",
        "Q5": "Finance and Procurement",
        "Q6": "Community, Place and Economy",
        "Q7": "Resources, Strategy and Transformation",
    }

    def _coalesce_service(row):
        for col in sa_cols:
            v = row[col]
            if pd.notna(v) and str(v).strip() not in ("", "Not Answered"):
                return str(v).strip(), col
        return "Not Answered", None

    _svc_result = df.apply(_coalesce_service, axis=1)
    df["service_area"] = _svc_result.apply(lambda x: x[0])
    _source_col    = _svc_result.apply(lambda x: x[1])

    # Retrofit Q1 (directorate) based on which service column was filled in,
    # prioritising the service data over the directorate selection for the ~10
    # respondents where these disagree.
    _inferred_dir = _source_col.map(_sa_col_to_dir)
    _mismatch = _inferred_dir.notna() & (_inferred_dir != df["Q1"])
    df.loc[_mismatch, "Q1"] = _inferred_dir[_mismatch]

    # Split service_area into svc_group (before " - " / " – ") and svc_name (after)
    _ICT_SERVICES = {
        "ICT Data and Analytics", "ICT Delivery", "ICT Infrastructure",
        "ICT Security Operations", "ICT Transformation",
    }
    # Services sitting directly under their directorate with no intermediate
    # service area — grouped under a placeholder so the two-level selector works
    _NA_RESOURCES = {
        "Democratic Services & Governance", "Legal",
    }
    _NA_COMMUNITY = {
        "Income & Tenancy Management", "Property", "HRA Property",
    }

    def _split_svc(val):
        if not val or val == "Not Answered":
            return "Not Answered", None
        # Custom groupings
        if val in _ICT_SERVICES:
            return "ICT", val
        if val in _NA_RESOURCES or val in _NA_COMMUNITY:
            return "N/A — go to service", val
        parts = re.split(r" [–\-] ", val, maxsplit=1)
        if len(parts) == 2:
            return parts[0].strip(), parts[1].strip()
        return val, None

    df[["svc_group", "svc_name"]] = df["service_area"].apply(
        lambda v: pd.Series(_split_svc(v))
    )

    # Convert Likert columns Q11–Q68 to numeric
    for col in [f"Q{i}" for i in range(11, 69)]:
        df[col] = df[col].map(LIKERT_MAP)

    # Convert Q69 (1–10) → 0–4.5; keep raw 1–10 values in Q69_raw
    def parse_q69(val):
        if pd.isna(val):
            return np.nan
        try:
            return float(str(val).split()[0]) / 2 - 0.5
        except (ValueError, IndexError):
            return np.nan

    df["Q69_raw"] = df["Q69"].apply(
        lambda v: float(str(v).split()[0]) if not pd.isna(v) else np.nan
    )
    df["Q69"] = df["Q69"].apply(parse_q69)

    return df


# ── Helpers ───────────────────────────────────────────────────────────────────────
def apply_filters(df: pd.DataFrame, selections: dict) -> pd.DataFrame:
    mask = pd.Series(True, index=df.index)
    for col, vals in selections.items():
        if vals:
            mask &= df[col].isin(vals)
    return df[mask]


@st.cache_data(show_spinner=False)
def spearman_matrix(df: pd.DataFrame, x_cols: list, y_cols: list) -> pd.DataFrame:
    from scipy.stats import spearmanr
    all_cols = list(dict.fromkeys(x_cols + y_cols))
    data = df[all_cols].astype(float).to_numpy()
    # scipy.stats.spearmanr on a 2D array computes the full matrix in one
    # vectorised C call — much faster than pandas' pairwise Python loop.
    result = spearmanr(data, nan_policy="omit")
    n = len(all_cols)
    corr_arr = result.statistic if n > 2 else np.array([[1.0, float(result.statistic)], [float(result.statistic), 1.0]])
    corr_df = pd.DataFrame(corr_arr, index=all_cols, columns=all_cols)
    return corr_df.loc[x_cols, y_cols]


def make_writable_matrix(mat: pd.DataFrame) -> pd.DataFrame:
    return pd.DataFrame(mat.to_numpy(copy=True), index=mat.index, columns=mat.columns)


def fill_diagonal_with_nan(mat: pd.DataFrame) -> pd.DataFrame:
    data = mat.to_numpy(copy=True)
    np.fill_diagonal(data, np.nan)
    mat.iloc[:, :] = data
    return mat


def get_filter_options(df: pd.DataFrame, col: str) -> list:
    return sorted(
        v for v in df[col].dropna().unique()
        if str(v).strip() not in ("Not Answered", "")
    )


# ── Visualisations ────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def make_heatmap(matrix: pd.DataFrame, y_labels: list, x_labels: list,
                 y_statements: dict = None, x_statements: dict = None) -> go.Figure:
    z = matrix.values.astype(float)
    text = np.where(np.isnan(z), "", np.round(z, 2).astype(str))
    # Build customdata: each cell gets [y_statement, x_statement]
    _all_stmts = {**WOW_PLACE_STATEMENTS, **WOW_IND_STATEMENTS, **OUTCOME_STATEMENTS}
    _y_st = y_statements or _all_stmts
    _x_st = x_statements or _all_stmts
    customdata = [
        [
            [_y_st.get(y_labels[i], ""), _x_st.get(x_labels[j], "")]
            for j in range(len(x_labels))
        ]
        for i in range(len(y_labels))
    ]
    hover = (
        "<b>%{y}</b><br><i>%{customdata[0]}</i><br><br>"
        "<b>%{x}</b><br><i>%{customdata[1]}</i><br><br>"
        "r = %{z:.3f}<extra></extra>"
    )
    fig = go.Figure(go.Heatmap(
        z=z,
        x=x_labels,
        y=y_labels,
        colorscale=HEATMAP_COLORSCALE,
        zmid=0, zmin=-1, zmax=1,
        text=text,
        texttemplate="%{text}",
        textfont={"size": 9, "color": "#1A2B3C"},
        customdata=customdata,
        hovertemplate=hover,
        colorbar=dict(title="r", tickvals=[-1, -0.5, 0, 0.5, 1], len=0.7),
    ))
    fig.update_layout(
        font=dict(family="Inter", color="#1A2B3C"),
        paper_bgcolor="#F7F9FC",
        plot_bgcolor="#F7F9FC",
        margin=dict(l=10, r=10, t=30, b=10),
        xaxis=dict(tickangle=-40, tickfont=dict(size=11, color="#1A2B3C"), side="bottom"),
        yaxis=dict(tickfont=dict(size=11, color="#1A2B3C"), autorange="reversed"),
        height=max(420, len(y_labels) * 26 + 100),
    )
    return fig


def make_wow_bar_chart(df: pd.DataFrame) -> go.Figure:
    """Grouped bar chart: Place, Individual, Δ (P−I) per WoW theme."""
    p_vals = [df[col].mean() for col in WOW_PLACE_COLS]
    i_vals = [df[col].mean() for col in WOW_IND_COLS]
    delta_vals = [
        p - i if not (np.isnan(p) or np.isnan(i)) else np.nan
        for p, i in zip(p_vals, i_vals)
    ]
    pos_delta = [d if not np.isnan(d) and d >= 0 else np.nan for d in delta_vals]
    neg_delta = [d if not np.isnan(d) and d <  0 else np.nan for d in delta_vals]

    _p_stmts = [WOW_PLACE_STATEMENTS.get(t, "") for t in WOW_THEMES]
    _i_stmts = [WOW_IND_STATEMENTS.get(t, "") for t in WOW_THEMES]

    fig = go.Figure()
    fig.add_trace(go.Bar(
        name="Place (P)", x=WOW_THEMES, y=p_vals,
        marker_color=PRIMARY,
        customdata=_p_stmts,
        hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Place: %{y:.2f}<extra></extra>",
    ))
    fig.add_trace(go.Bar(
        name="Individual (I)", x=WOW_THEMES, y=i_vals,
        marker_color="#5A9BB5",
        customdata=_i_stmts,
        hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Individual: %{y:.2f}<extra></extra>",
    ))
    fig.add_trace(go.Bar(
        name="Δ — Place higher", x=WOW_THEMES, y=pos_delta,
        marker_color=GREEN,
        hovertemplate="%{x}<br>Δ (P−I): %{y:.2f}<extra></extra>",
    ))
    fig.add_trace(go.Bar(
        name="Δ — Individual higher", x=WOW_THEMES, y=neg_delta,
        marker_color=RED,
        hovertemplate="%{x}<br>Δ (P−I): %{y:.2f}<extra></extra>",
    ))
    fig.update_layout(
        barmode="group",
        font=dict(family="Inter", color="#1A2B3C"),
        paper_bgcolor="#F7F9FC",
        plot_bgcolor="#F7F9FC",
        margin=dict(l=10, r=10, t=10, b=180),
        xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
        yaxis=dict(
            title=dict(text="Average Score", font=dict(color="#1A2B3C", size=12)),
            tickfont=dict(size=11, color="#1A2B3C"),
            zeroline=True, zerolinecolor="#9EB5C2", zerolinewidth=1.5,
            gridcolor="#E8EEF2",
        ),
        legend=dict(
            orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
            font=dict(color="#1A2B3C", size=12),
        ),
        height=560,
    )
    return fig


def make_outcome_bar_chart(df: pd.DataFrame, overall_df: pd.DataFrame = None,
                           council_df: pd.DataFrame = None,
                           svc_area_df: pd.DataFrame = None) -> go.Figure:
    """Bar chart: average score for each employee experience outcome.
    overall_df adds a directorate-level comparison bar.
    council_df adds a council-wide comparison bar.
    svc_area_df adds a service-area-level comparison bar."""
    _o_stmts = [OUTCOME_STATEMENTS.get(lbl, "") for lbl in OUTCOME_LABELS]
    vals = [df[col].mean() for col in OUTCOME_COLS]
    fig = go.Figure()
    fig.add_trace(go.Bar(
        name="Selected group",
        x=OUTCOME_LABELS, y=vals,
        marker_color=PRIMARY,
        customdata=_o_stmts,
        hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Score: %{y:.2f}<extra></extra>",
    ))
    if svc_area_df is not None:
        fig.add_trace(go.Bar(
            name="Service area overall",
            x=OUTCOME_LABELS,
            y=[svc_area_df[col].mean() for col in OUTCOME_COLS],
            marker_color="#5A9BB5",
            customdata=_o_stmts,
            hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Service area overall: %{y:.2f}<extra></extra>",
        ))
    if overall_df is not None:
        fig.add_trace(go.Bar(
            name="Directorate overall",
            x=OUTCOME_LABELS,
            y=[overall_df[col].mean() for col in OUTCOME_COLS],
            marker_color="#6BA3BA",
            customdata=_o_stmts,
            hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Directorate overall: %{y:.2f}<extra></extra>",
        ))
    if council_df is not None:
        fig.add_trace(go.Bar(
            name="Council overall",
            x=OUTCOME_LABELS,
            y=[council_df[col].mean() for col in OUTCOME_COLS],
            marker_color="#A8C8D8",
            customdata=_o_stmts,
            hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Council overall: %{y:.2f}<extra></extra>",
        ))
    fig.update_layout(
        barmode="group",
        font=dict(family="Inter", color="#1A2B3C"),
        paper_bgcolor="#F7F9FC",
        plot_bgcolor="#F7F9FC",
        margin=dict(l=10, r=10, t=10, b=160),
        xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
        yaxis=dict(
            title=dict(text="Average Score", font=dict(color="#1A2B3C", size=12)),
            range=[-0.2, 4.5],
            tickfont=dict(size=11, color="#1A2B3C"),
            gridcolor="#E8EEF2",
        ),
        legend=dict(
            orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
            font=dict(color="#1A2B3C", size=12),
        ),
        height=480,
    )
    return fig



def render_heatmap_cards(n: int, matrix: pd.DataFrame, headcount: Optional[int] = None):
    flat = matrix.stack().dropna()
    # Deduplicate symmetric pairs: (A, B) and (B, A) are the same correlation
    seen: set = set()
    unique: dict = {}
    for idx, val in flat.items():
        key = frozenset(idx)
        if key not in seen:
            seen.add(key)
            unique[idx] = val
    flat = pd.Series(unique)
    cols = st.columns(3)
    _pct = f" ({n / headcount:.0%})" if headcount else ""
    with cols[0]:
        st.markdown(f'<div class="metric-card"><p class="card-label">Respondents</p>'
                    f'<p class="card-value">n = {n:,}{_pct}</p></div>', unsafe_allow_html=True)
    if flat.empty:
        return
    top3_pos = flat.nlargest(3)
    top3_neg = flat.nsmallest(3)
    pos_html = "".join(
        f'<p class="card-sub">{idx[0]} × {idx[1]}</p><p class="card-value">r = {val:.3f}</p>'
        for idx, val in top3_pos.items()
    )
    neg_html = "".join(
        f'<p class="card-sub">{idx[0]} × {idx[1]}</p><p class="card-value">r = {val:.3f}</p>'
        for idx, val in top3_neg.items()
    )
    with cols[1]:
        st.markdown(
            f'<div class="metric-card"><p class="card-label">Strongest positive correlations</p>'
            f'{pos_html}</div>',
            unsafe_allow_html=True)
    with cols[2]:
        st.markdown(
            f'<div class="metric-card"><p class="card-label">Strongest negative correlations</p>'
            f'{neg_html}</div>',
            unsafe_allow_html=True)



def svc_selectors(tab_key: str, filtered_df: pd.DataFrame, directorates: list):
    """Directorate + two-level service area/service selectors.
    Returns (chart_df, dir_df, selected_dir, is_overall, svc_area_df).
    svc_area_df is the service area overall when a specific service is selected, else None."""
    _label_style = (
        '<p style="font-size:13px;font-weight:600;color:#5A7080;'
        'text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px{extra}">{text}</p>'
    )
    # ── Directorate ───────────────────────────────────────
    st.markdown(_label_style.format(extra="", text="Select Directorate"),
                unsafe_allow_html=True)
    _dir_prev = st.session_state.get(f"{tab_key}_dir", directorates[0])
    _dir_prev_df = filtered_df[filtered_df["Q1"] == _dir_prev]
    _dir_prev_hc = HEADCOUNT.get(_dir_prev)
    _dir_prev_pct = f" ({len(_dir_prev_df) / _dir_prev_hc:.0%} of headcount)" if _dir_prev_hc else ""
    st.caption(f"{len(_dir_prev_df):,} respondents in **{_dir_prev}**{_dir_prev_pct}")
    sel_dir = st.radio("Directorate", directorates, horizontal=True,
                       label_visibility="collapsed", key=f"{tab_key}_dir")
    dir_df = filtered_df[filtered_df["Q1"] == sel_dir]
    n_dir = len(dir_df)
    svc_groups = get_filter_options(dir_df, "svc_group")
    if not svc_groups:
        return dir_df, dir_df, sel_dir, True, None
    # ── Service area ──────────────────────────────────────
    st.markdown(_label_style.format(extra=";margin-top:8px", text="Select service area"),
                unsafe_allow_html=True)
    _SG_OVERALL = "Overall service area"
    _sg_prev = st.session_state.get(f"{tab_key}_sg", _SG_OVERALL)
    _sg_prev_df = dir_df if _sg_prev == _SG_OVERALL else dir_df[dir_df["svc_group"] == _sg_prev]
    _sg_prev_hc = SERVICE_AREA_HEADCOUNT.get(_sg_prev) if _sg_prev != _SG_OVERALL else HEADCOUNT.get(sel_dir)
    _sg_prev_pct = f" ({len(_sg_prev_df) / _sg_prev_hc:.0%} of headcount)" if _sg_prev_hc else ""
    st.caption(f"{len(_sg_prev_df):,} respondents in **{_sg_prev}**{_sg_prev_pct}")
    sg_opts = [_SG_OVERALL] + list(svc_groups)
    sel_sg = st.radio("Select service area", sg_opts, horizontal=True,
                      label_visibility="collapsed", key=f"{tab_key}_sg")
    if sel_sg == _SG_OVERALL:
        return dir_df, dir_df, sel_dir, True, None
    sg_df = dir_df[dir_df["svc_group"] == sel_sg]
    sub_svcs = get_filter_options(sg_df, "svc_name")
    if not sub_svcs:
        return sg_df, dir_df, sel_dir, False, None
    # ── Service ───────────────────────────────────────────
    st.markdown(_label_style.format(extra=";margin-top:8px", text="Select service"),
                unsafe_allow_html=True)
    _SVC_OVERALL = "Overall service"
    _svc_prev = st.session_state.get(f"{tab_key}_svc", _SVC_OVERALL)
    _svc_prev_df = sg_df if _svc_prev == _SVC_OVERALL else sg_df[sg_df["svc_name"] == _svc_prev] if _svc_prev in sg_df["svc_name"].values else sg_df
    st.caption(f"{len(_svc_prev_df):,} respondents in **{_svc_prev}**")
    svc_opts = [_SVC_OVERALL] + list(sub_svcs)
    sel_svc = st.radio("Select service", svc_opts, horizontal=True,
                       label_visibility="collapsed", key=f"{tab_key}_svc")
    if sel_svc == _SVC_OVERALL:
        return sg_df, dir_df, sel_dir, False, None
    return sg_df[sg_df["svc_name"] == sel_svc], dir_df, sel_dir, False, sg_df


def render_summary_cards(n: int, top3_high: list, top3_low: list, headcount: Optional[int] = None):
    """top3_high / top3_low are lists of (label, value_str) tuples."""
    _pct = f" ({n / headcount:.0%})" if headcount else ""
    cols = st.columns(3)
    with cols[0]:
        st.markdown(f'<div class="metric-card"><p class="card-label">Respondents</p>'
                    f'<p class="card-value">n = {n:,}{_pct}</p></div>', unsafe_allow_html=True)
    high_html = "".join(
        f'<p class="card-sub">{lbl}</p><p class="card-value">{val}</p>'
        for lbl, val in top3_high
    )
    low_html = "".join(
        f'<p class="card-sub">{lbl}</p><p class="card-value">{val}</p>'
        for lbl, val in top3_low
    )
    with cols[1]:
        st.markdown(
            f'<div class="metric-card"><p class="card-label">Highest scoring</p>'
            f'{high_html}</div>',
            unsafe_allow_html=True)
    with cols[2]:
        st.markdown(
            f'<div class="metric-card"><p class="card-label">Lowest scoring</p>'
            f'{low_html}</div>',
            unsafe_allow_html=True)


# ── SEM / EFA helpers ─────────────────────────────────────────────────────────────
@st.cache_data(show_spinner="Running EFA…")
def run_efa(outcome_df: pd.DataFrame, n_factors: int):
    clean = outcome_df.dropna()
    fa = FactorAnalyzer(n_factors=n_factors, rotation="varimax", method="minres")
    fa.fit(clean)
    loadings = pd.DataFrame(
        fa.loadings_,
        index=outcome_df.columns.tolist(),
        columns=[f"LV{i+1}" for i in range(n_factors)],
    )
    return loadings, len(clean)


@st.cache_data(show_spinner=False)
def compute_factor_scores(outcome_df: pd.DataFrame, n_factors: int) -> pd.DataFrame:
    """Return per-respondent factor scores (shape: n_respondents × n_factors)."""
    clean = outcome_df.dropna()
    fa = FactorAnalyzer(n_factors=n_factors, rotation="varimax", method="minres")
    fa.fit(clean)
    scores = fa.transform(clean)
    cols = [f"LV{i+1}" for i in range(n_factors)]
    return pd.DataFrame(scores, index=clean.index, columns=cols)


@st.cache_data(show_spinner=False)
def run_lf_driver_analysis(
    df: pd.DataFrame,
    factor_scores: pd.DataFrame,
    predictor_cols: tuple,
    predictor_labels: tuple,
    p_threshold: float,
):
    """Run backward elimination of WoW themes against each latent factor score.
    No correlation pre-filter — all predictors enter the model."""
    import statsmodels.api as sm

    col_to_lbl = dict(zip(predictor_cols, predictor_labels))
    lf_cols = list(factor_scores.columns)
    combined = df[list(predictor_cols)].join(factor_scores, how="inner").dropna()
    results = {}

    for lf_col in lf_cols:
        y = combined[lf_col].values.astype(float)
        eligible_cols = list(predictor_cols)

        # Full model — check Significance F
        X_full = sm.add_constant(
            combined[eligible_cols].values.astype(float), has_constant="add"
        )
        full_model = sm.OLS(y, X_full).fit()
        sig_f = float(full_model.f_pvalue)
        adj_r2_full = float(full_model.rsquared_adj)

        if sig_f >= 0.05:
            results[lf_col] = {
                "status": "fail_sig_f",
                "sig_f": sig_f,
                "adj_r2_full": adj_r2_full,
                "n": len(combined),
            }
            continue

        # Backward elimination
        elim_log, coef_df, fitted, resid, adj_r2_final, r2_final, n, retained = \
            run_backward_elimination(combined, lf_col, tuple(eligible_cols), p_threshold)

        if coef_df is None:
            results[lf_col] = {
                "status": "all_eliminated",
                "sig_f": sig_f,
                "adj_r2_full": adj_r2_full,
                "n": len(combined),
            }
            continue

        factor_std = float(combined[lf_col].std())
        coef_df = coef_df.copy()
        coef_df["label"] = coef_df["col"].map(col_to_lbl)
        coef_df["sig"] = coef_df["p_value"].apply(
            lambda p: "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else ""
        )
        coef_df["point_change"] = coef_df["β_std"] * factor_std
        coef_df["text"] = coef_df.apply(
            lambda r: f"{r['β_std']:.2f}{r['sig']}  (Δ{r['point_change']:+.2f})", axis=1
        )

        results[lf_col] = {
            "status": "ok",
            "sig_f": sig_f,
            "adj_r2_full": adj_r2_full,
            "adj_r2_final": adj_r2_final,
            "r2_final": r2_final,
            "n": n,
            "coef_df": coef_df,
            "elim_log": elim_log,
            "retained": retained,
        }

    return results


@st.cache_data(show_spinner="Fitting SEM — this may take a moment…")
def run_sem(fit_df: pd.DataFrame, wow_cols: tuple, factor_map_items: tuple):
    """factor_map_items: tuple of (outcome_col, lv_name) pairs (hashable for cache)."""
    factor_map = dict(factor_map_items)
    lv_names = sorted(set(factor_map.values()))

    meas_lines, valid_lvs = [], []
    for lv in lv_names:
        items = [k for k, v in factor_map.items() if v == lv]
        if len(items) >= 2:
            meas_lines.append(f"{lv} =~ {' + '.join(items)}")
            valid_lvs.append(lv)

    if not valid_lvs:
        return None, None, "No latent variable has ≥2 assigned outcomes.", ""

    struct_lines = [f"{lv} ~ {' + '.join(wow_cols)}" for lv in valid_lvs]
    model_desc = "\n".join(meas_lines + struct_lines)

    try:
        with _timed("SEM model fit"):
            model = semopy.Model(model_desc)
            model.fit(fit_df)
            results = model.inspect(mode="list")
            try:
                fit_stats = semopy.calc_stats(model)
            except Exception:
                fit_stats = None
        return results, fit_stats, None, model_desc
    except Exception as e:
        return None, None, str(e), model_desc


@st.cache_data(show_spinner=False)
def run_driver_analysis_batch(
    df: pd.DataFrame,
    outcome_cols: tuple,
    outcome_labels: tuple,
    predictor_cols: tuple,
    predictor_labels: tuple,
    r_threshold: float,
    p_threshold: float,
):
    """For each outcome: pre-filter predictors by |r| >= r_threshold, run full OLS,
    check Significance F, then run backward elimination. Returns dict keyed by outcome label."""
    import statsmodels.api as sm
    from scipy.stats import spearmanr

    col_to_lbl = dict(zip(predictor_cols, predictor_labels))
    results = {}

    for out_col, out_lbl in zip(outcome_cols, outcome_labels):
        # ── Step 1: pre-filter by Spearman r ──────────────────
        eligible_cols = []
        for pc in predictor_cols:
            valid = df[[pc, out_col]].dropna()
            if len(valid) < 5:
                continue
            r, _ = spearmanr(valid[pc], valid[out_col])
            if abs(r) >= r_threshold:
                eligible_cols.append(pc)

        if not eligible_cols:
            results[out_lbl] = {"status": "no_predictors"}
            continue

        fit_df = df[eligible_cols + [out_col]].dropna()
        y = fit_df[out_col].values.astype(float)

        # ── Step 2: full model ─────────────────────────────────
        X_full = sm.add_constant(fit_df[eligible_cols].values.astype(float), has_constant="add")
        full_model = sm.OLS(y, X_full).fit()
        sig_f = float(full_model.f_pvalue)
        adj_r2_full = float(full_model.rsquared_adj)

        if sig_f >= 0.05:
            results[out_lbl] = {
                "status": "fail_sig_f",
                "sig_f": sig_f,
                "adj_r2_full": adj_r2_full,
                "n": len(fit_df),
                "n_eligible": len(eligible_cols),
            }
            continue

        # ── Step 3: backward elimination ──────────────────────
        elim_log, coef_df, fitted, resid, adj_r2_final, r2_final, n, retained = \
            run_backward_elimination(df, out_col, tuple(eligible_cols), p_threshold)

        if coef_df is None:
            results[out_lbl] = {
                "status": "all_eliminated",
                "sig_f": sig_f,
                "adj_r2_full": adj_r2_full,
                "n": len(fit_df),
                "n_eligible": len(eligible_cols),
            }
            continue

        outcome_std = float(y.std())
        coef_df = coef_df.copy()
        coef_df["label"] = coef_df["col"].map(col_to_lbl)
        coef_df["sig"] = coef_df["p_value"].apply(
            lambda p: "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else ""
        )
        coef_df["point_change"] = coef_df["β_std"] * outcome_std
        coef_df["text"] = coef_df.apply(
            lambda r: f"{r['β_std']:.2f}{r['sig']}  (Δ{r['point_change']:+.2f} pts)", axis=1
        )

        results[out_lbl] = {
            "status": "ok",
            "sig_f": sig_f,
            "adj_r2_full": adj_r2_full,
            "adj_r2_final": adj_r2_final,
            "r2_final": r2_final,
            "n": n,
            "n_eligible": len(eligible_cols),
            "coef_df": coef_df,
            "outcome_std": outcome_std,
            "fitted": fitted,
            "elim_log": elim_log,
            "retained": retained,
        }

    return results


@st.cache_data(show_spinner="Running backward elimination…")
def run_backward_elimination(
    df: pd.DataFrame, outcome_col: str, predictor_cols: tuple, p_threshold: float
):
    import statsmodels.api as sm

    fit_df = df[list(predictor_cols) + [outcome_col]].dropna()
    predictors = list(predictor_cols)
    y = fit_df[outcome_col].values.astype(float)
    elim_log = []
    step = 0

    while predictors:
        X = sm.add_constant(fit_df[predictors].values.astype(float), has_constant="add")
        model = sm.OLS(y, X).fit()
        pvals = pd.Series(model.pvalues[1:], index=predictors).dropna()
        if pvals.empty:
            break
        max_p = pvals.max()
        worst = pvals.idxmax()
        if worst not in predictors:
            break

        elim_log.append({
            "Step": step,
            "Predictors in model": len(predictors),
            "Adj. R²": round(model.rsquared_adj, 4),
            "Removed this step": worst if max_p > p_threshold else "—",
            "p-value of removed": round(float(max_p), 4) if max_p > p_threshold else "—",
        })

        if max_p <= p_threshold:
            break

        predictors.remove(worst)
        step += 1

    if not predictors:
        return elim_log, None, None, None, None, None, len(fit_df), []

    X_final = fit_df[predictors].values.astype(float)
    X_const = sm.add_constant(X_final, has_constant="add")
    final = sm.OLS(y, X_const).fit()

    y_std = y.std()
    coef_df = pd.DataFrame({
        "col":       predictors,
        "β":         final.params[1:],
        "std_err":   final.bse[1:],
        "p_value":   final.pvalues[1:],
        "β_std":     [
            final.params[i + 1] * fit_df[predictors[i]].std() / y_std
            for i in range(len(predictors))
        ],
    })

    return (
        elim_log,
        coef_df,
        final.fittedvalues.tolist(),
        final.resid.tolist(),
        float(final.rsquared_adj),
        float(final.rsquared),
        len(fit_df),
        predictors,
    )


# ── Table builders ────────────────────────────────────────────────────────────────
def _rag(val, is_delta=False):
    if pd.isna(val):
        return ""
    if is_delta:
        if abs(val) > 1.0:
            return f"background-color: {RED}; color: white"
        if abs(val) > 0.5:
            return f"background-color: {AMBER}; color: #1A2B3C"
        return ""
    if val < 2.5:
        return f"background-color: {RED}; color: white"
    if val <= 3.5:
        return f"background-color: {AMBER}; color: #1A2B3C"
    return f"background-color: {GREEN}; color: #1A2B3C"


def build_wow_table(df: pd.DataFrame, breakdown_col: str,
                    groups: list) -> tuple[pd.DataFrame, object]:
    """Returns (raw DataFrame, Styler) for the WoW I/P/Δ table."""
    rows = []
    for idx, theme in enumerate(WOW_THEMES):
        p_col = WOW_PLACE_COLS[idx]
        i_col = WOW_IND_COLS[idx]
        row = {"Ways of Working": theme}

        p_ov = df[p_col].mean()
        i_ov = df[i_col].mean()
        row["Overall — P"] = round(p_ov, 2) if not np.isnan(p_ov) else None
        row["Overall — I"] = round(i_ov, 2) if not np.isnan(i_ov) else None
        row["Overall — Δ"] = (round(i_ov - p_ov, 2)
                              if not (np.isnan(p_ov) or np.isnan(i_ov)) else None)

        for g in groups:
            sub = df[df[breakdown_col] == g]
            pm = sub[p_col].mean() if len(sub) else np.nan
            im = sub[i_col].mean() if len(sub) else np.nan
            row[f"{g} — P"] = round(pm, 2) if not np.isnan(pm) else None
            row[f"{g} — I"] = round(im, 2) if not np.isnan(im) else None
            row[f"{g} — Δ"] = (round(im - pm, 2)
                                if not (np.isnan(pm) or np.isnan(im)) else None)
        rows.append(row)

    tdf = pd.DataFrame(rows).set_index("Ways of Working")

    def style_col(col):
        is_delta = "— Δ" in col.name
        return [_rag(v, is_delta) for v in col]

    styler = (tdf.style
              .apply(style_col, axis=0)
              .format("{:.2f}", na_rep="—")
              .set_table_styles([
                  {"selector": "th", "props": [("font-weight", "600"),
                                               ("background-color", "#F0F4F8"),
                                               ("color", "#1A2B3C")]},
                  {"selector": "tr:nth-child(even) td",
                   "props": [("background-color", "#FAFBFC")]},
              ]))
    return tdf, styler


def build_outcome_table(df: pd.DataFrame, breakdown_col: str,
                        groups: list) -> tuple[pd.DataFrame, object]:
    rows = []
    for label, col in zip(OUTCOME_LABELS, OUTCOME_COLS):
        row = {"Outcome": label}
        ov = df[col].mean()
        row["Overall"] = round(ov, 2) if not np.isnan(ov) else None
        for g in groups:
            sub = df[df[breakdown_col] == g]
            m = sub[col].mean() if len(sub) else np.nan
            row[g] = round(m, 2) if not np.isnan(m) else None
        rows.append(row)

    tdf = pd.DataFrame(rows).set_index("Outcome")

    def style_col(col):
        return [_rag(v, is_delta=False) for v in col]

    styler = (tdf.style
              .apply(style_col, axis=0)
              .format("{:.2f}", na_rep="—")
              .set_table_styles([
                  {"selector": "th", "props": [("font-weight", "600"),
                                               ("background-color", "#F0F4F8"),
                                               ("color", "#1A2B3C")]},
                  {"selector": "tr:nth-child(even) td",
                   "props": [("background-color", "#FAFBFC")]},
              ]))
    return tdf, styler


def wow_table_cards(n, tdf, headcount: Optional[int] = None):
    p = pd.to_numeric(tdf["Overall — P"], errors="coerce")
    i = pd.to_numeric(tdf["Overall — I"], errors="coerce")
    combined = pd.concat([
        p.rename(index=lambda x: f"P · {x}"),
        i.rename(index=lambda x: f"I · {x}"),
    ]).dropna()
    if combined.empty:
        render_summary_cards(n, [], [], headcount)
        return
    top3_high = [(lbl, f"{val:.2f}") for lbl, val in combined.nlargest(3).items()]
    top3_low  = [(lbl, f"{val:.2f}") for lbl, val in combined.nsmallest(3).items()]
    render_summary_cards(n, top3_high, top3_low, headcount)


def outcome_table_cards(n, tdf, headcount: Optional[int] = None):
    ov = pd.to_numeric(tdf["Overall"], errors="coerce").dropna()
    if ov.empty:
        render_summary_cards(n, [], [], headcount)
        return
    top3_high = [(lbl, f"{val:.2f}") for lbl, val in ov.nlargest(3).items()]
    top3_low  = [(lbl, f"{val:.2f}") for lbl, val in ov.nsmallest(3).items()]
    render_summary_cards(n, top3_high, top3_low, headcount)


# ── Sidebar ───────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(
        f'<div style="color:{PRIMARY};font-size:19px;font-weight:700;'
        f'padding-bottom:10px">Cultural Diagnostic</div>',
        unsafe_allow_html=True,
    )
    uploaded = st.file_uploader("Upload survey export (.xlsx)", type=["xlsx"])
    st.markdown("---")
    st.checkbox("Show performance timing", key="_show_timing", value=False)
    st.markdown("---")
    st.markdown("**EDI Filters**")
    filter_selections: dict = {}

# ── Landing screen ────────────────────────────────────────────────────────────────
if uploaded is None:
    st.markdown(
        f"""
        <div class="landing-hero" style="padding:24px 40px 12px;">
            <div style="font-size:52px;margin-bottom:16px">📊</div>
            <h2 style="color:{PRIMARY};margin-bottom:8px">Our Culture Discovery - survey analysis tool</h2>
            <p style="font-size:15px;color:#5A7080">
                Upload a survey export using the panel on the left to get started.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.stop()

# ── Load & process data ───────────────────────────────────────────────────────────
with _timed("load_data"):
    df = load_data(uploaded.read())

# Build EDI filter controls in sidebar
with st.sidebar:
    for q_col, label in EDI_FILTERS.items():
        opts = get_filter_options(df, q_col)
        if opts:
            sel = st.multiselect(label, opts, default=[], key=f"f_{q_col}")
            if sel:
                filter_selections[q_col] = sel

filtered = apply_filters(df, filter_selections)
n_total = len(filtered)

if n_total == 0:
    st.warning("No respondents match the current filters. Adjust the filters to see results.")
    st.stop()

# Dynamic lookup values (read from data, not hardcoded)
directorates = get_filter_options(df, "Q1")
q9_levels    = get_filter_options(df, "Q9")

# ── Top-level section tabs ────────────────────────────────────────────────────────
sec_a, sec_b, sec_c, sec_d = st.tabs([
    "Section A: Council-Wide",
    "Section B: Directorate & Service Deep Dive",
    "Section C: Org Health Analysis",
    "Section D: EDI Views",
])

# ═══════════════════════════════════════════════════════════════════════════════════
# SECTION A
# ═══════════════════════════════════════════════════════════════════════════════════
with sec_a:
    corr_group, desc_group = st.tabs(["A1: Drivers Analysis", "A2: Descriptive Analysis"])

    # ── Correlation Analysis group ────────────────────────
    with _timed("A — corr_group render"), corr_group:
        # ── ODA tabs at the top level ─────────────────────────
        oda_outcomes_tab, oda_lf_tab, heatmaps_tab = st.tabs([
            "A1.1: Individual Outcomes",
            "A1.2: Grouped Outcomes",
            "A1.3: Correlational Heatmaps",
        ])

        with oda_outcomes_tab:
            st.caption(
                "For each employee experience outcome, this analysis: (1) keeps only WoW themes "
                "with |r| ≥ the correlation floor — weak predictors are excluded before modelling "
                "begins; (2) runs a multiple regression with all remaining themes simultaneously "
                "and checks whether the overall model is statistically reliable (Significance F < 0.05); "
                "(3) progressively removes the weakest contributors until only those with an independent, "
                "meaningful relationship remain. The priority matrix shows the confirmed behavioural "
                "drivers across all outcomes at a glance."
            )

            # ── Controls ──────────────────────────────────────────────
            oda_c1, oda_c2, oda_c3 = st.columns([3, 2, 2])
            with oda_c1:
                wow_choice_oda = st.radio(
                    "WoW predictors", ["Place (P)", "Individual (I)"],
                    horizontal=True, key="oda_wow"
                )
            with oda_c2:
                r_thresh_oda = st.slider(
                    "Correlation floor |r| ≥", 0.00, 0.40, 0.20, 0.05, key="oda_r"
                )
                st.markdown(
                    '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                    'Recommended: 0.20</p>',
                    unsafe_allow_html=True)
            with oda_c3:
                p_thresh_oda = st.slider(
                    "P-value threshold", 0.00, 0.10, 0.05, 0.01, key="oda_p"
                )
                st.markdown(
                    '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                    'Recommended: 0.05</p>',
                    unsafe_allow_html=True)

            # ── Org level filter ───────────────────────────────────────
            _oda_q9_ordered = sorted(q9_levels, key=lambda x: next(
                (i for i, o in enumerate(Q9_ORDER) if o.lower() == x.lower()), len(Q9_ORDER)))
            st.markdown(
                '<p style="font-size:13px;font-weight:600;color:#5A7080;'
                'text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px">'
                'Select org level</p>', unsafe_allow_html=True)
            _oda_prev_q9 = st.session_state.get("oda_q9", "Overall")
            _oda_prev_q9_df = filtered if _oda_prev_q9 == "Overall" else filtered[filtered["Q9"] == _oda_prev_q9]
            _oda_q9_pct = f" ({len(_oda_prev_q9_df) / TOTAL_HEADCOUNT:.0%} of total)"
            st.caption(f"{len(_oda_prev_q9_df):,} respondents in **{_oda_prev_q9}**{_oda_q9_pct}")
            sel_oda_q9 = st.radio(
                "Select org level", ["Overall"] + _oda_q9_ordered,
                horizontal=True, label_visibility="collapsed", key="oda_q9"
            )
            oda_data = filtered if sel_oda_q9 == "Overall" else filtered[filtered["Q9"] == sel_oda_q9]

            if "Place" in wow_choice_oda:
                oda_pred_cols = tuple(WOW_PLACE_COLS)
                oda_col_to_lbl = dict(zip(WOW_PLACE_COLS, WOW_THEMES))
            else:
                oda_pred_cols = tuple(WOW_IND_COLS)
                oda_col_to_lbl = dict(zip(WOW_IND_COLS, WOW_THEMES))
            oda_pred_labels = tuple(oda_col_to_lbl[c] for c in oda_pred_cols)

            with st.spinner("Running outcome driver analysis…"):
                oda_results = run_driver_analysis_batch(
                    oda_data,
                    tuple(OUTCOME_COLS), tuple(OUTCOME_LABELS),
                    oda_pred_cols, oda_pred_labels,
                    r_thresh_oda, p_thresh_oda,
                )

            ok_outcomes = {lbl: res for lbl, res in oda_results.items()
                           if res["status"] == "ok"}

            _oda_combined = {}
            for _res in ok_outcomes.values():
                for _, _row in _res["coef_df"].iterrows():
                    _lbl = _row["label"]
                    _oda_combined[_lbl] = _oda_combined.get(_lbl, 0.0) + _row["β_std"]
            if _oda_combined:
                _oda_sorted = sorted(_oda_combined.items(), key=lambda x: x[1], reverse=True)
                _oda_labels = [x[0] for x in _oda_sorted]
                _oda_scores = [x[1] for x in _oda_sorted]
                st.markdown("#### Top WoW Drivers — Combined Across All Outcomes")
                st.caption(
                    "Each bar shows the sum of standardised β values for that WoW theme across all "
                    "individual outcome models where it was a confirmed driver, preserving sign. "
                    "A consistently positive theme scores high; mixed effects partially cancel out. "
                    "Red = net positive effect; blue = net negative."
                )
                _oda_c_stmts = WOW_PLACE_STATEMENTS if "Place" in wow_choice_oda else WOW_IND_STATEMENTS
                _fig_oda_c = go.Figure(go.Bar(
                    x=_oda_labels,
                    y=_oda_scores,
                    text=[f"{s:+.2f}" for s in _oda_scores],
                    textposition="outside",
                    textfont=dict(color="#1A2B3C", size=10),
                    marker_color=[RED if s >= 0 else PRIMARY for s in _oda_scores],
                    customdata=[_oda_c_stmts.get(lbl, "") for lbl in _oda_labels],
                    hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Combined β = %{y:.3f}<extra></extra>",
                ))
                _fig_oda_c.update_layout(
                    font=dict(family="Inter", color="#1A2B3C"),
                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                    margin=dict(l=10, r=10, t=40, b=260),
                    xaxis=dict(tickangle=-40, tickfont=dict(color="#1A2B3C", size=10)),
                    yaxis=dict(
                        title=dict(text="Combined β", font=dict(color="#1A2B3C")),
                        zeroline=True, zerolinecolor="#D6E0EA",
                        tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                    ),
                    height=560,
                )
                st.plotly_chart(_fig_oda_c, use_container_width=True, key="oda_combined_bar")

            # ── Priority matrix ────────────────────────────────────────
            st.markdown("---")
            st.markdown("#### Top Ways of Working Drivers — Per Outcome")
            st.caption(
                "Each cell shows the standardised β for a WoW theme that survived backward "
                "elimination for that outcome's model. Blank = not a confirmed driver (below "
                "the correlation floor or eliminated). Blue = positive effect (more of this "
                "behaviour → better outcome score); red = negative effect."
            )

            if not ok_outcomes:
                st.warning(
                    "No outcomes produced a statistically reliable model with the current settings. "
                    "Try lowering the correlation floor or raising the p-value threshold."
                )
            else:
                all_retained_labels = sorted(set(
                    oda_col_to_lbl[c]
                    for res in ok_outcomes.values()
                    for c in res["retained"]
                ))
                matrix_rows = {}
                for out_lbl, res in ok_outcomes.items():
                    lbl_to_beta = {
                        row["label"]: row["β_std"]
                        for _, row in res["coef_df"].iterrows()
                    }
                    matrix_rows[out_lbl] = {
                        theme: lbl_to_beta.get(theme, np.nan)
                        for theme in all_retained_labels
                    }
                priority_df = pd.DataFrame(matrix_rows, index=all_retained_labels)
                _ordered_outcomes = [lbl for lbl in OUTCOME_LABELS if lbl in priority_df.columns]
                priority_df = priority_df[_ordered_outcomes]

                beta_abs_max = priority_df.abs().max().max()
                beta_abs_max = max(beta_abs_max, 0.01)

                priority_df_T = priority_df.T  # outcomes as rows, WoW as columns
                fig_matrix = go.Figure(go.Heatmap(
                    z=priority_df_T.values.tolist(),
                    x=list(priority_df_T.columns),
                    y=list(priority_df_T.index),
                    colorscale=[
                        [0.0,  "#0F4C6B"],
                        [0.5,  "#F7F9FC"],
                        [1.0,  "#C0392B"],
                    ],
                    zmid=0,
                    zmin=-beta_abs_max,
                    zmax=beta_abs_max,
                    text=[[
                        f"{v:.2f}" if not np.isnan(v) else ""
                        for v in row
                    ] for row in priority_df_T.values.tolist()],
                    texttemplate="%{text}",
                    textfont=dict(size=11, color="#1A2B3C"),
                    hoverongaps=False,
                    hovertemplate="<b>%{x}</b><br>%{y}<br>β = %{z:.3f}<extra></extra>",
                    colorbar=dict(title="β", len=0.7),
                ))
                fig_matrix.update_layout(
                    font=dict(family="Inter", color="#1A2B3C"),
                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                    margin=dict(l=10, r=10, t=10, b=160),
                    xaxis=dict(
                        tickfont=dict(color="#1A2B3C", size=11),
                        tickangle=-40,
                        side="bottom",
                    ),
                    yaxis=dict(tickfont=dict(color="#1A2B3C", size=11), autorange="reversed"),
                    height=max(350, 36 * len(ok_outcomes)),
                )
                st.plotly_chart(fig_matrix, use_container_width=True, key="oda_matrix")

            # ── Per-outcome detail ─────────────────────────────────────
            st.markdown("---")
            st.markdown('<p style="font-size:14px;font-weight:600;color:#5A7080;margin-top:8px;margin-bottom:4px">Breakdown per Outcome</p>', unsafe_allow_html=True)

            for out_lbl in OUTCOME_LABELS:
                res = oda_results.get(out_lbl, {"status": "no_predictors"})
                status = res["status"]
                icon = "✓" if status == "ok" else "✗"
                with st.expander(f"{icon}  {out_lbl}"):
                    if status == "no_predictors":
                        st.info(
                            f"No WoW themes reached |r| ≥ {r_thresh_oda:.2f} for this outcome. "
                            "No regression was run."
                        )
                    elif status == "fail_sig_f":
                        st.error(
                            f"**Model not statistically reliable** — Significance F = {res['sig_f']:.4f} "
                            f"(> 0.05). The {res['n_eligible']} predictors that cleared the correlation "
                            f"floor did not collectively explain this outcome reliably. "
                            "Do not read into the individual coefficients."
                        )
                        mc1, mc2, mc3 = st.columns(3)
                        for col_ui, label, value in [
                            (mc1, "Respondents", f"n = {res['n']:,}"),
                            (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                            (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                        ]:
                            with col_ui:
                                st.markdown(
                                    f'<div class="metric-card"><p class="card-label">{label}</p>'
                                    f'<p class="card-value">{value}</p></div>',
                                    unsafe_allow_html=True)
                    elif status == "all_eliminated":
                        st.warning(
                            "All predictors were eliminated during backward stepwise. "
                            "Try raising the p-value threshold."
                        )
                        mc1, mc2, mc3 = st.columns(3)
                        for col_ui, label, value in [
                            (mc1, "Respondents", f"n = {res['n']:,}"),
                            (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                            (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                        ]:
                            with col_ui:
                                st.markdown(
                                    f'<div class="metric-card"><p class="card-label">{label}</p>'
                                    f'<p class="card-value">{value}</p></div>',
                                    unsafe_allow_html=True)
                    else:
                        coef_df = res["coef_df"]
                        # Summary cards
                        mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                        for col_ui, label, value in [
                            (mc1, "Respondents",         f"n = {res['n']:,}"),
                            (mc2, "Significance F",      f"{res['sig_f']:.4f}"),
                            (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                            (mc4, "Adj R² (final model)",f"{res['adj_r2_final']:.3f}"),
                            (mc5, "Confirmed drivers",   str(len(res["retained"]))),
                        ]:
                            with col_ui:
                                st.markdown(
                                    f'<div class="metric-card"><p class="card-label">{label}</p>'
                                    f'<p class="card-value">{value}</p></div>',
                                    unsafe_allow_html=True)

                        # β coefficient chart
                        st.markdown("##### Confirmed Drivers — Standardised β Coefficients")
                        st.caption(
                            "Each bar shows a retained WoW theme's standardised β and the equivalent "
                            "point change on the 1–5 outcome scale (Δ pts). β reflects the independent "
                            "contribution of that theme after controlling for all others — a 1 SD shift "
                            "in the WoW theme produces the shown point change in the outcome. "
                            "* p<0.05  ** p<0.01  *** p<0.001"
                        )
                        coef_sorted = coef_df.reindex(
                            coef_df["β_std"].abs().sort_values(ascending=False).index
                        )
                        _coef_stmts = WOW_PLACE_STATEMENTS if "Place" in wow_choice_oda else WOW_IND_STATEMENTS
                        fig_coef = go.Figure(go.Bar(
                            x=coef_sorted["label"],
                            y=coef_sorted["β_std"],
                            text=coef_sorted["text"],
                            textposition="outside",
                            textfont=dict(color="#1A2B3C", size=10),
                            marker_color=[
                                RED if v >= 0 else PRIMARY
                                for v in coef_sorted["β_std"]
                            ],
                            customdata=[_coef_stmts.get(lbl, "") for lbl in coef_sorted["label"]],
                            hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>β = %{y:.3f}<extra></extra>",
                        ))
                        fig_coef.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=40, b=160),
                            xaxis=dict(
                                tickangle=-40,
                                tickfont=dict(color="#1A2B3C", size=10),
                            ),
                            yaxis=dict(
                                title=dict(text="Standardised β", font=dict(color="#1A2B3C")),
                                zeroline=True, zerolinecolor="#D6E0EA",
                                tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                            ),
                            height=420,
                        )
                        st.plotly_chart(fig_coef, use_container_width=True,
                                        key=f"oda_coef_{out_lbl}")

                        # Elimination path
                        st.markdown("##### Elimination Path — Adj R² at each step")
                        st.caption(
                            "Each point shows the model's Adjusted R² after removing the least "
                            "significant predictor at that step. A flat or rising line means those "
                            "predictors weren't earning their place. A drop signals the core has "
                            "been reached."
                        )
                        log_df = pd.DataFrame(res["elim_log"])
                        if "Predictors in model" in log_df.columns:
                            fig_elim = go.Figure(go.Scatter(
                                x=log_df["Predictors in model"],
                                y=log_df["Adj. R²"],
                                mode="lines+markers",
                                line=dict(color=PRIMARY, width=2),
                                marker=dict(color=PRIMARY, size=8),
                                hovertemplate="Predictors: %{x}<br>Adj. R² = %{y:.4f}<extra></extra>",
                            ))
                            fig_elim.update_layout(
                                font=dict(family="Inter", color="#1A2B3C"),
                                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                                margin=dict(l=10, r=10, t=10, b=10),
                                xaxis=dict(
                                    title=dict(text="Predictors in model", font=dict(color="#1A2B3C")),
                                    autorange="reversed",
                                    tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                ),
                                yaxis=dict(
                                    title=dict(text="Adjusted R²", font=dict(color="#1A2B3C")),
                                    tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                ),
                                height=300,
                            )
                            st.plotly_chart(fig_elim, use_container_width=True,
                                            key=f"oda_elim_{out_lbl}")

        with oda_lf_tab:
            st.caption(
                "Runs the same multiple regression pipeline as the Outcomes tab, but uses "
                "the latent factor scores (LV1, LV2 …) derived from the EFA in the "
                "Outcome Clusters tab as the dependent variables. All WoW themes enter the "
                "model directly — no correlation floor is applied since there is no "
                "pre-existing heatmap for latent factors. Backward elimination then removes "
                "non-significant predictors until only confirmed drivers remain."
            )

            # ── Controls row 1: WoW type + p-value ────────────────
            lf_c1, lf_c2, lf_c3 = st.columns([3, 2, 2])
            with lf_c1:
                wow_choice_lf = st.radio(
                    "WoW predictors", ["Place (P)", "Individual (I)"],
                    horizontal=True, key="lf_wow"
                )
            with lf_c3:
                p_thresh_lf = st.slider(
                    "P-value threshold", 0.00, 0.10, 0.05, 0.01, key="lf_p"
                )
                st.markdown(
                    '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                    'Recommended: 0.05</p>',
                    unsafe_allow_html=True)

            # ── Controls row 2: n_factors + factor preview ─────────
            lf_r2_left, _ = st.columns([2, 5])
            with lf_r2_left:
                lf_n_factors = st.slider(
                    "Number of latent factors", 1, 6, 3, key="lf_nfactors"
                )
            with st.expander("Scree Plot — How many hidden factors exist in your outcome data?"):
                _sem_clean_a = filtered[OUTCOME_COLS].dropna()
                if len(_sem_clean_a) < 20:
                    st.warning("Too few complete responses for factor analysis.")
                else:
                    _corr_a = np.corrcoef(_sem_clean_a.values.T)
                    _eig_a = sorted(np.linalg.eigvalsh(_corr_a).tolist(), reverse=True)
                    _nk_a = max(2, sum(1 for e in _eig_a if e > 1))
                    st.caption(
                        "Each point is one potential factor and its height shows how much shared information "
                        "it captures. Factors above the dashed line (eigenvalue > 1) are worth keeping. "
                        "Look for where the curve flattens — that bend marks your real underlying themes."
                    )
                    _fig_scree_a = go.Figure([go.Scatter(
                        x=list(range(1, len(_eig_a) + 1)), y=_eig_a,
                        mode="lines+markers",
                        line=dict(color=PRIMARY, width=2),
                        marker=dict(color=PRIMARY, size=7),
                        hovertemplate="Factor %{x}<br>Eigenvalue = %{y:.3f}<extra></extra>",
                    )])
                    _fig_scree_a.add_hline(y=1, line_dash="dash", line_color="#C0392B",
                        annotation_text="Kaiser criterion (λ = 1)",
                        annotation_font=dict(color="#C0392B", size=11))
                    _fig_scree_a.update_layout(
                        font=dict(family="Inter", color="#1A2B3C"),
                        paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                        margin=dict(l=10, r=10, t=20, b=10),
                        xaxis=dict(title=dict(text="Factor", font=dict(color="#1A2B3C")),
                                   tickvals=list(range(1, len(_eig_a) + 1)),
                                   tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2"),
                        yaxis=dict(title=dict(text="Eigenvalue", font=dict(color="#1A2B3C")),
                                   tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2"),
                        height=300,
                    )
                    st.plotly_chart(_fig_scree_a, use_container_width=True, key="a1_sem_scree")
            if "Place" in wow_choice_lf:
                lf_pred_cols = tuple(WOW_PLACE_COLS)
                lf_col_to_lbl = dict(zip(WOW_PLACE_COLS, WOW_THEMES))
            else:
                lf_pred_cols = tuple(WOW_IND_COLS)
                lf_col_to_lbl = dict(zip(WOW_IND_COLS, WOW_THEMES))
            lf_pred_labels = tuple(lf_col_to_lbl[c] for c in lf_pred_cols)

            outcome_clean_lf = filtered[OUTCOME_COLS].dropna()
            if len(outcome_clean_lf) < 20:
                st.warning("Too few complete responses for factor analysis.")
            else:
                with st.spinner("Computing factor scores and running regressions…"):
                    factor_scores_df = compute_factor_scores(
                        filtered[OUTCOME_COLS], lf_n_factors
                    )
                    lf_results = run_lf_driver_analysis(
                        filtered, factor_scores_df,
                        lf_pred_cols, lf_pred_labels, p_thresh_lf
                    )

                # ── Factor preview + WoW drivers ──────────────────
                with _timed(f"A — run_efa ({lf_n_factors} factors)"):
                    _lf_loadings, _ = run_efa(filtered[OUTCOME_COLS], lf_n_factors)
                _lf_loadings.index = OUTCOME_LABELS
                _lf_cols = [f"LV{i+1}" for i in range(lf_n_factors)]
                _lf_max = _lf_loadings.abs().max(axis=1)
                _lf_primary = _lf_loadings.abs().idxmax(axis=1).copy()
                _lf_primary[_lf_max < 0.3] = "Unassigned"
                _lf_members = {lv: [] for lv in _lf_cols}
                for _outcome, _lv in _lf_primary.items():
                    if _lv != "Unassigned":
                        _lf_members[_lv].append(
                            (_outcome, _lf_loadings.loc[_outcome, _lv])
                        )
                for _lv in _lf_members:
                    _lf_members[_lv].sort(key=lambda x: abs(x[1]), reverse=True)

                # ── Combined top WoW drivers across all latent factors ──
                _combined_betas: dict = {}
                for _lv, _res in lf_results.items():
                    if _res.get("status") == "ok":
                        for _, _row in _res["coef_df"].iterrows():
                            _lbl = _row["label"]
                            _combined_betas[_lbl] = _combined_betas.get(_lbl, 0.0) + _row["β_std"]
                _preview_cols = st.columns(lf_n_factors)
                for _col_ui, _lv in zip(_preview_cols, _lf_cols):
                    _items_html = "".join(
                        f'<p class="card-sub">· {m} ({v:.2f})</p>'
                        for m, v in _lf_members[_lv]
                    ) or '<p class="card-sub"><em>None assigned</em></p>'
                    _res = lf_results.get(_lv, {})
                    if _res.get("status") == "ok":
                        _coef = _res["coef_df"].copy()
                        _coef_sorted = _coef.reindex(
                            _coef["β_std"].abs().sort_values(ascending=False).index
                        )
                        _drivers_html = "".join(
                            f'<p class="card-sub">· {row["label"]} '
                            f'(<span style="color:{"#C0392B" if row["β_std"] >= 0 else "#2980B9"};font-weight:600">'
                            f'{row["β_std"]:+.2f}</span>)</p>'
                            for _, row in _coef_sorted.iterrows()
                        )
                    elif _res.get("status") == "fail_sig_f":
                        _drivers_html = (
                            f'<p class="card-sub"><em>Model not reliable '
                            f'(Sig. F = {_res["sig_f"]:.3f})</em></p>'
                        )
                    else:
                        _drivers_html = '<p class="card-sub"><em>No confirmed drivers</em></p>'
                    with _col_ui:
                        st.markdown(
                            f'<div class="metric-card-lg">'
                            f'<p class="card-label">{_lv}</p>'
                            f'{_items_html}'
                            f'<p class="card-label" style="margin-top:10px">'
                            f'Top WoW drivers (β)</p>'
                            f'{_drivers_html}'
                            f'</div>',
                            unsafe_allow_html=True)

                if _combined_betas:
                    _sorted_combined = sorted(_combined_betas.items(), key=lambda x: x[1], reverse=True)
                    _cb_labels = [x[0] for x in _sorted_combined]
                    _cb_scores = [x[1] for x in _sorted_combined]
                    st.markdown("#### Overall Top WoW Drivers Across All Outcome Groups (Latent Factors)")
                    st.caption(
                        "Each bar shows the sum of standardised β values for that WoW theme across all "
                        "latent factors where it was a confirmed driver, preserving sign. A consistently "
                        "positive theme scores high; one with mixed effects across factors will partially "
                        "cancel out. Red = net positive effect; blue = net negative."
                    )
                    fig_combined = go.Figure(go.Bar(
                        x=_cb_labels,
                        y=_cb_scores,
                        text=[f"{s:+.2f}" for s in _cb_scores],
                        textposition="outside",
                        textfont=dict(color="#1A2B3C", size=10),
                        marker_color=[RED if s >= 0 else PRIMARY for s in _cb_scores],
                        hovertemplate="<b>%{x}</b><br>Combined β = %{y:.3f}<extra></extra>",
                    ))
                    fig_combined.update_layout(
                        font=dict(family="Inter", color="#1A2B3C"),
                        paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                        margin=dict(l=10, r=10, t=40, b=260),
                        xaxis=dict(
                            tickangle=-40,
                            tickfont=dict(color="#1A2B3C", size=10),
                        ),
                        yaxis=dict(
                            title=dict(text="Combined β", font=dict(color="#1A2B3C")),
                            zeroline=True, zerolinecolor="#D6E0EA",
                            tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                        ),
                        height=560,
                    )
                    st.plotly_chart(fig_combined, use_container_width=True, key="lf_combined_bar")

                ok_lf = {lv: res for lv, res in lf_results.items()
                         if res["status"] == "ok"}

                # ── Priority matrix ────────────────────────────────
                st.markdown("---")
                st.markdown("#### WoW Drivers by Outcome Group (Latent Factor)")
                st.caption(
                    "Each cell shows the standardised β for a WoW theme retained in that "
                    "latent factor's final model. Blank = not a confirmed driver. "
                    "Red = positive effect; blue = negative effect."
                )

                if not ok_lf:
                    st.warning(
                        "No latent factors produced a statistically reliable model. "
                        "Try raising the p-value threshold."
                    )
                else:
                    all_lf_retained = sorted(set(
                        lf_col_to_lbl[c]
                        for res in ok_lf.values()
                        for c in res["retained"]
                    ))
                    lf_matrix_rows = {}
                    for lv, res in ok_lf.items():
                        lbl_to_beta = {
                            row["label"]: row["β_std"]
                            for _, row in res["coef_df"].iterrows()
                        }
                        lf_matrix_rows[lv] = {
                            theme: lbl_to_beta.get(theme, np.nan)
                            for theme in all_lf_retained
                        }
                    lf_priority_df = pd.DataFrame(lf_matrix_rows, index=all_lf_retained)
                    lf_priority_df_T = lf_priority_df.T

                    lf_beta_max = max(lf_priority_df.abs().max().max(), 0.01)

                    fig_lf_matrix = go.Figure(go.Heatmap(
                        z=lf_priority_df_T.values.tolist(),
                        x=list(lf_priority_df_T.columns),
                        y=list(lf_priority_df_T.index),
                        colorscale=[
                            [0.0, "#0F4C6B"],
                            [0.5, "#F7F9FC"],
                            [1.0, "#C0392B"],
                        ],
                        zmid=0,
                        zmin=-lf_beta_max,
                        zmax=lf_beta_max,
                        text=[[
                            f"{v:.2f}" if not np.isnan(v) else ""
                            for v in row
                        ] for row in lf_priority_df_T.values.tolist()],
                        texttemplate="%{text}",
                        textfont=dict(size=11, color="#1A2B3C"),
                        hoverongaps=False,
                        hovertemplate="<b>%{x}</b><br>%{y}<br>β = %{z:.3f}<extra></extra>",
                        colorbar=dict(title="β", len=0.7),
                    ))
                    fig_lf_matrix.update_layout(
                        font=dict(family="Inter", color="#1A2B3C"),
                        paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                        margin=dict(l=10, r=10, t=10, b=160),
                        xaxis=dict(
                            tickfont=dict(color="#1A2B3C", size=11),
                            tickangle=-40,
                            side="bottom",
                        ),
                        yaxis=dict(tickfont=dict(color="#1A2B3C", size=11)),
                        height=max(300, 36 * len(ok_lf)),
                    )
                    st.plotly_chart(fig_lf_matrix, use_container_width=True,
                                    key="lf_matrix")

                # ── Per-factor detail ──────────────────────────────
                st.markdown("---")
                st.markdown("#### Breakdown per Outcome Group")

                for lv in factor_scores_df.columns:
                    res = lf_results.get(lv, {"status": "all_eliminated"})
                    icon = "✓" if res["status"] == "ok" else "✗"
                    with st.expander(f"{icon}  {lv}"):
                        if res["status"] == "fail_sig_f":
                            st.error(
                                f"**Model not statistically reliable** — "
                                f"Significance F = {res['sig_f']:.4f} (> 0.05)."
                            )
                            mc1, mc2, mc3 = st.columns(3)
                            for col_ui, label, value in [
                                (mc1, "Respondents", f"n = {res['n']:,}"),
                                (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                                (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                            ]:
                                with col_ui:
                                    st.markdown(
                                        f'<div class="metric-card"><p class="card-label">{label}</p>'
                                        f'<p class="card-value">{value}</p></div>',
                                        unsafe_allow_html=True)
                        elif res["status"] == "all_eliminated":
                            st.warning("All predictors were eliminated. Try raising the p-value threshold.")
                        else:
                            coef_df = res["coef_df"]
                            mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                            for col_ui, label, value in [
                                (mc1, "Respondents",          f"n = {res['n']:,}"),
                                (mc2, "Significance F",       f"{res['sig_f']:.4f}"),
                                (mc3, "Adj R² (full model)",  f"{res['adj_r2_full']:.3f}"),
                                (mc4, "Adj R² (final model)", f"{res['adj_r2_final']:.3f}"),
                                (mc5, "Confirmed drivers",    str(len(res["retained"]))),
                            ]:
                                with col_ui:
                                    st.markdown(
                                        f'<div class="metric-card"><p class="card-label">{label}</p>'
                                        f'<p class="card-value">{value}</p></div>',
                                        unsafe_allow_html=True)

                            st.markdown("##### Confirmed Drivers — Standardised β Coefficients")
                            coef_sorted = coef_df.reindex(
                                coef_df["β_std"].abs().sort_values(ascending=False).index
                            )
                            _lf_coef_stmts = WOW_PLACE_STATEMENTS if "Place" in wow_choice_lf else WOW_IND_STATEMENTS
                            fig_lf_coef = go.Figure(go.Bar(
                                x=coef_sorted["label"],
                                y=coef_sorted["β_std"],
                                text=coef_sorted["text"],
                                textposition="outside",
                                textfont=dict(color="#1A2B3C", size=10),
                                marker_color=[
                                    RED if v >= 0 else PRIMARY
                                    for v in coef_sorted["β_std"]
                                ],
                                customdata=[_lf_coef_stmts.get(lbl, "") for lbl in coef_sorted["label"]],
                                hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>β = %{y:.3f}<extra></extra>",
                            ))
                            fig_lf_coef.update_layout(
                                font=dict(family="Inter", color="#1A2B3C"),
                                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                                margin=dict(l=10, r=10, t=40, b=160),
                                xaxis=dict(
                                    tickangle=-40,
                                    tickfont=dict(color="#1A2B3C", size=10),
                                ),
                                yaxis=dict(
                                    title=dict(text="Standardised β", font=dict(color="#1A2B3C")),
                                    zeroline=True, zerolinecolor="#D6E0EA",
                                    tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                ),
                                height=420,
                            )
                            st.plotly_chart(fig_lf_coef, use_container_width=True,
                                            key=f"lf_coef_{lv}")


        with heatmaps_tab:
                h1, h2, h3 = st.tabs([
                    "WoW × Outcomes",
                    "WoW × WoW",
                    "Outcomes × Outcomes",
                ])

                # ── H1: Ways of Working × Outcomes ───────────────────
                with h1:
                    a1_place, a1_ind = st.tabs(["Place (P)", "Individual (I)"])
                    with a1_place:
                        st.markdown("#### Correlational Heatmap: Ways of Working (Place) × Employee Experience")
                        st.caption(
                            "Each cell shows the Spearman correlation (r) between a Place-level Ways of Working "
                            "theme (rows) and an employee experience outcome (columns). Place themes reflect the "
                            "shared working environment — norms and behaviours that feel consistent across the team "
                            "or directorate. Darker blue = stronger positive relationship; darker red = stronger "
                            "negative. Values above ~0.2 are worth noting; above ~0.4 indicate a meaningful pattern. "
                            "Read across a row to see which outcomes a given working norm connects to most strongly."
                        )
                        mat = spearman_matrix(filtered, WOW_PLACE_COLS, OUTCOME_COLS)
                        mat.index   = WOW_THEMES
                        mat.columns = OUTCOME_LABELS
                        render_heatmap_cards(n_total, mat, TOTAL_HEADCOUNT)
                        st.plotly_chart(make_heatmap(mat.T, OUTCOME_LABELS, WOW_THEMES, y_statements=OUTCOME_STATEMENTS, x_statements=WOW_PLACE_STATEMENTS),
                                        use_container_width=True, key="a1_place")
                    with a1_ind:
                        st.markdown("#### Correlational Heatmap: Ways of Working (Individual) × Employee Experience")
                        st.caption(
                            "Each cell shows the Spearman correlation (r) between an Individual-level Ways of Working "
                            "theme (rows) and an employee experience outcome (columns). Individual themes reflect how "
                            "each person personally operates — their own habits and approach, regardless of what the "
                            "team around them does. Darker blue = stronger positive relationship; darker red = stronger "
                            "negative. Compare this map with the Place heatmap to see whether environmental or personal "
                            "behaviours are more strongly linked to each outcome."
                        )
                        mat = spearman_matrix(filtered, WOW_IND_COLS, OUTCOME_COLS)
                        mat.index   = WOW_THEMES
                        mat.columns = OUTCOME_LABELS
                        render_heatmap_cards(n_total, mat, TOTAL_HEADCOUNT)
                        st.plotly_chart(make_heatmap(mat.T, OUTCOME_LABELS, WOW_THEMES, y_statements=OUTCOME_STATEMENTS, x_statements=WOW_IND_STATEMENTS),
                                        use_container_width=True, key="a1_ind")

                # ── H2: Ways of Working × Ways of Working ────────────
                with h2:
                    a2_place, a2_ind = st.tabs(["Place (P)", "Individual (I)"])
                    with a2_place:
                        st.markdown("#### Correlational Heatmap: Ways of Working (Place) × Ways of Working (Place)")
                        st.caption(
                            "Each cell shows the Spearman correlation between two Place-level Ways of Working themes. "
                            "Strong positive correlations (dark blue) mean those norms tend to co-exist — groups that "
                            "display one tend to display the other too. Strong negative correlations (dark red) indicate "
                            "genuinely opposing orientations. Use this to identify clusters of related working norms "
                            "and to flag themes that may be measuring the same underlying dimension before regression modelling."
                        )
                        mat = make_writable_matrix(spearman_matrix(filtered, WOW_PLACE_COLS, WOW_PLACE_COLS))
                        mat.index = mat.columns = WOW_THEMES
                        fill_diagonal_with_nan(mat)
                        render_heatmap_cards(n_total, mat, TOTAL_HEADCOUNT)
                        st.plotly_chart(make_heatmap(mat, WOW_THEMES, WOW_THEMES, y_statements=WOW_PLACE_STATEMENTS, x_statements=WOW_PLACE_STATEMENTS),
                                        use_container_width=True, key="a2_place")
                    with a2_ind:
                        st.markdown("#### Correlational Heatmap: Ways of Working (Individual) × Ways of Working (Individual)")
                        st.caption(
                            "Each cell shows the Spearman correlation between two Individual-level Ways of Working themes. "
                            "Strong positive correlations mean those personal behaviours tend to appear together in the same "
                            "people. Use this alongside the Place heatmap to see whether individual habits cluster differently "
                            "from the team-level norms — a divergence can point to tension between personal style and the "
                            "shared working environment."
                        )
                        mat = make_writable_matrix(spearman_matrix(filtered, WOW_IND_COLS, WOW_IND_COLS))
                        mat.index = mat.columns = WOW_THEMES
                        fill_diagonal_with_nan(mat)
                        render_heatmap_cards(n_total, mat, TOTAL_HEADCOUNT)
                        st.plotly_chart(make_heatmap(mat, WOW_THEMES, WOW_THEMES, y_statements=WOW_IND_STATEMENTS, x_statements=WOW_IND_STATEMENTS),
                                        use_container_width=True, key="a2_ind")

                # ── H3: Outcomes × Outcomes ───────────────────────────
                with h3:
                    st.markdown("#### Correlational Heatmap: Employee Experience × Employee Experience")
                    st.caption(
                        "Each cell shows the Spearman correlation between two employee experience outcomes. "
                        "Strong positive correlations indicate outcomes that tend to rise and fall together — "
                        "suggesting they may be capturing the same underlying dimension of experience. "
                        "Use this to understand how interconnected the outcome measures are before interpreting "
                        "individual scores in isolation, and to inform how outcomes might be grouped for further analysis."
                    )
                    mat = make_writable_matrix(spearman_matrix(filtered, OUTCOME_COLS, OUTCOME_COLS))
                    mat.index = mat.columns = OUTCOME_LABELS
                    fill_diagonal_with_nan(mat)
                    render_heatmap_cards(n_total, mat, TOTAL_HEADCOUNT)
                    st.plotly_chart(make_heatmap(mat, OUTCOME_LABELS, OUTCOME_LABELS, y_statements=OUTCOME_STATEMENTS, x_statements=OUTCOME_STATEMENTS),
                                    use_container_width=True, key="a3")

        # ── Descriptive Analysis group ────────────────────────
        with _timed("A — desc_group render"), desc_group:
            a4, a5 = st.tabs([
                "A2.1: Ways of Working",
                "A2.2: Employee Experience",
            ])

            # ── A4: Ways of Working descriptive table ─────────────
            with a4:
                st.caption("Average scores for each Way of Working theme, split by Place and Individual, for the selected organisational level.")
                _a4_q9 = sorted(q9_levels, key=lambda x: next(
                    (i for i, o in enumerate(Q9_ORDER) if o.lower() == x.lower()), len(Q9_ORDER)))
                st.markdown(
                    '<p style="font-size:13px;font-weight:600;color:#5A7080;'
                    'text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px">'
                    'Select org level</p>', unsafe_allow_html=True)
                _a4_prev = st.session_state.get("a4_q9_chart", "Overall")
                _a4_prev_df = filtered if _a4_prev == "Overall" else filtered[filtered["Q9"] == _a4_prev]
                _a4_prev_pct = f" ({len(_a4_prev_df) / TOTAL_HEADCOUNT:.0%} of total)"
                st.caption(f"{len(_a4_prev_df):,} respondents in **{_a4_prev}**{_a4_prev_pct}")
                sel_a4 = st.radio("Select org level", ["Overall"] + _a4_q9, horizontal=True,
                                  label_visibility="collapsed", key="a4_q9_chart")
                chart_df = filtered if sel_a4 == "Overall" else filtered[filtered["Q9"] == sel_a4]
                st.markdown("#### Ways of Working — Place vs. Individual per Org Level")
                st.plotly_chart(make_wow_bar_chart(chart_df), use_container_width=True, key="a4_bar")

                st.markdown("---")
                st.markdown("#### Ways of Working — Comparison Across Org Level")
                st.caption(
                    "For each Way of Working theme, this chart shows the average score broken down by "
                    "organisational level, so you can see where scores vary most across the hierarchy."
                )
                _a4_pi = st.radio("WoW type", ["Place (P)", "Individual (I)"], horizontal=True, key="a4_pi")
                _a4_wow_cols   = WOW_PLACE_COLS if "Place" in _a4_pi else WOW_IND_COLS
                _a4_wow_labels = WOW_THEMES
                _a4_levels = ["Overall"] + _a4_q9
                _a4_colours = ["#1A2B3C", "#0F4C6B", "#2E7096", "#5A9BB5", "#8DC0D4"]
                fig_a4_cmp = go.Figure()
                for lvl, col in zip(_a4_levels, _a4_colours):
                    lvl_df = filtered if lvl == "Overall" else filtered[filtered["Q9"] == lvl]
                    fig_a4_cmp.add_trace(go.Bar(
                        name=lvl,
                        x=_a4_wow_labels,
                        y=[lvl_df[c].mean() for c in _a4_wow_cols],
                        marker_color=col,
                        hovertemplate=f"{lvl}<br>%{{x}}<br>Score: %{{y:.2f}}<extra></extra>",
                    ))
                fig_a4_cmp.update_layout(
                    barmode="group",
                    font=dict(family="Inter", color="#1A2B3C"),
                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                    margin=dict(l=10, r=10, t=10, b=180),
                    xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
                    yaxis=dict(
                        title=dict(text="Average Score", font=dict(color="#1A2B3C", size=12)),
                        range=[-0.2, 4.5],
                        tickfont=dict(size=11, color="#1A2B3C"),
                        gridcolor="#E8EEF2",
                    ),
                    legend=dict(
                        orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                        font=dict(color="#1A2B3C", size=12),
                    ),
                    height=560,
                )
                st.plotly_chart(fig_a4_cmp, use_container_width=True, key="a4_cmp_bar")

            # ── A5: Sentiment Outcomes descriptive table ──────────
            with a5:
                st.caption("Average scores for each employee experience outcome for the selected organisational level, compared against the council-wide average.")
                _a5_q9 = sorted(q9_levels, key=lambda x: next(
                    (i for i, o in enumerate(Q9_ORDER) if o.lower() == x.lower()), len(Q9_ORDER)))
                st.markdown(
                    '<p style="font-size:13px;font-weight:600;color:#5A7080;'
                    'text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px">'
                    'Select org level</p>', unsafe_allow_html=True)
                _a5_prev = st.session_state.get("a5_q9_chart", "Overall")
                _a5_prev_df = filtered if _a5_prev == "Overall" else filtered[filtered["Q9"] == _a5_prev]
                _a5_prev_pct = f" ({len(_a5_prev_df) / TOTAL_HEADCOUNT:.0%} of total)"
                st.caption(f"{len(_a5_prev_df):,} respondents in **{_a5_prev}**{_a5_prev_pct}")
                sel_a5 = st.radio("Select org level", ["Overall"] + _a5_q9, horizontal=True,
                                  label_visibility="collapsed", key="a5_q9_chart")
                chart_df = filtered if sel_a5 == "Overall" else filtered[filtered["Q9"] == sel_a5]
                overall_df = None if sel_a5 == "Overall" else filtered
                st.markdown("#### Employee Experience — Average Scores by Organisational Level")
                st.plotly_chart(make_outcome_bar_chart(chart_df, overall_df), use_container_width=True,
                                key=f"a5_bar_{sel_a5}")

                st.markdown("---")
                st.markdown("#### Employee Experience — Comparison Across Org Level")
                st.caption(
                    "For each employee experience outcome, this chart shows the average score broken "
                    "down by organisational level, so you can see where scores vary most across the hierarchy."
                )
                _a5_levels = ["Overall"] + _a5_q9
                _a5_colours = ["#1A2B3C", "#0F4C6B", "#2E7096", "#5A9BB5", "#8DC0D4"]
                fig_a5_cmp = go.Figure()
                for lvl, col in zip(_a5_levels, _a5_colours):
                    lvl_df = filtered if lvl == "Overall" else filtered[filtered["Q9"] == lvl]
                    fig_a5_cmp.add_trace(go.Bar(
                        name=lvl,
                        x=OUTCOME_LABELS,
                        y=[lvl_df[c].mean() for c in OUTCOME_COLS],
                        marker_color=col,
                        hovertemplate=f"{lvl}<br>%{{x}}<br>Score: %{{y:.2f}}<extra></extra>",
                    ))
                fig_a5_cmp.update_layout(
                    barmode="group",
                    font=dict(family="Inter", color="#1A2B3C"),
                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                    margin=dict(l=10, r=10, t=10, b=180),
                    xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
                    yaxis=dict(
                        title=dict(text="Average Score", font=dict(color="#1A2B3C", size=12)),
                        range=[-0.2, 4.5],
                        tickfont=dict(size=11, color="#1A2B3C"),
                        gridcolor="#E8EEF2",
                    ),
                    legend=dict(
                        orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                        font=dict(color="#1A2B3C", size=12),
                    ),
                    height=560,
                )
                st.plotly_chart(fig_a5_cmp, use_container_width=True, key="a5_cmp_bar")

            # ── A6: By Q9 Organisational Level ───────────────────


    # ═══════════════════════════════════════════════════════════════════════════════════
    # SECTION B
    # ═══════════════════════════════════════════════════════════════════════════════════
    with sec_b:
        if not directorates:
            st.warning("No directorate data found in this file.")
            st.stop()

        b_corr_group, b_desc_group = st.tabs(["B1: Drivers Analysis", "B2: Descriptive Analysis"])

        # ── Correlation Analysis group ────────────────────────
        with _timed("B — b_corr_group render"), b_corr_group:
            # ── ODA tabs at the top level ─────────────────────────
            oda_outcomes_tab, oda_lf_tab, heatmaps_tab = st.tabs([
                "B1.1: Individual Outcomes",
                "B1.2: Grouped Outcomes",
                "B1.3: Correlational Heatmaps",
            ])

            with oda_outcomes_tab:
                b11_main_tab, b112_dir_tab, b111_tab, b112_tab = st.tabs([
                    "B1.1.1: Top Drivers by Org Breakdown",
                    "B1.1.2: Top Drivers Summary — Directorates",
                    "B1.1.3: Top Drivers Summary — Service Areas",
                    "B1.1.4: Top Drivers Summary — Services",
                ])

                with b112_dir_tab:
                    st.caption(
                        "For each directorate, shows the confirmed WoW drivers with a combined β ≥ 1 (positive) "
                        "or ≤ -0.25 (negative) across all outcome models. "
                        "Uses default settings: correlation floor 0.20, p-value threshold 0.05, Place WoW predictors."
                    )
                    _b_dir_pos, _b_dir_neg = {}, {}
                    _b_dir_pos_freq, _b_dir_neg_freq = {}, {}
                    _b_dir_pos_detail = {}
                    _b_dir_neg_detail = {}
                    _b_dir_total = 0
                    with st.spinner("Computing drivers across all directorates…"):
                        for _dir in directorates:
                            _dir_df2 = filtered[filtered["Q1"] == _dir]
                            if len(_dir_df2) < 5:
                                continue
                            _b_dir_total += 1
                            _dir_res = run_driver_analysis_batch(
                                _dir_df2,
                                tuple(OUTCOME_COLS), tuple(OUTCOME_LABELS),
                                tuple(WOW_PLACE_COLS), tuple(WOW_THEMES),
                                0.20, 0.05,
                            )
                            _dir_comb = {}
                            for _r in _dir_res.values():
                                if _r.get("status") == "ok":
                                    for _, _row in _r["coef_df"].iterrows():
                                        _lbl = _row["label"]
                                        _dir_comb[_lbl] = _dir_comb.get(_lbl, 0.0) + _row["β_std"]
                            _dpos = sorted([(l, b) for l, b in _dir_comb.items() if b >= 1],
                                           key=lambda x: x[1], reverse=True)
                            _dneg = sorted([(l, b) for l, b in _dir_comb.items() if b <= -0.25],
                                           key=lambda x: x[1])
                            if _dpos:
                                _b_dir_pos[_dir] = [f"{l} ({b:+.2f})" for l, b in _dpos]
                                for l, b in _dpos:
                                    _b_dir_pos_freq[l] = _b_dir_pos_freq.get(l, 0) + 1
                                    if l not in _b_dir_pos_detail:
                                        _b_dir_pos_detail[l] = []
                                    _wow_i = WOW_THEMES.index(l) if l in WOW_THEMES else -1
                                    _ddelta = (_dir_df2[WOW_PLACE_COLS[_wow_i]].mean() - _dir_df2[WOW_IND_COLS[_wow_i]].mean()) if _wow_i >= 0 else 0.0
                                    _b_dir_pos_detail[l].append((_dir, b, _ddelta))
                            if _dneg:
                                _b_dir_neg[_dir] = [f"{l} ({b:+.2f})" for l, b in _dneg]
                                for l, b in _dneg:
                                    _b_dir_neg_freq[l] = _b_dir_neg_freq.get(l, 0) + 1
                                    if l not in _b_dir_neg_detail:
                                        _b_dir_neg_detail[l] = []
                                    _wow_i = WOW_THEMES.index(l) if l in WOW_THEMES else -1
                                    _ddelta = (_dir_df2[WOW_PLACE_COLS[_wow_i]].mean() - _dir_df2[WOW_IND_COLS[_wow_i]].mean()) if _wow_i >= 0 else 0.0
                                    _b_dir_neg_detail[l].append((_dir, b, _ddelta))
                    # ── Shortlisted drivers (beta + delta cut-offs) ────────────
                    def _shortlist_table_dir(detail_dict, label, is_pos):
                        shortlist = {}
                        for theme, entries in detail_dict.items():
                            for unit, b, delta in entries:
                                if abs(delta) > 0.5:
                                    if unit not in shortlist:
                                        shortlist[unit] = []
                                    shortlist[unit].append((theme, b, delta))
                        if not shortlist:
                            st.info(f"No {label} drivers meet both beta and delta (>0.5) cut-offs.")
                            return
                        for unit in shortlist:
                            shortlist[unit].sort(key=lambda x: x[1], reverse=is_pos)
                        max_rows = max(len(v) for v in shortlist.values())
                        table = {}
                        for unit, drivers in shortlist.items():
                            cells = []
                            for theme, b, delta in drivers:
                                cells.append(f"{theme} (β={b:+.2f}, Δ={delta:+.2f})")
                            cells += [""] * (max_rows - len(cells))
                            table[unit] = cells
                        df_short = pd.DataFrame(table)
                        df_short.index = [f"#{i+1}" for i in range(max_rows)]
                        _STANDARD = {"Recognise Contributions", "Challenge Decisions"}
                        def _hl_cell(val):
                            if not val:
                                return ""
                            theme_name = val.split(" (β=")[0]
                            if theme_name not in _STANDARD:
                                return "background-color: #FFF3CD; color: #1A2B3C"
                            return ""
                        st.dataframe(df_short.style.map(_hl_cell), use_container_width=True)
                    st.markdown("#### Shortlisted Drivers — Directorates (β + Δ > 0.5)")
                    st.caption("Positive shortlist: drivers with combined β ≥ 1 AND |Δ| > 0.5. Negative: β ≤ -0.25 AND |Δ| > 0.5.")
                    if _b_dir_pos_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Positive drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_dir(_b_dir_pos_detail, "positive", True)
                    if _b_dir_neg_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Negative drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_dir(_b_dir_neg_detail, "negative", False)
                    if _b_dir_pos:
                        st.markdown("#### Where Each Positive Driver Appears")
                        for _theme, _count in sorted(_b_dir_pos_freq.items(), key=lambda x: x[1], reverse=True):
                            with st.expander(f"{_theme} — {_count} directorate{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b_dir_total} directorates analysed</p>', unsafe_allow_html=True)
                                for _dv, _b, _dd in sorted(_b_dir_pos_detail.get(_theme, []), key=lambda x: x[1], reverse=True):
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_dv}</strong> '
                                        f'(β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Positive Drivers by Directorate (combined β ≥ 1)")
                        _mpd = max(len(v) for v in _b_dir_pos.values())
                        _pos_dfd = pd.DataFrame({k: v + [""] * (_mpd - len(v)) for k, v in _b_dir_pos.items()})
                        _pos_dfd.index = [f"#{i+1}" for i in range(_mpd)]
                        st.dataframe(_pos_dfd, use_container_width=True)
                    else:
                        st.info("No positive drivers with combined β ≥ 1 found at directorate level.")

                    if _b_dir_neg:
                        st.markdown("#### Where Each Negative Driver Appears")
                        for _theme, _count in sorted(_b_dir_neg_freq.items(), key=lambda x: x[1], reverse=True):
                            with st.expander(f"{_theme} — {_count} directorate{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b_dir_total} directorates analysed</p>', unsafe_allow_html=True)
                                for _dv, _b, _dd in sorted(_b_dir_neg_detail.get(_theme, []), key=lambda x: x[1]):
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_dv}</strong> (β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Negative Drivers by Directorate (combined β ≤ -0.25)")
                        _mnd = max(len(v) for v in _b_dir_neg.values())
                        _neg_dfd = pd.DataFrame({k: v + [""] * (_mnd - len(v)) for k, v in _b_dir_neg.items()})
                        _neg_dfd.index = [f"#{i+1}" for i in range(_mnd)]
                        st.dataframe(_neg_dfd, use_container_width=True)
                    else:
                        st.info("No negative drivers with combined β ≤ -0.25 found at directorate level.")

                with b111_tab:
                    st.caption(
                        "For each service area across all directorates, shows the confirmed WoW drivers "
                        "with a combined β ≥ 1 (positive) or ≤ -0.25 (negative) across all outcome models. "
                        "Uses default settings: correlation floor 0.20, p-value threshold 0.05, Place WoW predictors."
                    )
                    _DIR_ABBREV = {
                        "Adult Services and Housing": "ASH",
                        "Chief Exec's Office (including Executive and Service Directors)": "CEO",
                        "Children, Families and Education": "CFE",
                        "Community, Place and Economy": "CPE",
                        "Finance and Procurement": "Finance",
                        "Resources, Strategy and Transformation": "RST",
                    }
                    _b111_svc_map = []
                    for _dir in directorates:
                        _dir_df = filtered[filtered["Q1"] == _dir]
                        _abbrev = _DIR_ABBREV.get(_dir, _dir[:3])
                        for _sg in get_filter_options(_dir_df, "svc_group"):
                            _display_sg = f"{_sg} ({_abbrev})" if _sg == "N/A — go to service" else _sg
                            _b111_svc_map.append((_dir, _display_sg, _dir_df[_dir_df["svc_group"] == _sg]))
                    _b111_pred_cols = tuple(WOW_PLACE_COLS)
                    _b111_pred_labels = tuple(WOW_THEMES)
                    _b111_pos, _b111_neg = {}, {}
                    _b111_pos_freq, _b111_neg_freq = {}, {}
                    _b111_pos_detail = {}  # {theme: [(service_area, beta), ...]}
                    _b111_neg_detail = {}
                    _b111_total = sum(1 for _, _, _sdf in _b111_svc_map if len(_sdf) >= 5)
                    with st.spinner("Computing drivers across all service areas…"):
                        for _dir, _sg, _sg_df in _b111_svc_map:
                            if len(_sg_df) < 5:
                                continue
                            _sg_res = run_driver_analysis_batch(
                                _sg_df,
                                tuple(OUTCOME_COLS), tuple(OUTCOME_LABELS),
                                _b111_pred_cols, _b111_pred_labels,
                                0.20, 0.05,
                            )
                            _sg_comb = {}
                            for _r in _sg_res.values():
                                if _r.get("status") == "ok":
                                    for _, _row in _r["coef_df"].iterrows():
                                        _lbl = _row["label"]
                                        _sg_comb[_lbl] = _sg_comb.get(_lbl, 0.0) + _row["β_std"]
                            _pos = sorted([(l, b) for l, b in _sg_comb.items() if b >= 1],
                                          key=lambda x: x[1], reverse=True)
                            _neg = sorted([(l, b) for l, b in _sg_comb.items() if b <= -0.25],
                                          key=lambda x: x[1])
                            if _pos:
                                _b111_pos[_sg] = [f"{l} ({b:+.2f})" for l, b in _pos]
                                for l, b in _pos:
                                    _b111_pos_freq[l] = _b111_pos_freq.get(l, 0) + 1
                                    if l not in _b111_pos_detail:
                                        _b111_pos_detail[l] = []
                                    _wow_i = WOW_THEMES.index(l) if l in WOW_THEMES else -1
                                    _sadelta = (_sg_df[WOW_PLACE_COLS[_wow_i]].mean() - _sg_df[WOW_IND_COLS[_wow_i]].mean()) if _wow_i >= 0 else 0.0
                                    _b111_pos_detail[l].append((_sg, b, _sadelta))
                            if _neg:
                                _b111_neg[_sg] = [f"{l} ({b:+.2f})" for l, b in _neg]
                                for l, _ in _neg:
                                    _b111_neg_freq[l] = _b111_neg_freq.get(l, 0) + 1
                    # ── Shortlisted drivers (beta + delta cut-offs) ────────────
                    def _shortlist_table_sa(detail_dict, label, is_pos):
                        shortlist = {}
                        for theme, entries in detail_dict.items():
                            for unit, b, delta in entries:
                                if abs(delta) > 0.5:
                                    if unit not in shortlist:
                                        shortlist[unit] = []
                                    shortlist[unit].append((theme, b, delta))
                        if not shortlist:
                            st.info(f"No {label} drivers meet both beta and delta (>0.5) cut-offs.")
                            return
                        for unit in shortlist:
                            shortlist[unit].sort(key=lambda x: x[1], reverse=is_pos)
                        max_rows = max(len(v) for v in shortlist.values())
                        table = {}
                        for unit, drivers in shortlist.items():
                            cells = []
                            for theme, b, delta in drivers:
                                cells.append(f"{theme} (β={b:+.2f}, Δ={delta:+.2f})")
                            cells += [""] * (max_rows - len(cells))
                            table[unit] = cells
                        df_short = pd.DataFrame(table)
                        df_short.index = [f"#{i+1}" for i in range(max_rows)]
                        _STANDARD = {"Recognise Contributions", "Challenge Decisions"}
                        def _hl_cell(val):
                            if not val:
                                return ""
                            theme_name = val.split(" (β=")[0]
                            if theme_name not in _STANDARD:
                                return "background-color: #FFF3CD; color: #1A2B3C"
                            return ""
                        st.dataframe(df_short.style.map(_hl_cell), use_container_width=True)
                    st.markdown("#### Shortlisted Drivers — Service Areas (β + Δ > 0.5)")
                    st.caption("Positive shortlist: combined β ≥ 1 AND |Δ| > 0.5. Negative: β ≤ -0.25 AND |Δ| > 0.5.")
                    if _b111_pos_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Positive drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_sa(_b111_pos_detail, "positive", True)
                    if _b111_neg_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Negative drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_sa(_b111_neg_detail, "negative", False)

                    def _freq_summary_html(freq_dict, label):
                        if not freq_dict:
                            return ""
                        ranked = sorted(freq_dict.items(), key=lambda x: x[1], reverse=True)
                        items = "".join(
                            f'<p class="card-sub">· {theme} — appears in <strong>{count}</strong> service area{"s" if count != 1 else ""}</p>'
                            for theme, count in ranked
                        )
                        return (
                            f'<div class="metric-card-lg" style="margin-bottom:16px">'
                            f'<p class="card-label">Most frequently occurring {label} drivers</p>'
                            f'{items}</div>'
                        )

                    if _b111_pos:
                        st.markdown("#### Where Each Positive Driver Appears")
                        _ranked_pos = sorted(_b111_pos_freq.items(), key=lambda x: x[1], reverse=True)
                        for _theme, _count in _ranked_pos:
                            with st.expander(f"{_theme} — {_count} service area{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b111_total} service areas analysed</p>', unsafe_allow_html=True)
                                _svc_betas = sorted(_b111_pos_detail.get(_theme, []), key=lambda x: x[1], reverse=True)
                                for _sa, _b, _dd in _svc_betas:
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_sa}</strong> '
                                        f'(β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Positive Drivers by Service Area (combined β ≥ 1)")
                        _mp = max(len(v) for v in _b111_pos.values())
                        _pos_df = pd.DataFrame(
                            {k: v + [""] * (_mp - len(v)) for k, v in _b111_pos.items()}
                        )
                        _pos_df.index = [f"#{i+1}" for i in range(_mp)]
                        st.dataframe(_pos_df, use_container_width=True)
                    else:
                        st.info("No positive drivers with combined β ≥ 1 found.")

                    if _b111_neg:
                        st.markdown("#### Where Each Negative Driver Appears")
                        for _theme, _count in sorted(_b111_neg_freq.items(), key=lambda x: x[1], reverse=True):
                            with st.expander(f"{_theme} — {_count} service area{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b111_total} service areas analysed</p>', unsafe_allow_html=True)
                                for _sa, _b, _dd in sorted(_b111_neg_detail.get(_theme, []), key=lambda x: x[1]):
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_sa}</strong> (β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Negative Drivers by Service Area (combined β ≤ -0.25)")
                        _mn = max(len(v) for v in _b111_neg.values())
                        _neg_df = pd.DataFrame(
                            {k: v + [""] * (_mn - len(v)) for k, v in _b111_neg.items()}
                        )
                        _neg_df.index = [f"#{i+1}" for i in range(_mn)]
                        st.dataframe(_neg_df, use_container_width=True)
                    else:
                        st.info("No negative drivers with combined β ≤ -0.25 found.")

                with b112_tab:
                    st.caption(
                        "For each individual service across all directorates, shows the confirmed WoW drivers "
                        "with a combined β ≥ 1 (positive) or ≤ -0.25 (negative) across all outcome models. "
                        "Uses default settings: correlation floor 0.20, p-value threshold 0.05, Place WoW predictors."
                    )
                    _b112_svc_map = []
                    for _dir in directorates:
                        _dir_df = filtered[filtered["Q1"] == _dir]
                        _abbrev = _DIR_ABBREV.get(_dir, _dir[:3])
                        for _sg in get_filter_options(_dir_df, "svc_group"):
                            _sg_df2 = _dir_df[_dir_df["svc_group"] == _sg]
                            for _svc in get_filter_options(_sg_df2, "svc_name"):
                                _svc_df = _sg_df2[_sg_df2["svc_name"] == _svc]
                                _display_svc = f"{_svc} ({_abbrev})"
                                _b112_svc_map.append((_dir, _display_svc, _svc_df))
                    _b112_pos, _b112_neg = {}, {}
                    _b112_pos_freq, _b112_neg_freq = {}, {}
                    _b112_pos_detail = {}
                    _b112_neg_detail = {}
                    _b112_total = sum(1 for _, _, _sdf in _b112_svc_map if len(_sdf) >= 5)
                    with st.spinner("Computing drivers across all services…"):
                        for _dir, _dsvc, _svc_df in _b112_svc_map:
                            if len(_svc_df) < 5:
                                continue
                            _svc_res = run_driver_analysis_batch(
                                _svc_df,
                                tuple(OUTCOME_COLS), tuple(OUTCOME_LABELS),
                                tuple(WOW_PLACE_COLS), tuple(WOW_THEMES),
                                0.20, 0.05,
                            )
                            _svc_comb = {}
                            for _r in _svc_res.values():
                                if _r.get("status") == "ok":
                                    for _, _row in _r["coef_df"].iterrows():
                                        _lbl = _row["label"]
                                        _svc_comb[_lbl] = _svc_comb.get(_lbl, 0.0) + _row["β_std"]
                            _pos2 = sorted([(l, b) for l, b in _svc_comb.items() if b >= 1],
                                           key=lambda x: x[1], reverse=True)
                            _neg2 = sorted([(l, b) for l, b in _svc_comb.items() if b <= -0.25],
                                           key=lambda x: x[1])
                            if _pos2:
                                _b112_pos[_dsvc] = [f"{l} ({b:+.2f})" for l, b in _pos2]
                                for l, b in _pos2:
                                    _b112_pos_freq[l] = _b112_pos_freq.get(l, 0) + 1
                                    if l not in _b112_pos_detail:
                                        _b112_pos_detail[l] = []
                                    _wow_i = WOW_THEMES.index(l) if l in WOW_THEMES else -1
                                    _svcdelta = (_svc_df[WOW_PLACE_COLS[_wow_i]].mean() - _svc_df[WOW_IND_COLS[_wow_i]].mean()) if _wow_i >= 0 else 0.0
                                    _b112_pos_detail[l].append((_dsvc, b, _svcdelta))
                            if _neg2:
                                _b112_neg[_dsvc] = [f"{l} ({b:+.2f})" for l, b in _neg2]
                                for l, b in _neg2:
                                    _b112_neg_freq[l] = _b112_neg_freq.get(l, 0) + 1
                                    if l not in _b112_neg_detail:
                                        _b112_neg_detail[l] = []
                                    _wow_i = WOW_THEMES.index(l) if l in WOW_THEMES else -1
                                    _svcdelta = (_svc_df[WOW_PLACE_COLS[_wow_i]].mean() - _svc_df[WOW_IND_COLS[_wow_i]].mean()) if _wow_i >= 0 else 0.0
                                    _b112_neg_detail[l].append((_dsvc, b, _svcdelta))
                    # ── Shortlisted drivers (beta + delta cut-offs) ────────────
                    def _shortlist_table_svc(detail_dict, label, is_pos):
                        shortlist = {}
                        for theme, entries in detail_dict.items():
                            for unit, b, delta in entries:
                                if abs(delta) > 0.5:
                                    if unit not in shortlist:
                                        shortlist[unit] = []
                                    shortlist[unit].append((theme, b, delta))
                        if not shortlist:
                            st.info(f"No {label} drivers meet both beta and delta (>0.5) cut-offs.")
                            return
                        for unit in shortlist:
                            shortlist[unit].sort(key=lambda x: x[1], reverse=is_pos)
                        max_rows = max(len(v) for v in shortlist.values())
                        table = {}
                        for unit, drivers in shortlist.items():
                            cells = []
                            for theme, b, delta in drivers:
                                cells.append(f"{theme} (β={b:+.2f}, Δ={delta:+.2f})")
                            cells += [""] * (max_rows - len(cells))
                            table[unit] = cells
                        df_short = pd.DataFrame(table)
                        df_short.index = [f"#{i+1}" for i in range(max_rows)]
                        _STANDARD = {"Recognise Contributions", "Challenge Decisions"}
                        def _hl_cell(val):
                            if not val:
                                return ""
                            theme_name = val.split(" (β=")[0]
                            if theme_name not in _STANDARD:
                                return "background-color: #FFF3CD; color: #1A2B3C"
                            return ""
                        st.dataframe(df_short.style.map(_hl_cell), use_container_width=True)
                    st.markdown("#### Shortlisted Drivers — Services (β + Δ > 0.5)")
                    st.caption("Positive shortlist: combined β ≥ 1 AND |Δ| > 0.5. Negative: β ≤ -0.25 AND |Δ| > 0.5.")
                    if _b112_pos_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Positive drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_svc(_b112_pos_detail, "positive", True)
                    if _b112_neg_detail:
                        st.markdown('<p style="font-size:14px;font-weight:700;color:#1A2B3C;margin:8px 0 4px">Negative drivers</p>', unsafe_allow_html=True)
                        _shortlist_table_svc(_b112_neg_detail, "negative", False)

                    if _b112_pos:
                        st.markdown("#### Where Each Positive Driver Appears")
                        for _theme, _count in sorted(_b112_pos_freq.items(), key=lambda x: x[1], reverse=True):
                            with st.expander(f"{_theme} — {_count} service{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b112_total} services analysed</p>', unsafe_allow_html=True)
                                for _sv, _b, _dd in sorted(_b112_pos_detail.get(_theme, []), key=lambda x: x[1], reverse=True):
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_sv}</strong> '
                                        f'(β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Positive Drivers by Service (combined β ≥ 1)")
                        _mp2 = max(len(v) for v in _b112_pos.values())
                        _pos_df2 = pd.DataFrame({k: v + [""] * (_mp2 - len(v)) for k, v in _b112_pos.items()})
                        _pos_df2.index = [f"#{i+1}" for i in range(_mp2)]
                        st.dataframe(_pos_df2, use_container_width=True)
                    else:
                        st.info("No positive drivers with combined β ≥ 1 found at service level.")

                    if _b112_neg:
                        st.markdown("#### Where Each Negative Driver Appears")
                        for _theme, _count in sorted(_b112_neg_freq.items(), key=lambda x: x[1], reverse=True):
                            with st.expander(f"{_theme} — {_count} service{'s' if _count != 1 else ''}"):
                                st.markdown(f'<p style="font-size:11px;color:#8FA3B1;margin-bottom:6px">{_count} of {_b112_total} services analysed</p>', unsafe_allow_html=True)
                                for _sv, _b, _dd in sorted(_b112_neg_detail.get(_theme, []), key=lambda x: x[1]):
                                    _delta_str = ""
                                    if abs(_dd) > 0.5:
                                        _dc = "#27AE60" if _dd > 0 else "#C0392B"
                                        _delta_str = f', <span style="color:{_dc};font-weight:600">Δ = {_dd:+.2f}</span>'
                                    st.markdown(
                                        f'<p style="margin:2px 0;font-size:13px;color:#1A2B3C">· <strong style="color:#1A2B3C">{_sv}</strong> (β = <span style="color:#2980B9;font-weight:600">{_b:+.2f}</span>{_delta_str})</p>',
                                        unsafe_allow_html=True
                                    )
                        st.markdown("#### Top Negative Drivers by Service (combined β ≤ -0.25)")
                        _mn2 = max(len(v) for v in _b112_neg.values())
                        _neg_df2 = pd.DataFrame({k: v + [""] * (_mn2 - len(v)) for k, v in _b112_neg.items()})
                        _neg_df2.index = [f"#{i+1}" for i in range(_mn2)]
                        st.dataframe(_neg_df2, use_container_width=True)
                    else:
                        st.info("No negative drivers with combined β ≤ -0.25 found at service level.")

                with b11_main_tab:
                    st.caption(
                        "For each employee experience outcome, this analysis: (1) keeps only WoW themes "
                        "with |r| ≥ the correlation floor — weak predictors are excluded before modelling "
                        "begins; (2) runs a multiple regression with all remaining themes simultaneously "
                        "and checks whether the overall model is statistically reliable (Significance F < 0.05); "
                        "(3) progressively removes the weakest contributors until only those with an independent, "
                        "meaningful relationship remain. The priority matrix shows the confirmed behavioural "
                        "drivers across all outcomes at a glance."
                    )
                    st.markdown(
                        '<p style="font-size:13px;font-weight:600;color:#5A7080;text-transform:uppercase;'
                        'letter-spacing:0.06em;margin-bottom:4px">Select Directorate</p>',
                        unsafe_allow_html=True,
                    )
                    _preview_dir = st.session_state.get("dir_selector_corr", directorates[0])
                    _preview_df  = filtered[filtered["Q1"] == _preview_dir]
                    _preview_hc  = HEADCOUNT.get(_preview_dir)
                    _preview_pct = f" ({len(_preview_df) / _preview_hc:.0%} of headcount)" if _preview_hc else ""
                    st.caption(f"{len(_preview_df):,} respondents in **{_preview_dir}**{_preview_pct}")
                    selected_dir = st.radio(
                        "Directorate", directorates, horizontal=True,
                        label_visibility="collapsed", key="dir_selector_corr",
                    )
                    dir_df = filtered[filtered["Q1"] == selected_dir]
                    n_dir  = len(dir_df)

                    # ── Service area selector ──────────────────────────────────
                    _b1_svc_groups = get_filter_options(dir_df, "svc_group")
                    st.markdown(
                        '<p style="font-size:13px;font-weight:600;color:#5A7080;text-transform:uppercase;'
                        'letter-spacing:0.06em;margin-bottom:4px;margin-top:8px">Select Service Area</p>',
                        unsafe_allow_html=True,
                    )
                    _b1_prev_sg = st.session_state.get("b1_oda_sg", "Overall service area")
                    _b1_prev_sg_df = dir_df if _b1_prev_sg == "Overall service area" else dir_df[dir_df["svc_group"] == _b1_prev_sg] if _b1_prev_sg in _b1_svc_groups else dir_df
                    _b1_prev_sg_hc = SERVICE_AREA_HEADCOUNT.get(_b1_prev_sg) if _b1_prev_sg != "Overall service area" else HEADCOUNT.get(selected_dir)
                    _b1_prev_sg_pct = f" ({len(_b1_prev_sg_df) / _b1_prev_sg_hc:.0%} of headcount)" if _b1_prev_sg_hc else ""
                    st.caption(f"{len(_b1_prev_sg_df):,} respondents in **{_b1_prev_sg}**{_b1_prev_sg_pct}")
                    sel_b1_sg = st.radio(
                        "Select service area", ["Overall service area"] + list(_b1_svc_groups),
                        horizontal=True, label_visibility="collapsed", key="b1_oda_sg"
                    )
                    _sg_df = dir_df if sel_b1_sg == "Overall service area" else dir_df[dir_df["svc_group"] == sel_b1_sg]

                    # ── Service selector ───────────────────────────────────────
                    _b1_svcs = get_filter_options(_sg_df, "svc_name")
                    if _b1_svcs:
                        st.markdown(
                            '<p style="font-size:13px;font-weight:600;color:#5A7080;text-transform:uppercase;'
                            'letter-spacing:0.06em;margin-bottom:4px;margin-top:8px">Select Service</p>',
                            unsafe_allow_html=True,
                        )
                        _b1_prev_svc = st.session_state.get("b1_oda_svc", "Overall service")
                        _b1_prev_svc_df = _sg_df if _b1_prev_svc == "Overall service" else _sg_df[_sg_df["svc_name"] == _b1_prev_svc] if _b1_prev_svc in _b1_svcs else _sg_df
                        st.caption(f"{len(_b1_prev_svc_df):,} respondents in **{_b1_prev_svc}**")
                        sel_b1_svc = st.radio(
                            "Select service", ["Overall service"] + list(_b1_svcs),
                            horizontal=True, label_visibility="collapsed", key="b1_oda_svc"
                        )
                        oda_input_df = _sg_df if sel_b1_svc == "Overall service" else _sg_df[_sg_df["svc_name"] == sel_b1_svc]
                    else:
                        oda_input_df = _sg_df

                    # ── Controls ──────────────────────────────────────────────
                    oda_c1, oda_c2, oda_c3 = st.columns([3, 2, 2])
                    with oda_c1:
                        wow_choice_oda = st.radio(
                            "WoW predictors", ["Place (P)", "Individual (I)"],
                            horizontal=True, key="b_oda_wow"
                        )
                    with oda_c2:
                        r_thresh_oda = st.slider(
                            "Correlation floor |r| ≥", 0.00, 0.40, 0.20, 0.05, key="b_oda_r"
                        )
                        st.markdown(
                            '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                            'Recommended: 0.20</p>',
                            unsafe_allow_html=True)
                    with oda_c3:
                        p_thresh_oda = st.slider(
                            "P-value threshold", 0.00, 0.10, 0.05, 0.01, key="b_oda_p"
                        )
                        st.markdown(
                            '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                            'Recommended: 0.05</p>',
                            unsafe_allow_html=True)

                    if "Place" in wow_choice_oda:
                        oda_pred_cols = tuple(WOW_PLACE_COLS)
                        oda_col_to_lbl = dict(zip(WOW_PLACE_COLS, WOW_THEMES))
                    else:
                        oda_pred_cols = tuple(WOW_IND_COLS)
                        oda_col_to_lbl = dict(zip(WOW_IND_COLS, WOW_THEMES))
                    oda_pred_labels = tuple(oda_col_to_lbl[c] for c in oda_pred_cols)

                    with st.spinner("Running outcome driver analysis…"):
                        oda_results = run_driver_analysis_batch(
                            oda_input_df,
                            tuple(OUTCOME_COLS), tuple(OUTCOME_LABELS),
                            oda_pred_cols, oda_pred_labels,
                            r_thresh_oda, p_thresh_oda,
                        )

                    ok_outcomes = {lbl: res for lbl, res in oda_results.items()
                                   if res["status"] == "ok"}

                    _oda_combined = {}
                    for _res in ok_outcomes.values():
                        for _, _row in _res["coef_df"].iterrows():
                            _lbl = _row["label"]
                            _oda_combined[_lbl] = _oda_combined.get(_lbl, 0.0) + _row["β_std"]
                    if _oda_combined:
                        _oda_sorted = sorted(_oda_combined.items(), key=lambda x: x[1], reverse=True)
                        _oda_labels = [x[0] for x in _oda_sorted]
                        _oda_scores = [x[1] for x in _oda_sorted]
                        st.markdown("#### Top WoW Drivers — Combined Across All Outcomes")
                        st.caption(
                            "Each bar shows the sum of standardised β values for that WoW theme across all "
                            "individual outcome models where it was a confirmed driver, preserving sign. "
                            "A consistently positive theme scores high; mixed effects partially cancel out. "
                            "Red = net positive effect; blue = net negative."
                        )
                        _fig_oda_c = go.Figure(go.Bar(
                            x=_oda_labels,
                            y=_oda_scores,
                            text=[f"{s:+.2f}" for s in _oda_scores],
                            textposition="outside",
                            textfont=dict(color="#1A2B3C", size=10),
                            marker_color=[RED if s >= 0 else PRIMARY for s in _oda_scores],
                            hovertemplate="<b>%{x}</b><br>Combined β = %{y:.3f}<extra></extra>",
                        ))
                        _fig_oda_c.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=40, b=260),
                            xaxis=dict(tickangle=-40, tickfont=dict(color="#1A2B3C", size=10)),
                            yaxis=dict(
                                title=dict(text="Combined β", font=dict(color="#1A2B3C")),
                                zeroline=True, zerolinecolor="#D6E0EA",
                                tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                            ),
                            height=560,
                        )
                        st.plotly_chart(_fig_oda_c, use_container_width=True, key="b_oda_combined_bar")

                    # ── Priority matrix ────────────────────────────────────────
                    st.markdown("---")
                    st.markdown("#### Top Ways of Working Drivers — Per Outcome")
                    st.caption(
                        "Each cell shows the standardised β for a WoW theme that survived backward "
                        "elimination for that outcome's model. Blank = not a confirmed driver (below "
                        "the correlation floor or eliminated). Blue = positive effect (more of this "
                        "behaviour → better outcome score); red = negative effect."
                    )

                    if not ok_outcomes:
                        st.warning(
                            "No outcomes produced a statistically reliable model with the current settings. "
                            "Try lowering the correlation floor or raising the p-value threshold."
                        )
                    else:
                        all_retained_labels = sorted(set(
                            oda_col_to_lbl[c]
                            for res in ok_outcomes.values()
                            for c in res["retained"]
                        ))
                        matrix_rows = {}
                        for out_lbl, res in ok_outcomes.items():
                            lbl_to_beta = {
                                row["label"]: row["β_std"]
                                for _, row in res["coef_df"].iterrows()
                            }
                            matrix_rows[out_lbl] = {
                                theme: lbl_to_beta.get(theme, np.nan)
                                for theme in all_retained_labels
                            }
                        priority_df = pd.DataFrame(matrix_rows, index=all_retained_labels)
                        _ordered_outcomes = [lbl for lbl in OUTCOME_LABELS if lbl in priority_df.columns]
                        priority_df = priority_df[_ordered_outcomes]

                        beta_abs_max = priority_df.abs().max().max()
                        beta_abs_max = max(beta_abs_max, 0.01)

                        priority_df_T = priority_df.T  # outcomes as rows, WoW as columns
                        fig_matrix = go.Figure(go.Heatmap(
                            z=priority_df_T.values.tolist(),
                            x=list(priority_df_T.columns),
                            y=list(priority_df_T.index),
                            colorscale=[
                                [0.0,  "#0F4C6B"],
                                [0.5,  "#F7F9FC"],
                                [1.0,  "#C0392B"],
                            ],
                            zmid=0,
                            zmin=-beta_abs_max,
                            zmax=beta_abs_max,
                            text=[[
                                f"{v:.2f}" if not np.isnan(v) else ""
                                for v in row
                            ] for row in priority_df_T.values.tolist()],
                            texttemplate="%{text}",
                            textfont=dict(size=11, color="#1A2B3C"),
                            hoverongaps=False,
                            hovertemplate="<b>%{x}</b><br>%{y}<br>β = %{z:.3f}<extra></extra>",
                            colorbar=dict(title="β", len=0.7),
                        ))
                        fig_matrix.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=10, b=160),
                            xaxis=dict(
                                tickfont=dict(color="#1A2B3C", size=11),
                                tickangle=-40,
                                side="bottom",
                            ),
                            yaxis=dict(tickfont=dict(color="#1A2B3C", size=11), autorange="reversed"),
                            height=max(350, 36 * len(ok_outcomes)),
                        )
                        st.plotly_chart(fig_matrix, use_container_width=True, key="b_oda_matrix")

                    # ── Per-outcome detail ─────────────────────────────────────
                    st.markdown("---")
                    st.markdown('<p style="font-size:14px;font-weight:600;color:#5A7080;margin-top:8px;margin-bottom:4px">Breakdown per Outcome</p>', unsafe_allow_html=True)

                    for out_lbl in OUTCOME_LABELS:
                        res = oda_results.get(out_lbl, {"status": "no_predictors"})
                        status = res["status"]
                        icon = "✓" if status == "ok" else "✗"
                        with st.expander(f"{icon}  {out_lbl}"):
                            if status == "no_predictors":
                                st.info(
                                    f"No WoW themes reached |r| ≥ {r_thresh_oda:.2f} for this outcome. "
                                    "No regression was run."
                                )
                            elif status == "fail_sig_f":
                                st.error(
                                    f"**Model not statistically reliable** — Significance F = {res['sig_f']:.4f} "
                                    f"(> 0.05). The {res['n_eligible']} predictors that cleared the correlation "
                                    f"floor did not collectively explain this outcome reliably. "
                                    "Do not read into the individual coefficients."
                                )
                                mc1, mc2, mc3 = st.columns(3)
                                for col_ui, label, value in [
                                    (mc1, "Respondents", f"n = {res['n']:,}"),
                                    (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                                    (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                                ]:
                                    with col_ui:
                                        st.markdown(
                                            f'<div class="metric-card"><p class="card-label">{label}</p>'
                                            f'<p class="card-value">{value}</p></div>',
                                            unsafe_allow_html=True)
                            elif status == "all_eliminated":
                                st.warning(
                                    "All predictors were eliminated during backward stepwise. "
                                    "Try raising the p-value threshold."
                                )
                                mc1, mc2, mc3 = st.columns(3)
                                for col_ui, label, value in [
                                    (mc1, "Respondents", f"n = {res['n']:,}"),
                                    (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                                    (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                                ]:
                                    with col_ui:
                                        st.markdown(
                                            f'<div class="metric-card"><p class="card-label">{label}</p>'
                                            f'<p class="card-value">{value}</p></div>',
                                            unsafe_allow_html=True)
                            else:
                                coef_df = res["coef_df"]
                                # Summary cards
                                mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                                for col_ui, label, value in [
                                    (mc1, "Respondents",         f"n = {res['n']:,}"),
                                    (mc2, "Significance F",      f"{res['sig_f']:.4f}"),
                                    (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                                    (mc4, "Adj R² (final model)",f"{res['adj_r2_final']:.3f}"),
                                    (mc5, "Confirmed drivers",   str(len(res["retained"]))),
                                ]:
                                    with col_ui:
                                        st.markdown(
                                            f'<div class="metric-card"><p class="card-label">{label}</p>'
                                            f'<p class="card-value">{value}</p></div>',
                                            unsafe_allow_html=True)

                                # β coefficient chart
                                st.markdown("##### Confirmed Drivers — Standardised β Coefficients")
                                st.caption(
                                    "Each bar shows a retained WoW theme's standardised β and the equivalent "
                                    "point change on the 1–5 outcome scale (Δ pts). β reflects the independent "
                                    "contribution of that theme after controlling for all others — a 1 SD shift "
                                    "in the WoW theme produces the shown point change in the outcome. "
                                    "* p<0.05  ** p<0.01  *** p<0.001"
                                )
                                coef_sorted = coef_df.reindex(
                                    coef_df["β_std"].abs().sort_values(ascending=False).index
                                )
                                fig_coef = go.Figure(go.Bar(
                                    x=coef_sorted["label"],
                                    y=coef_sorted["β_std"],
                                    text=coef_sorted["text"],
                                    textposition="outside",
                                    textfont=dict(color="#1A2B3C", size=10),
                                    marker_color=[
                                        RED if v >= 0 else PRIMARY
                                        for v in coef_sorted["β_std"]
                                    ],
                                    hovertemplate="<b>%{x}</b><br>β (std) = %{y:.3f}<extra></extra>",
                                ))
                                fig_coef.update_layout(
                                    font=dict(family="Inter", color="#1A2B3C"),
                                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                                    margin=dict(l=10, r=10, t=40, b=160),
                                    xaxis=dict(
                                        tickangle=-40,
                                        tickfont=dict(color="#1A2B3C", size=10),
                                    ),
                                    yaxis=dict(
                                        title=dict(text="Standardised β", font=dict(color="#1A2B3C")),
                                        zeroline=True, zerolinecolor="#D6E0EA",
                                        tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                    ),
                                    height=420,
                                )
                                st.plotly_chart(fig_coef, use_container_width=True,
                                                key=f"b_oda_coef_{out_lbl}")

                                # Elimination path
                                st.markdown("##### Elimination Path — Adj R² at each step")
                                st.caption(
                                    "Each point shows the model's Adjusted R² after removing the least "
                                    "significant predictor at that step. A flat or rising line means those "
                                    "predictors weren't earning their place. A drop signals the core has "
                                    "been reached."
                                )
                                log_df = pd.DataFrame(res["elim_log"])
                                if "Predictors in model" in log_df.columns:
                                    fig_elim = go.Figure(go.Scatter(
                                        x=log_df["Predictors in model"],
                                        y=log_df["Adj. R²"],
                                        mode="lines+markers",
                                        line=dict(color=PRIMARY, width=2),
                                        marker=dict(color=PRIMARY, size=8),
                                        hovertemplate="Predictors: %{x}<br>Adj. R² = %{y:.4f}<extra></extra>",
                                    ))
                                    fig_elim.update_layout(
                                        font=dict(family="Inter", color="#1A2B3C"),
                                        paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                                        margin=dict(l=10, r=10, t=10, b=10),
                                        xaxis=dict(
                                            title=dict(text="Predictors in model", font=dict(color="#1A2B3C")),
                                            autorange="reversed",
                                            tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                        ),
                                        yaxis=dict(
                                            title=dict(text="Adjusted R²", font=dict(color="#1A2B3C")),
                                            tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                        ),
                                        height=300,
                                    )
                                    st.plotly_chart(fig_elim, use_container_width=True,
                                                    key=f"b_oda_elim_{out_lbl}")

            with oda_lf_tab:
                st.caption(
                    "Runs the same multiple regression pipeline as the Outcomes tab, but uses "
                    "the latent factor scores (LV1, LV2 …) derived from the EFA in the "
                    "Outcome Clusters tab as the dependent variables. All WoW themes enter the "
                    "model directly — no correlation floor is applied since there is no "
                    "pre-existing heatmap for latent factors. Backward elimination then removes "
                    "non-significant predictors until only confirmed drivers remain."
                )
                st.markdown(
                    '<p style="font-size:13px;font-weight:600;color:#5A7080;text-transform:uppercase;'
                    'letter-spacing:0.06em;margin-bottom:4px">Select Directorate</p>',
                    unsafe_allow_html=True,
                )
                _lf_prev_dir = st.session_state.get("b_dir_selector_lf", directorates[0])
                _lf_prev_df  = filtered[filtered["Q1"] == _lf_prev_dir]
                _lf_prev_hc  = HEADCOUNT.get(_lf_prev_dir)
                _lf_prev_pct = f" ({len(_lf_prev_df) / _lf_prev_hc:.0%} of headcount)" if _lf_prev_hc else ""
                st.caption(f"{len(_lf_prev_df):,} respondents in **{_lf_prev_dir}**{_lf_prev_pct}")
                selected_dir = st.radio(
                    "Directorate", directorates, horizontal=True,
                    label_visibility="collapsed", key="b_dir_selector_lf",
                )
                dir_df = filtered[filtered["Q1"] == selected_dir]
                n_dir  = len(dir_df)

                # ── Controls row 1: WoW type + p-value ────────────────
                lf_c1, lf_c2, lf_c3 = st.columns([3, 2, 2])
                with lf_c1:
                    wow_choice_lf = st.radio(
                        "WoW predictors", ["Place (P)", "Individual (I)"],
                        horizontal=True, key="b_lf_wow"
                    )
                with lf_c3:
                    p_thresh_lf = st.slider(
                        "P-value threshold", 0.00, 0.10, 0.05, 0.01, key="b_lf_p"
                    )
                    st.markdown(
                        '<p style="font-size:11px;color:#8FA3B1;margin-top:-10px">'
                        'Recommended: 0.05</p>',
                        unsafe_allow_html=True)

                # ── Controls row 2: n_factors + factor preview ─────────
                lf_r2_left, _ = st.columns([2, 5])
                with lf_r2_left:
                    lf_n_factors = st.slider(
                        "Number of latent factors", 1, 6, 3, key="b_lf_nfactors"
                    )
                with st.expander("Scree Plot — How many hidden factors exist in your outcome data?"):
                    _sem_clean_b = dir_df[OUTCOME_COLS].dropna()
                    if len(_sem_clean_b) < 20:
                        st.warning("Too few complete responses for factor analysis.")
                    else:
                        _corr_b = np.corrcoef(_sem_clean_b.values.T)
                        _eig_b = sorted(np.linalg.eigvalsh(_corr_b).tolist(), reverse=True)
                        st.caption(
                            "Each point is one potential factor and its height shows how much shared information "
                            "it captures. Factors above the dashed line (eigenvalue > 1) are worth keeping. "
                            "Look for where the curve flattens — that bend marks your real underlying themes."
                        )
                        _fig_scree_b = go.Figure([go.Scatter(
                            x=list(range(1, len(_eig_b) + 1)), y=_eig_b,
                            mode="lines+markers",
                            line=dict(color=PRIMARY, width=2),
                            marker=dict(color=PRIMARY, size=7),
                            hovertemplate="Factor %{x}<br>Eigenvalue = %{y:.3f}<extra></extra>",
                        )])
                        _fig_scree_b.add_hline(y=1, line_dash="dash", line_color="#C0392B",
                            annotation_text="Kaiser criterion (λ = 1)",
                            annotation_font=dict(color="#C0392B", size=11))
                        _fig_scree_b.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=20, b=10),
                            xaxis=dict(title=dict(text="Factor", font=dict(color="#1A2B3C")),
                                       tickvals=list(range(1, len(_eig_b) + 1)),
                                       tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2"),
                            yaxis=dict(title=dict(text="Eigenvalue", font=dict(color="#1A2B3C")),
                                       tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2"),
                            height=300,
                        )
                        st.plotly_chart(_fig_scree_b, use_container_width=True, key="b_sem_scree")
                if "Place" in wow_choice_lf:
                    lf_pred_cols = tuple(WOW_PLACE_COLS)
                    lf_col_to_lbl = dict(zip(WOW_PLACE_COLS, WOW_THEMES))
                else:
                    lf_pred_cols = tuple(WOW_IND_COLS)
                    lf_col_to_lbl = dict(zip(WOW_IND_COLS, WOW_THEMES))
                lf_pred_labels = tuple(lf_col_to_lbl[c] for c in lf_pred_cols)

                outcome_clean_lf = dir_df[OUTCOME_COLS].dropna()
                if len(outcome_clean_lf) < 20:
                    st.warning("Too few complete responses for factor analysis.")
                else:
                    with st.spinner("Computing factor scores and running regressions…"):
                        factor_scores_df = compute_factor_scores(
                            dir_df[OUTCOME_COLS], lf_n_factors
                        )
                        lf_results = run_lf_driver_analysis(
                            dir_df, factor_scores_df,
                            lf_pred_cols, lf_pred_labels, p_thresh_lf
                        )

                    # ── Factor preview + WoW drivers ──────────────────
                    with _timed(f"B — run_efa ({lf_n_factors} factors)"):
                        _lf_loadings, _ = run_efa(dir_df[OUTCOME_COLS], lf_n_factors)
                    _lf_loadings.index = OUTCOME_LABELS
                    _lf_cols = [f"LV{i+1}" for i in range(lf_n_factors)]
                    _lf_max = _lf_loadings.abs().max(axis=1)
                    _lf_primary = _lf_loadings.abs().idxmax(axis=1).copy()
                    _lf_primary[_lf_max < 0.3] = "Unassigned"
                    _lf_members = {lv: [] for lv in _lf_cols}
                    for _outcome, _lv in _lf_primary.items():
                        if _lv != "Unassigned":
                            _lf_members[_lv].append(
                                (_outcome, _lf_loadings.loc[_outcome, _lv])
                            )
                    for _lv in _lf_members:
                        _lf_members[_lv].sort(key=lambda x: abs(x[1]), reverse=True)

                    # ── Combined top WoW drivers across all latent factors ──
                    _combined_betas: dict = {}
                    for _lv, _res in lf_results.items():
                        if _res.get("status") == "ok":
                            for _, _row in _res["coef_df"].iterrows():
                                _lbl = _row["label"]
                                _combined_betas[_lbl] = _combined_betas.get(_lbl, 0.0) + _row["β_std"]
                    _preview_cols = st.columns(lf_n_factors)
                    for _col_ui, _lv in zip(_preview_cols, _lf_cols):
                        _items_html = "".join(
                            f'<p class="card-sub">· {m} ({v:.2f})</p>'
                            for m, v in _lf_members[_lv]
                        ) or '<p class="card-sub"><em>None assigned</em></p>'
                        _res = lf_results.get(_lv, {})
                        if _res.get("status") == "ok":
                            _coef = _res["coef_df"].copy()
                            _coef_sorted = _coef.reindex(
                                _coef["β_std"].abs().sort_values(ascending=False).index
                            )
                            _drivers_html = "".join(
                                f'<p class="card-sub">· {row["label"]} '
                                f'(<span style="color:{"#C0392B" if row["β_std"] >= 0 else "#2980B9"};font-weight:600">'
                                f'{row["β_std"]:+.2f}</span>)</p>'
                                for _, row in _coef_sorted.iterrows()
                            )
                        elif _res.get("status") == "fail_sig_f":
                            _drivers_html = (
                                f'<p class="card-sub"><em>Model not reliable '
                                f'(Sig. F = {_res["sig_f"]:.3f})</em></p>'
                            )
                        else:
                            _drivers_html = '<p class="card-sub"><em>No confirmed drivers</em></p>'
                        with _col_ui:
                            st.markdown(
                                f'<div class="metric-card-lg">'
                                f'<p class="card-label">{_lv}</p>'
                                f'{_items_html}'
                                f'<p class="card-label" style="margin-top:10px">'
                                f'Top WoW drivers (β)</p>'
                                f'{_drivers_html}'
                                f'</div>',
                                unsafe_allow_html=True)

                    if _combined_betas:
                        _sorted_combined = sorted(_combined_betas.items(), key=lambda x: x[1], reverse=True)
                        _cb_labels = [x[0] for x in _sorted_combined]
                        _cb_scores = [x[1] for x in _sorted_combined]
                        st.markdown("#### Overall Top WoW Drivers Across All Outcome Groups (Latent Factors)")
                        st.caption(
                            "Each bar shows the sum of standardised β values for that WoW theme across all "
                            "latent factors where it was a confirmed driver, preserving sign. A consistently "
                            "positive theme scores high; one with mixed effects across factors will partially "
                            "cancel out. Red = net positive effect; blue = net negative."
                        )
                        _lf_c_stmts = WOW_PLACE_STATEMENTS if "Place" in wow_choice_lf else WOW_IND_STATEMENTS
                        fig_combined = go.Figure(go.Bar(
                            x=_cb_labels,
                            y=_cb_scores,
                            text=[f"{s:+.2f}" for s in _cb_scores],
                            textposition="outside",
                            textfont=dict(color="#1A2B3C", size=10),
                            marker_color=[RED if s >= 0 else PRIMARY for s in _cb_scores],
                            customdata=[_lf_c_stmts.get(lbl, "") for lbl in _cb_labels],
                            hovertemplate="<b>%{x}</b><br><i>%{customdata}</i><br>Combined β = %{y:.3f}<extra></extra>",
                        ))
                        fig_combined.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=40, b=260),
                            xaxis=dict(
                                tickangle=-40,
                                tickfont=dict(color="#1A2B3C", size=10),
                            ),
                            yaxis=dict(
                                title=dict(text="Combined β", font=dict(color="#1A2B3C")),
                                zeroline=True, zerolinecolor="#D6E0EA",
                                tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                            ),
                            height=560,
                        )
                        st.plotly_chart(fig_combined, use_container_width=True, key="b_lf_combined_bar")

                    ok_lf = {lv: res for lv, res in lf_results.items()
                             if res["status"] == "ok"}

                    # ── Priority matrix ────────────────────────────────
                    st.markdown("---")
                    st.markdown("#### WoW Drivers by Outcome Group (Latent Factor)")
                    st.caption(
                        "Each cell shows the standardised β for a WoW theme retained in that "
                        "latent factor's final model. Blank = not a confirmed driver. "
                        "Red = positive effect; blue = negative effect."
                    )

                    if not ok_lf:
                        st.warning(
                            "No latent factors produced a statistically reliable model. "
                            "Try raising the p-value threshold."
                        )
                    else:
                        all_lf_retained = sorted(set(
                            lf_col_to_lbl[c]
                            for res in ok_lf.values()
                            for c in res["retained"]
                        ))
                        lf_matrix_rows = {}
                        for lv, res in ok_lf.items():
                            lbl_to_beta = {
                                row["label"]: row["β_std"]
                                for _, row in res["coef_df"].iterrows()
                            }
                            lf_matrix_rows[lv] = {
                                theme: lbl_to_beta.get(theme, np.nan)
                                for theme in all_lf_retained
                            }
                        lf_priority_df = pd.DataFrame(lf_matrix_rows, index=all_lf_retained)
                        lf_priority_df_T = lf_priority_df.T

                        lf_beta_max = max(lf_priority_df.abs().max().max(), 0.01)

                        fig_lf_matrix = go.Figure(go.Heatmap(
                            z=lf_priority_df_T.values.tolist(),
                            x=list(lf_priority_df_T.columns),
                            y=list(lf_priority_df_T.index),
                            colorscale=[
                                [0.0, "#0F4C6B"],
                                [0.5, "#F7F9FC"],
                                [1.0, "#C0392B"],
                            ],
                            zmid=0,
                            zmin=-lf_beta_max,
                            zmax=lf_beta_max,
                            text=[[
                                f"{v:.2f}" if not np.isnan(v) else ""
                                for v in row
                            ] for row in lf_priority_df_T.values.tolist()],
                            texttemplate="%{text}",
                            textfont=dict(size=11, color="#1A2B3C"),
                            hoverongaps=False,
                            hovertemplate="<b>%{x}</b><br>%{y}<br>β = %{z:.3f}<extra></extra>",
                            colorbar=dict(title="β", len=0.7),
                        ))
                        fig_lf_matrix.update_layout(
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=10, b=160),
                            xaxis=dict(
                                tickfont=dict(color="#1A2B3C", size=11),
                                tickangle=-40,
                                side="bottom",
                            ),
                            yaxis=dict(tickfont=dict(color="#1A2B3C", size=11)),
                            height=max(300, 36 * len(ok_lf)),
                        )
                        st.plotly_chart(fig_lf_matrix, use_container_width=True,
                                        key="b_lf_matrix")

                    # ── Per-factor detail ──────────────────────────────
                    st.markdown("---")
                    st.markdown("#### Per-Factor Detail")

                    for lv in factor_scores_df.columns:
                        res = lf_results.get(lv, {"status": "all_eliminated"})
                        icon = "✓" if res["status"] == "ok" else "✗"
                        with st.expander(f"{icon}  {lv}"):
                            if res["status"] == "fail_sig_f":
                                st.error(
                                    f"**Model not statistically reliable** — "
                                    f"Significance F = {res['sig_f']:.4f} (> 0.05)."
                                )
                                mc1, mc2, mc3 = st.columns(3)
                                for col_ui, label, value in [
                                    (mc1, "Respondents", f"n = {res['n']:,}"),
                                    (mc2, "Significance F", f"{res['sig_f']:.4f}"),
                                    (mc3, "Adj R² (full model)", f"{res['adj_r2_full']:.3f}"),
                                ]:
                                    with col_ui:
                                        st.markdown(
                                            f'<div class="metric-card"><p class="card-label">{label}</p>'
                                            f'<p class="card-value">{value}</p></div>',
                                            unsafe_allow_html=True)
                            elif res["status"] == "all_eliminated":
                                st.warning("All predictors were eliminated. Try raising the p-value threshold.")
                            else:
                                coef_df = res["coef_df"]
                                mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                                for col_ui, label, value in [
                                    (mc1, "Respondents",          f"n = {res['n']:,}"),
                                    (mc2, "Significance F",       f"{res['sig_f']:.4f}"),
                                    (mc3, "Adj R² (full model)",  f"{res['adj_r2_full']:.3f}"),
                                    (mc4, "Adj R² (final model)", f"{res['adj_r2_final']:.3f}"),
                                    (mc5, "Confirmed drivers",    str(len(res["retained"]))),
                                ]:
                                    with col_ui:
                                        st.markdown(
                                            f'<div class="metric-card"><p class="card-label">{label}</p>'
                                            f'<p class="card-value">{value}</p></div>',
                                            unsafe_allow_html=True)

                                st.markdown("##### Confirmed Drivers — Standardised β Coefficients")
                                coef_sorted = coef_df.reindex(
                                    coef_df["β_std"].abs().sort_values(ascending=False).index
                                )
                                fig_lf_coef = go.Figure(go.Bar(
                                    x=coef_sorted["label"],
                                    y=coef_sorted["β_std"],
                                    text=coef_sorted["text"],
                                    textposition="outside",
                                    textfont=dict(color="#1A2B3C", size=10),
                                    marker_color=[
                                        RED if v >= 0 else PRIMARY
                                        for v in coef_sorted["β_std"]
                                    ],
                                    hovertemplate="<b>%{x}</b><br>β (std) = %{y:.3f}<extra></extra>",
                                ))
                                fig_lf_coef.update_layout(
                                    font=dict(family="Inter", color="#1A2B3C"),
                                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                                    margin=dict(l=10, r=10, t=40, b=160),
                                    xaxis=dict(
                                        tickangle=-40,
                                        tickfont=dict(color="#1A2B3C", size=10),
                                    ),
                                    yaxis=dict(
                                        title=dict(text="Standardised β", font=dict(color="#1A2B3C")),
                                        zeroline=True, zerolinecolor="#D6E0EA",
                                        tickfont=dict(color="#1A2B3C"), gridcolor="#E8EEF2",
                                    ),
                                    height=420,
                                )
                                st.plotly_chart(fig_lf_coef, use_container_width=True,
                                                key=f"b_lf_coef_{lv}")

                st.markdown("---")

            # Use whichever directorate was last selected in either ODA tab
            _b_dir = st.session_state.get('dir_selector_corr',
                     st.session_state.get('b_dir_selector_lf', directorates[0]))
            dir_df = filtered[filtered['Q1'] == _b_dir]
            n_dir  = len(dir_df)
            selected_dir = _b_dir

        with heatmaps_tab:
            h1, h2, h3 = st.tabs([
                "WoW × Outcomes",
                "WoW × WoW",
                "Outcomes × Outcomes",
            ])

            # ── H1: Ways of Working × Outcomes ───────────────────
            with h1:
                a1_place, a1_ind = st.tabs(["Place (P)", "Individual (I)"])
                with a1_place:
                    st.markdown("#### Correlational Heatmap: Ways of Working (Place) × Employee Experience")
                    st.caption(
                        "Each cell shows the Spearman correlation (r) between a Place-level Ways of Working "
                        "theme (rows) and an employee experience outcome (columns). Place themes reflect the "
                        "shared working environment — norms and behaviours that feel consistent across the team "
                        "or directorate. Darker blue = stronger positive relationship; darker red = stronger "
                        "negative. Values above ~0.2 are worth noting; above ~0.4 indicate a meaningful pattern. "
                        "Read across a row to see which outcomes a given working norm connects to most strongly."
                    )
                    mat = spearman_matrix(dir_df, WOW_PLACE_COLS, OUTCOME_COLS)
                    mat.index   = WOW_THEMES
                    mat.columns = OUTCOME_LABELS
                    render_heatmap_cards(n_dir, mat, HEADCOUNT.get(selected_dir))
                    st.plotly_chart(make_heatmap(mat.T, OUTCOME_LABELS, WOW_THEMES, y_statements=OUTCOME_STATEMENTS, x_statements=WOW_PLACE_STATEMENTS),
                                    use_container_width=True, key="bh1_place")
                with a1_ind:
                    st.markdown("#### Correlational Heatmap: Ways of Working (Individual) × Employee Experience")
                    st.caption(
                        "Each cell shows the Spearman correlation (r) between an Individual-level Ways of Working "
                        "theme (rows) and an employee experience outcome (columns). Individual themes reflect how "
                        "each person personally operates — their own habits and approach, regardless of what the "
                        "team around them does. Darker blue = stronger positive relationship; darker red = stronger "
                        "negative. Compare this map with the Place heatmap to see whether environmental or personal "
                        "behaviours are more strongly linked to each outcome."
                    )
                    mat = spearman_matrix(dir_df, WOW_IND_COLS, OUTCOME_COLS)
                    mat.index   = WOW_THEMES
                    mat.columns = OUTCOME_LABELS
                    render_heatmap_cards(n_dir, mat, HEADCOUNT.get(selected_dir))
                    st.plotly_chart(make_heatmap(mat.T, OUTCOME_LABELS, WOW_THEMES, y_statements=OUTCOME_STATEMENTS, x_statements=WOW_IND_STATEMENTS),
                                    use_container_width=True, key="bh1_ind")

            # ── H2: Ways of Working × Ways of Working ────────────
            with h2:
                a2_place, a2_ind = st.tabs(["Place (P)", "Individual (I)"])
                with a2_place:
                    st.markdown("#### Correlational Heatmap: Ways of Working (Place) × Ways of Working (Place)")
                    st.caption(
                        "Each cell shows the Spearman correlation between two Place-level Ways of Working themes. "
                        "Strong positive correlations (dark blue) mean those norms tend to co-exist — groups that "
                        "display one tend to display the other too. Strong negative correlations (dark red) indicate "
                        "genuinely opposing orientations. Use this to identify clusters of related working norms "
                        "and to flag themes that may be measuring the same underlying dimension before regression modelling."
                    )
                    mat = make_writable_matrix(spearman_matrix(dir_df, WOW_PLACE_COLS, WOW_PLACE_COLS))
                    mat.index = mat.columns = WOW_THEMES
                    fill_diagonal_with_nan(mat)
                    render_heatmap_cards(n_dir, mat, HEADCOUNT.get(selected_dir))
                    st.plotly_chart(make_heatmap(mat, WOW_THEMES, WOW_THEMES, y_statements=WOW_PLACE_STATEMENTS, x_statements=WOW_PLACE_STATEMENTS),
                                    use_container_width=True, key="bh2_place")
                with a2_ind:
                    st.markdown("#### Correlational Heatmap: Ways of Working (Individual) × Ways of Working (Individual)")
                    st.caption(
                        "Each cell shows the Spearman correlation between two Individual-level Ways of Working themes. "
                        "Strong positive correlations mean those personal behaviours tend to appear together in the same "
                        "people. Use this alongside the Place heatmap to see whether individual habits cluster differently "
                        "from the team-level norms — a divergence can point to tension between personal style and the "
                        "shared working environment."
                    )
                    mat = make_writable_matrix(spearman_matrix(dir_df, WOW_IND_COLS, WOW_IND_COLS))
                    mat.index = mat.columns = WOW_THEMES
                    fill_diagonal_with_nan(mat)
                    render_heatmap_cards(n_dir, mat, HEADCOUNT.get(selected_dir))
                    st.plotly_chart(make_heatmap(mat, WOW_THEMES, WOW_THEMES, y_statements=WOW_IND_STATEMENTS, x_statements=WOW_IND_STATEMENTS),
                                    use_container_width=True, key="bh2_ind")

            # ── H3: Outcomes × Outcomes ───────────────────────────
            with h3:
                st.markdown("#### Correlational Heatmap: Employee Experience × Employee Experience")
                st.caption(
                    "Each cell shows the Spearman correlation between two employee experience outcomes. "
                    "Strong positive correlations indicate outcomes that tend to rise and fall together — "
                    "suggesting they may be capturing the same underlying dimension of experience. "
                    "Use this to understand how interconnected the outcome measures are before interpreting "
                    "individual scores in isolation, and to inform how outcomes might be grouped for further analysis."
                )
                mat = make_writable_matrix(spearman_matrix(dir_df, OUTCOME_COLS, OUTCOME_COLS))
                mat.index = mat.columns = OUTCOME_LABELS
                fill_diagonal_with_nan(mat)
                render_heatmap_cards(n_dir, mat, HEADCOUNT.get(selected_dir))
                st.plotly_chart(make_heatmap(mat, OUTCOME_LABELS, OUTCOME_LABELS, y_statements=OUTCOME_STATEMENTS, x_statements=OUTCOME_STATEMENTS),
                                use_container_width=True, key="bh3")

    # ── Descriptive Analysis group ────────────────────────
    with _timed("B — b_desc_group render"), b_desc_group:
        b4, b5 = st.tabs([
            "B2.1: Ways of Working",
            "B2.2: Sentiment Outcomes",
        ])

        # ── B4 ────────────────────────────────────────────────
        with b4:
            b4_sub, b4_org = st.tabs([
                "B2.1.1: WoW — Place vs. Individual",
                "B2.1.2: WoW — Org Breakdown",
            ])

            with b4_sub:
                st.caption("Average scores for each Way of Working theme, split by Place and Individual, for the selected directorate, service area, and service.")
                chart_df, dir_df, selected_dir, _b4_overall, _b4_svc_area_df = svc_selectors("b4", filtered, directorates)
                st.markdown("#### Ways of Working — Average Scores by Service Area")
                st.plotly_chart(make_wow_bar_chart(chart_df), use_container_width=True, key="b4_bar")

            with b4_org:
                _b4_org_p, _b4_org_i = st.tabs(["Place (P)", "Individual (I)"])

                def _render_hier_wow(wow_cols, sel_key, chart_key):
                    st.caption(
                        "For each Way of Working theme, compares the selected service against its "
                        "service area overall, directorate overall, and council overall."
                    )
                    _cdf, _ddf, _sel_dir, _is_overall, _svc_area_df = svc_selectors(sel_key, filtered, directorates)

                    def _bar(name, df, colour):
                        return go.Bar(
                            name=name, x=WOW_THEMES,
                            y=[df[c].mean() for c in wow_cols],
                            marker_color=colour,
                            hovertemplate=f"{name}<br>%{{x}}<br>Score: %{{y:.2f}}<extra></extra>",
                        )

                    fig = go.Figure()
                    fig.add_trace(_bar("Selected group", _cdf, PRIMARY))
                    if _svc_area_df is not None:
                        fig.add_trace(_bar("Service area overall", _svc_area_df, "#5A9BB5"))
                    if not _is_overall:
                        fig.add_trace(_bar("Directorate overall", _ddf, "#6BA3BA"))
                    fig.add_trace(_bar("Council overall", filtered, "#A8C8D8"))
                    fig.update_layout(
                        barmode="group",
                        font=dict(family="Inter", color="#1A2B3C"),
                        paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                        margin=dict(l=10, r=10, t=10, b=180),
                        xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
                        yaxis=dict(
                            title=dict(text="Average Score", font=dict(color="#1A2B3C", size=12)),
                            range=[-0.2, 4.5], tickfont=dict(size=11, color="#1A2B3C"), gridcolor="#E8EEF2",
                        ),
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                                    font=dict(color="#1A2B3C", size=12)),
                        height=560,
                    )
                    st.markdown("#### Ways of Working — Average Score by Directorate / Service Area / Service")
                    st.plotly_chart(fig, use_container_width=True, key=chart_key)

                with _b4_org_p:
                    _render_hier_wow(WOW_PLACE_COLS, "b4_org_p", "b4_org_p_bar")
                with _b4_org_i:
                    _render_hier_wow(WOW_IND_COLS, "b4_org_i", "b4_org_i_bar")

        # ── B5 ────────────────────────────────────────────────
        with b5:
            st.caption("Average scores for each employee experience outcome for the selected service area, compared against the directorate-wide average.")
            chart_df, dir_df, selected_dir, _b5_overall, _b5_svc_area_df = svc_selectors("b5", filtered, directorates)
            overall_df = None if _b5_overall else dir_df
            council_df = filtered
            _hc_b5 = HEADCOUNT.get(selected_dir) if _b5_overall else None
            _pct_b5 = f" ({len(chart_df) / _hc_b5:.0%} of headcount)" if _hc_b5 else ""
            st.markdown("#### Employee Experience — Average Score by Directorate / Service Area / Service")
            st.plotly_chart(make_outcome_bar_chart(chart_df, overall_df, council_df, _b5_svc_area_df), use_container_width=True,
                            key="b5_bar")


# ── Section C service maps (module-level so @st.cache_data can reference them) ────
C1_SERVICE_MAP = [
    ("AdultsCommg",                   lambda d: d[d["svc_group"] == "Adults Commissioning"]),
    ("AdltOps HIS & IC",              lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["svc_name"] == "CHC, HIS and Intermediate Care")]),
    ("AdltOps Nbhd MH & LD",          lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["svc_name"].isin(["Neighbourhoods", "Mental Health", "Learning Disabilities"]))]),
    ("AdltOps OT",                    lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["svc_name"] == "Occupational Therapy: Staff Dev, SILC and DFG")]),
    ("AdltOps Social Work",           lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["svc_name"] == "Social Work: Staff Dev, Safeguarding, MCA & DoLS")]),
    ("Adlts Housing General Fund",    lambda d: d[(d["svc_group"] == "Adults Housing General Fund") & (d["Q1"] == "Adult Services and Housing")]),
    ("CEO-ExecDir-ServDir",           lambda d: d[(d["svc_group"] == "CEO and Electoral Services") & (d["svc_name"] == "Exec and Service Directors")]),
    ("CEO & Elections",               lambda d: d[d["svc_group"] == "CEO and Electoral Services"]),
    ("CEO-Elections",                 lambda d: d[(d["svc_group"] == "CEO and Electoral Services") & (d["svc_name"] == "Elections")]),
    ("C&F CLA Leaving Care & CWD",    lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "CLA Leaving Care and CWD")]),
    ("C&F Fostering & Permanence",    lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "Fostering and Kinship")]),
    ("C&F Help & Protect",            lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "Help and Protect")]),
    ("C&F Prevention & YJ",           lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "Prevention and Youth Justice")]),
    ("C&F Public Health Nursing",     lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "Public Health Nursing")]),
    ("C&F QA & Partnerships",         lambda d: d[(d["svc_group"] == "Children and Families") & (d["svc_name"] == "QA and Partnerships")]),
    ("CFE-C&F",                       lambda d: d[(d["svc_group"] == "Children and Families") & (d["Q1"] == "Children, Families and Education")]),
    ("Chld Commissioning & Perf",     lambda d: d[(d["svc_group"] == "Commissioning & Performance") & (d["Q1"] == "Children, Families and Education")]),
    ("Chld Comm BusinessOps",         lambda d: d[(d["svc_group"] == "Commissioning & Performance") & (d["svc_name"] == "Business Operations")]),
    ("Educ Curriculum & Training",    lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "Curriculum and Training")]),
    ("Educ EP & SEND",                lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "SEND Assessment and Review and Principal Educational Psychologist")]),
    ("Educ Leadership",               lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "Education Leadership")]),
    ("Educ Operations",               lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "Education Operations")]),
    ("Educ Places",                   lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "Education Places")]),
    ("Educ Virtual School",           lambda d: d[(d["svc_group"] == "Education") & (d["svc_name"] == "Virtual School")]),
    ("Finance & Procurement",         lambda d: d[d["Q1"] == "Finance and Procurement"]),
    ("Revs & Bens",                   lambda d: d[d["svc_name"] == "Revenues and Benefits"]),
    ("EcDev Skills & Climate",        lambda d: d[d["svc_group"] == "Economic Development, Skills and Climate"]),
    ("HRA Housing",                   lambda d: d[d["svc_name"] == "HRA Property"]),
    ("Infr&Trans Highways",           lambda d: d[(d["svc_name"] == "Highways") & (d["Q1"] == "Community, Place and Economy")]),
    ("Infr&Trans Infra Programmes",   lambda d: d[d["svc_name"] == "Infrastructure Programmes"]),
    ("Infr&Trans Traffic RS&Parking", lambda d: d[d["svc_name"] == "Traffic Management RS and Parking"]),
    ("Infr&Trans Transportation",     lambda d: d[d["svc_name"] == "Transportation"]),
    ("Partnership Loc & Culture",     lambda d: d[d["svc_group"] == "Partnerships, Localities and Culture"]),
    ("Planning",                      lambda d: d[d["svc_group"] == "Planning"]),
    ("Property",                      lambda d: d[d["svc_name"] == "Property"]),
    ("Reg&Ops Customer Ops",          lambda d: d[d["svc_name"] == "Customer Operations"]),
    ("Reg&Ops Lifeline & OOH",        lambda d: d[d["svc_name"] == "Lifeline & OOH"]),
    ("Reg&Ops Operations",            lambda d: d[d["svc_name"] == "Operations"]),
    ("Reg&Ops Regulatory Services",   lambda d: d[d["svc_name"] == "Regulatory Services"]),
    ("Regs&Ops Waste & Scientific",   lambda d: d[d["svc_name"].isin(["Waste", "Scientific Services"])]),
    ("Democratic Services & Gov",     lambda d: d[d["svc_name"] == "Democratic Services & Governance"]),
    ("ICT",                           lambda d: d[d["svc_group"] == "ICT"]),
    ("HR&OD",                         lambda d: d[d["svc_group"] == "HR and OD"]),
    ("Legal",                         lambda d: d[(d["svc_name"] == "Legal") & (d["Q1"] == "Resources, Strategy and Transformation")]),
    ("Public Health",                 lambda d: d[d["svc_group"] == "Public Health"]),
    ("Strategy & Performance",        lambda d: d[d["svc_group"] == "Strategy, Performance and Communications"]),
]
C1_TRAINING = {
    "AdultsCommg": 93, "AdltOps HIS & IC": 82, "AdltOps Nbhd MH & LD": 87,
    "AdltOps OT": 89, "AdltOps Social Work": 82, "Adlts Housing General Fund": 84,
    "CEO-ExecDir-ServDir": 87, "CEO & Elections": 100, "CEO-Elections": 91,
    "C&F CLA Leaving Care & CWD": 78, "C&F Fostering & Permanence": 83,
    "C&F Help & Protect": 71, "C&F Prevention & YJ": 93,
    "C&F Public Health Nursing": 96, "C&F QA & Partnerships": 83,
    "CFE-C&F": 52, "Chld Commissioning & Perf": 89, "Chld Comm BusinessOps": 99,
    "Educ Curriculum & Training": 82, "Educ EP & SEND": 80,
    "Educ Leadership": 96, "Educ Operations": 90, "Educ Places": 95,
    "Educ Virtual School": 90, "Finance & Procurement": 97, "Revs & Bens": 93,
    "EcDev Skills & Climate": 82, "HRA Housing": 96, "Infr&Trans Highways": 83,
    "Infr&Trans Infra Programmes": 90, "Infr&Trans Traffic RS&Parking": 94,
    "Infr&Trans Transportation": 37, "Partnership Loc & Culture": 92,
    "Planning": 75, "Property": 94, "Reg&Ops Customer Ops": 99,
    "Reg&Ops Lifeline & OOH": 96, "Reg&Ops Operations": 92,
    "Reg&Ops Regulatory Services": 88, "Regs&Ops Waste & Scientific": 98,
    "Democratic Services & Gov": 74, "ICT": 94, "HR&OD": 97,
    "Legal": 69, "Public Health": 99, "Strategy & Performance": 95,
}
C23_SERVICE_MAP = [
    ("Legal (RST)",               lambda d: d[(d["svc_name"] == "Legal") & (d["Q1"] == "Resources, Strategy and Transformation")]),
    ("Housing (ASH)",             lambda d: d[(d["svc_group"] == "Adults Housing General Fund") & (d["Q1"] == "Adult Services and Housing")]),
    ("Electoral Services",        lambda d: d[(d["svc_group"] == "CEO and Electoral Services") & (d["svc_name"] == "Elections")]),
    ("HRA Property (CPE)",        lambda d: d[d["svc_name"] == "HRA Property"]),
    ("Education (CFE)",           lambda d: d[(d["svc_group"] == "Education") & (d["Q1"] == "Children, Families and Education")]),
    ("Infra & Transport (CPE)",   lambda d: d[(d["svc_group"] == "Infrastructure and Transport") & (d["Q1"] == "Community, Place and Economy")]),
    ("Children & Families (CFE)", lambda d: d[(d["svc_group"] == "Children and Families") & (d["Q1"] == "Children, Families and Education")]),
    ("Adults Operations (ASH)",   lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["Q1"] == "Adult Services and Housing")]),
    ("CEO Office & Directors",    lambda d: d[d["svc_group"] == "CEO and Electoral Services"]),
    ("Property (CPE)",            lambda d: d[d["svc_name"] == "Property"]),
    ("Public Health (RST)",       lambda d: d[d["svc_group"] == "Public Health"]),
    ("Reg & Ops (CPE)",           lambda d: d[(d["svc_group"] == "Regulatory and Operations") & (d["Q1"] == "Community, Place and Economy")]),
    ("Adults Commissioning",      lambda d: d[d["svc_group"] == "Adults Commissioning"]),
    ("Democratic & Gov (RST)",    lambda d: d[d["svc_name"] == "Democratic Services & Governance"]),
    ("EcDev (CPE)",               lambda d: d[d["svc_group"] == "Economic Development, Skills and Climate"]),
    ("HR&OD (RST)",               lambda d: d[d["svc_group"] == "HR and OD"]),
    ("Planning (CPE)",            lambda d: d[d["svc_group"] == "Planning"]),
    ("Finance & Procurement",     lambda d: d[d["Q1"] == "Finance and Procurement"]),
    ("Partnership Loc & Culture", lambda d: d[d["svc_group"] == "Partnerships, Localities and Culture"]),
    ("Commissioning & Perf (CFE)",lambda d: d[(d["svc_group"] == "Commissioning & Performance") & (d["Q1"] == "Children, Families and Education")]),
    ("ICT (RST)",                 lambda d: d[d["svc_group"] == "ICT"]),
    ("Strategy & Comms (RST)",    lambda d: d[d["svc_group"] == "Strategy, Performance and Communications"]),
]
C2_SICKNESS = {
    "Legal (RST)": 7.77, "Housing (ASH)": 6.83, "Electoral Services": 6.41,
    "HRA Property (CPE)": 6.08, "Education (CFE)": 5.98, "Infra & Transport (CPE)": 5.75,
    "Children & Families (CFE)": 5.62, "Adults Operations (ASH)": 5.21,
    "CEO Office & Directors": 4.98, "Property (CPE)": 4.92,
    "Public Health (RST)": 4.90, "Reg & Ops (CPE)": 4.89,
    "Adults Commissioning": 4.79, "Democratic & Gov (RST)": 4.72,
    "EcDev (CPE)": 4.65, "HR&OD (RST)": 4.45, "Planning (CPE)": 4.44,
    "Finance & Procurement": 4.39, "Partnership Loc & Culture": 4.31,
    "Commissioning & Perf (CFE)": 4.31, "ICT (RST)": 4.21,
    "Strategy & Comms (RST)": 3.64,
}
C3_TURNOVER = {
    "Legal (RST)": 1.20, "Housing (ASH)": 0.67, "Electoral Services": None,
    "HRA Property (CPE)": 0.72, "Education (CFE)": 1.92, "Infra & Transport (CPE)": 0.68,
    "Children & Families (CFE)": 0.81, "Adults Operations (ASH)": 0.93,
    "CEO Office & Directors": 1.21, "Property (CPE)": 0.76,
    "Public Health (RST)": 0.63, "Reg & Ops (CPE)": 0.99,
    "Adults Commissioning": None, "Democratic & Gov (RST)": 1.52,
    "EcDev (CPE)": 1.04, "HR&OD (RST)": 0.36, "Planning (CPE)": 0.46,
    "Finance & Procurement": 0.46, "Partnership Loc & Culture": 0.51,
    "Commissioning & Perf (CFE)": 0.55, "ICT (RST)": 0.41,
    "Strategy & Comms (RST)": 0.35,
}
C4_SERVICE_MAP = [
    ("C&F",                       lambda d: d[(d["svc_group"] == "Children and Families") & (d["Q1"] == "Children, Families and Education")]),
    ("Adults Services (ASH)",     lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["Q1"] == "Adult Services and Housing")]),
    ("HRA (ASH)",                 lambda d: d[(d["svc_group"] == "Adults Housing General Fund") & (d["Q1"] == "Adult Services and Housing")]),
    ("Housing (ASH)",             lambda d: d[(d["svc_group"] == "Adults Housing General Fund") & (d["Q1"] == "Adult Services and Housing")]),
    ("Infra & Transport (CPE)",   lambda d: d[(d["svc_group"] == "Infrastructure and Transport") & (d["Q1"] == "Community, Place and Economy")]),
    ("Planning (CPE)",            lambda d: d[d["svc_group"] == "Planning"]),
    ("Climate (CPE)",             lambda d: d[(d["svc_name"] == "Climate and Natural Environment")]),
    ("Reg & Ops (CPE)",           lambda d: d[(d["svc_group"] == "Regulatory and Operations") & (d["Q1"] == "Community, Place and Economy")]),
    ("Partnership & Loc (CPE)",   lambda d: d[d["svc_group"] == "Partnerships, Localities and Culture"]),
    ("EcDev Skills & Climate",    lambda d: d[d["svc_group"] == "Economic Development, Skills and Climate"]),
    ("HR&OD (RST)",               lambda d: d[d["svc_group"] == "HR and OD"]),
    ("Digital Services (RST)",    lambda d: d[d["svc_group"] == "ICT"]),
    ("Strategy & Comms (RST)",    lambda d: d[d["svc_group"] == "Strategy, Performance and Communications"]),
    ("Public Health (RST)",       lambda d: d[d["svc_group"] == "Public Health"]),
    ("Procurement (CFO)",         lambda d: d[d["svc_name"] == "Procurement"]),
    ("Finance (CFO)",             lambda d: d[(d["Q1"] == "Finance and Procurement") & (d["svc_name"] != "Procurement")]),
    ("CEO Office",                lambda d: d[d["svc_group"] == "CEO and Electoral Services"]),
]
C4_PROCUREMENT = {
    "C&F": 6, "Adults Services (ASH)": 20, "HRA (ASH)": 0, "Housing (ASH)": 23,
    "Infra & Transport (CPE)": 11, "Planning (CPE)": 8, "Climate (CPE)": 6,
    "Reg & Ops (CPE)": 5, "Partnership & Loc (CPE)": 0, "EcDev Skills & Climate": 3,
    "HR&OD (RST)": 0, "Digital Services (RST)": 4, "Strategy & Comms (RST)": 0,
    "Public Health (RST)": 0, "Procurement (CFO)": 0, "Finance (CFO)": 1,
    "CEO Office": 0,
}
C5_SERVICE_MAP = [
    ("Adults Commissioning",      lambda d: d[d["svc_group"] == "Adults Commissioning"]),
    ("Adults Operations",         lambda d: d[(d["svc_group"] == "Adults Health & Operations") & (d["Q1"] == "Adult Services and Housing")]),
    ("Housing",                   lambda d: d[(d["svc_group"] == "Adults Housing General Fund") & (d["Q1"] == "Adult Services and Housing")]),
    ("Electoral Services",        lambda d: d[(d["svc_group"] == "CEO and Electoral Services") & (d["svc_name"] == "Elections")]),
    ("CEO Office & Directors",    lambda d: d[d["svc_group"] == "CEO and Electoral Services"]),
    ("Children & Families",       lambda d: d[(d["svc_group"] == "Children and Families") & (d["Q1"] == "Children, Families and Education")]),
    ("Commissioning & Perf",      lambda d: d[(d["svc_group"] == "Commissioning & Performance") & (d["Q1"] == "Children, Families and Education")]),
    ("Education",                 lambda d: d[(d["svc_group"] == "Education") & (d["Q1"] == "Children, Families and Education")]),
    ("EcDev Skills & Climate",    lambda d: d[d["svc_group"] == "Economic Development, Skills and Climate"]),
    ("HRA & Property",            lambda d: d[d["svc_name"].isin(["HRA Property", "Property"])]),
    ("Infra & Transport",         lambda d: d[(d["svc_group"] == "Infrastructure and Transport") & (d["Q1"] == "Community, Place and Economy")]),
    ("Partnership Loc & Culture", lambda d: d[d["svc_group"] == "Partnerships, Localities and Culture"]),
    ("Planning",                  lambda d: d[d["svc_group"] == "Planning"]),
    ("Reg & Ops",                 lambda d: d[(d["svc_group"] == "Regulatory and Operations") & (d["Q1"] == "Community, Place and Economy")]),
    ("Democratic & Gov",          lambda d: d[d["svc_name"] == "Democratic Services & Governance"]),
    ("ICT Services",              lambda d: d[d["svc_group"] == "ICT"]),
    ("HR&OD",                     lambda d: d[d["svc_group"] == "HR and OD"]),
    ("Legal",                     lambda d: d[(d["svc_name"] == "Legal") & (d["Q1"] == "Resources, Strategy and Transformation")]),
    ("Public Health",             lambda d: d[d["svc_group"] == "Public Health"]),
    ("Strategy & Comms",          lambda d: d[d["svc_group"] == "Strategy, Performance and Communications"]),
    ("Finance & Procurement",     lambda d: d[d["Q1"] == "Finance and Procurement"]),
]
C5_COMPLETION = {
    "Adults Commissioning": 91.4, "Adults Operations": 34.1, "Housing": 45.8,
    "Electoral Services": 33.3, "CEO Office & Directors": 155.0,
    "Children & Families": 22.2, "Commissioning & Perf": 32.9,
    "Education": 36.7, "EcDev Skills & Climate": 42.4, "HRA & Property": 22.9,
    "Infra & Transport": 43.5, "Partnership Loc & Culture": 52.1,
    "Planning": 51.9, "Reg & Ops": 34.7, "Democratic & Gov": 59.1,
    "ICT Services": 38.4, "HR&OD": 72.0, "Legal": 47.1,
    "Public Health": 47.0, "Strategy & Comms": 76.0,
    "Finance & Procurement": 43.9,
}


def _precompute_service(df, service_map):
    """Apply service map filters, sort by EE score, return aggregated values."""
    rows = [(lbl, fn(df)) for lbl, fn in service_map]
    rows = [(lbl, sub) for lbl, sub in rows if len(sub) > 0]
    labels = [r[0] for r in rows]
    ee_pre = [sub[OUTCOME_COLS].mean(axis=1).mean() for _, sub in rows]
    idx = sorted(range(len(labels)), key=lambda i: ee_pre[i])
    rows = [rows[i] for i in idx]
    labels = [labels[i] for i in idx]
    ee_scores = [ee_pre[i] for i in idx]
    wow_p = {c: [sub[c].mean() for _, sub in rows] for c in WOW_PLACE_COLS}
    wow_i = {c: [sub[c].mean() for _, sub in rows] for c in WOW_IND_COLS}
    return labels, ee_scores, wow_p, wow_i


@st.cache_data(show_spinner=False)
def _precompute_c1(df):
    return _precompute_service(df, C1_SERVICE_MAP)


@st.cache_data(show_spinner=False)
def _precompute_c23(df):
    return _precompute_service(df, C23_SERVICE_MAP)


@st.cache_data(show_spinner=False)
def _precompute_c4(df):
    return _precompute_service(df, C4_SERVICE_MAP)


@st.cache_data(show_spinner=False)
def _precompute_c5(df):
    return _precompute_service(df, C5_SERVICE_MAP)


with _timed("Section C render"), sec_c:
    c1, c2, c3, c4, c5 = st.tabs([
        "C1: Mandatory Training",
        "C2: Sickness",
        "C3: Turnover",
        "C4: Procurement",
        "C5: Survey Completion",
    ])
    with c1:
        c1_1, c1_2 = st.tabs([
            "C1.1: Employee Experience",
            "C1.2: Ways of Working",
        ])

        _c1_labels, _c1_ee_scores, _c1_wow_p, _c1_wow_i = _precompute_c1(filtered)
        _c1_training = [C1_TRAINING.get(lbl) for lbl in _c1_labels]

        def _c1_dual_chart(y_vals, y_title, chart_key, caption_text, x_labels=None, comp_override=None):
            _x = x_labels if x_labels is not None else _c1_labels
            _comp = comp_override if comp_override is not None else _c1_training
            fig = go.Figure()
            fig.add_trace(go.Bar(
                name=y_title, x=_x, y=y_vals,
                text=[f"{v:.2f}" if v is not None else "" for v in y_vals],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color=PRIMARY, yaxis="y1", offsetgroup=0,
                hovertemplate="<b>%{x}</b><br>" + y_title + ": %{y:.2f}<extra></extra>",
            ))
            fig.add_trace(go.Bar(
                name="Training Completion %", x=_x, y=_comp,
                text=[f"{t}%" if t is not None else "" for t in _comp],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color="#E07070", yaxis="y2", offsetgroup=1,
                hovertemplate="<b>%{x}</b><br>Training completion: %{y}%<extra></extra>",
            ))
            fig.update_layout(
                barmode="group",
                font=dict(family="Inter", color="#1A2B3C"),
                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                margin=dict(l=10, r=60, t=40, b=200),
                xaxis=dict(tickangle=-45, tickfont=dict(size=9, color="#1A2B3C")),
                yaxis=dict(
                    title=dict(text=y_title, font=dict(color=PRIMARY, size=10)),
                    range=[-0.2, 4.5], tickfont=dict(size=10, color=PRIMARY), gridcolor="#E8EEF2",
                ),
                yaxis2=dict(
                    title=dict(text="Training Completion (%)", font=dict(color="#C0392B", size=10)),
                    overlaying="y", side="right", range=[0, 130],
                    tickfont=dict(size=10, color="#C0392B"), showgrid=False, ticksuffix="%",
                ),
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                            font=dict(color="#1A2B3C", size=11)),
                height=600,
            )
            st.caption(caption_text)
            st.plotly_chart(fig, use_container_width=True, key=chart_key)

        with c1_1:
            st.markdown("#### Employee Experience vs Training Completion — by Service")
            _c1_dual_chart(
                _c1_ee_scores,
                "Avg EE Score (0 = Strongly Disagree → 4 = Strongly Agree)",
                "c1_1_bar",
                "Blue bars = average employee experience score across all 15 outcome questions (left axis, higher = more positive). Red bars = mandatory training completion rate (right axis).",
            )

        with c1_2:
            st.markdown("#### Ways of Working vs Training Completion — by Service")
            for _wi, _wt in enumerate(WOW_THEMES):
                _wc = WOW_PLACE_COLS[_wi]
                _ws = WOW_PLACE_STATEMENTS.get(_wt, "")
                st.markdown(f"#### {_wt}")
                if _ws:
                    st.caption(f'"{_ws}"')
                _wv = _c1_wow_p[_wc]
                _wsort = sorted(range(len(_c1_labels)), key=lambda i: _wv[i])
                _c1_dual_chart(
                    [_wv[i] for i in _wsort],
                    f"{_wt} (Place) — Avg Score",
                    f"c1_2_bar_{_wi}",
                    f"Blue = average score for '{_wt}' (Place). Red = mandatory training completion rate.",
                    x_labels=[_c1_labels[i] for i in _wsort],
                    comp_override=[_c1_training[i] for i in _wsort],
                )
    with c2:
        c2_1, c2_2 = st.tabs(["C2.1: Employee Experience", "C2.2: Ways of Working"])

        _c23_labels, _c23_ee_scores, _c23_wow_p, _c23_wow_i = _precompute_c23(filtered)
        _c2_comp = [C2_SICKNESS.get(lbl) for lbl in _c23_labels]

        def _c23_dual_chart(y_vals, y_title, comp_vals, comp_title, comp_suffix, comp_color, y_range, comp_range, chart_key, caption_text, x_labels=None):
            _x = x_labels if x_labels is not None else _c23_labels
            fig = go.Figure()
            fig.add_trace(go.Bar(
                name=y_title, x=_x, y=y_vals,
                text=[f"{v:.2f}" if v is not None else "" for v in y_vals],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color=PRIMARY, yaxis="y1", offsetgroup=0,
                hovertemplate="<b>%{x}</b><br>" + y_title + ": %{y:.2f}<extra></extra>",
            ))
            fig.add_trace(go.Bar(
                name=comp_title, x=_x, y=comp_vals,
                text=[f"{v}{comp_suffix}" if v is not None else "" for v in comp_vals],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color=comp_color, yaxis="y2", offsetgroup=1,
                hovertemplate="<b>%{x}</b><br>" + comp_title + ": %{y}" + comp_suffix + "<extra></extra>",
            ))
            fig.update_layout(
                barmode="group",
                font=dict(family="Inter", color="#1A2B3C"),
                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                margin=dict(l=10, r=60, t=40, b=200),
                xaxis=dict(tickangle=-45, tickfont=dict(size=9, color="#1A2B3C")),
                yaxis=dict(
                    title=dict(text=y_title, font=dict(color=PRIMARY, size=10)),
                    range=y_range, tickfont=dict(size=10, color=PRIMARY), gridcolor="#E8EEF2",
                ),
                yaxis2=dict(
                    title=dict(text=comp_title, font=dict(color=comp_color, size=10)),
                    overlaying="y", side="right", range=comp_range,
                    tickfont=dict(size=10, color=comp_color), showgrid=False,
                    ticksuffix=comp_suffix,
                ),
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                            font=dict(color="#1A2B3C", size=11)),
                height=600,
            )
            st.caption(caption_text)
            st.plotly_chart(fig, use_container_width=True, key=chart_key)

        with c2_1:
            st.markdown("#### Employee Experience vs Sickness Absence — by Service")
            _c23_dual_chart(
                _c23_ee_scores, "Avg EE Score (0 = Strongly Disagree → 4 = Strongly Agree)",
                _c2_comp, "Avg Sickness Days", " days", "#E07070",
                [-0.2, 4.5], [0, 10],
                "c2_1_bar",
                "Blue = average EE score (left axis, higher = more positive). Red = average sickness absence days (right axis).",
            )

        with c2_2:
            st.markdown("#### Ways of Working vs Sickness Absence — by Service")
            for _wi, _wt in enumerate(WOW_THEMES):
                _wc = WOW_PLACE_COLS[_wi]
                _ws = WOW_PLACE_STATEMENTS.get(_wt, "")
                st.markdown(f"#### {_wt}")
                if _ws:
                    st.caption(f'"{_ws}"')
                _wv = _c23_wow_p[_wc]
                _wsort = sorted(range(len(_c23_labels)), key=lambda i: _wv[i])
                _c23_dual_chart(
                    [_wv[i] for i in _wsort],
                    f"{_wt} (Place) — Avg Score",
                    [_c2_comp[i] for i in _wsort], "Avg Sickness Days", " days", "#E07070",
                    [-0.2, 4.5], [0, 10],
                    f"c2_2_bar_{_wi}",
                    f"Blue = average score for '{_wt}' (Place). Red = average sickness absence days.",
                    x_labels=[_c23_labels[i] for i in _wsort],
                )
    with c3:
        c3_1, c3_2 = st.tabs(["C3.1: Employee Experience", "C3.2: Ways of Working"])

        C3_TURNOVER = {
            "Legal (RST)": 1.20, "Housing (ASH)": 0.67, "Electoral Services": None,
            "HRA Property (CPE)": 0.72, "Education (CFE)": 1.92, "Infra & Transport (CPE)": 0.68,
            "Children & Families (CFE)": 0.81, "Adults Operations (ASH)": 0.93,
            "CEO Office & Directors": 1.21, "Property (CPE)": 0.76,
            "Public Health (RST)": 0.63, "Reg & Ops (CPE)": 0.99,
            "Adults Commissioning": None, "Democratic & Gov (RST)": 1.52,
            "EcDev (CPE)": 1.04, "HR&OD (RST)": 0.36, "Planning (CPE)": 0.46,
            "Finance & Procurement": 0.46, "Partnership Loc & Culture": 0.51,
            "Commissioning & Perf (CFE)": 0.55, "ICT (RST)": 0.41,
            "Strategy & Comms (RST)": 0.35,
        }

        _c3_comp = [C3_TURNOVER.get(lbl) for lbl in _c23_labels]

        with c3_1:
            st.markdown("#### Employee Experience vs Turnover — by Service")
            _c23_dual_chart(
                _c23_ee_scores, "Avg EE Score (0 = Strongly Disagree → 4 = Strongly Agree)",
                _c3_comp, "Turnover (%)", "%", "#E07070",
                [-0.2, 4.5], [0, 3],
                "c3_1_bar",
                "Blue = average EE score (left axis, higher = more positive). Red = monthly turnover rate (right axis). No bar shown where data unavailable.",
            )

        with c3_2:
            st.markdown("#### Ways of Working vs Turnover — by Service")
            for _wi, _wt in enumerate(WOW_THEMES):
                _wc = WOW_PLACE_COLS[_wi]
                _ws = WOW_PLACE_STATEMENTS.get(_wt, "")
                st.markdown(f"#### {_wt}")
                if _ws:
                    st.caption(f'"{_ws}"')
                _wv = _c23_wow_p[_wc]
                _wsort = sorted(range(len(_c23_labels)), key=lambda i: _wv[i])
                _c23_dual_chart(
                    [_wv[i] for i in _wsort],
                    f"{_wt} (Place) — Avg Score",
                    [_c3_comp[i] for i in _wsort], "Turnover (%)", "%", "#E07070",
                    [-0.2, 4.5], [0, 3],
                    f"c3_2_bar_{_wi}",
                    f"Blue = average score for '{_wt}' (Place). Red = monthly turnover rate.",
                    x_labels=[_c23_labels[i] for i in _wsort],
                )
    with c4:
        c4_1, c4_2 = st.tabs(["C4.1: Employee Experience", "C4.2: Ways of Working"])

        _c4_labels, _c4_ee_scores, _c4_wow_p, _c4_wow_i = _precompute_c4(filtered)
        _c4_comp = [C4_PROCUREMENT.get(lbl) for lbl in _c4_labels]

        def _c4_dual_chart(y_vals, y_title, chart_key, caption_text, x_labels=None, comp_override=None):
            _x = x_labels if x_labels is not None else _c4_labels
            _comp = comp_override if comp_override is not None else _c4_comp
            fig = go.Figure()
            fig.add_trace(go.Bar(
                name=y_title, x=_x, y=y_vals,
                text=[f"{v:.2f}" if v is not None else "" for v in y_vals],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color=PRIMARY, yaxis="y1", offsetgroup=0,
                hovertemplate="<b>%{x}</b><br>" + y_title + ": %{y:.2f}<extra></extra>",
            ))
            fig.add_trace(go.Bar(
                name="No. of Procurement Breaches", x=_x, y=_comp,
                text=[str(v) if v is not None else "" for v in _comp],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color="#E07070", yaxis="y2", offsetgroup=1,
                hovertemplate="<b>%{x}</b><br>Procurement breaches: %{y}<extra></extra>",
            ))
            fig.update_layout(
                barmode="group",
                font=dict(family="Inter", color="#1A2B3C"),
                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                margin=dict(l=10, r=60, t=40, b=200),
                xaxis=dict(tickangle=-45, tickfont=dict(size=9, color="#1A2B3C")),
                yaxis=dict(
                    title=dict(text=y_title, font=dict(color=PRIMARY, size=10)),
                    range=[-0.2, 4.5], tickfont=dict(size=10, color=PRIMARY), gridcolor="#E8EEF2",
                ),
                yaxis2=dict(
                    title=dict(text="No. of Procurement Breaches", font=dict(color="#C0392B", size=10)),
                    overlaying="y", side="right", range=[0, 30],
                    tickfont=dict(size=10, color="#C0392B"), showgrid=False,
                ),
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                            font=dict(color="#1A2B3C", size=11)),
                height=600,
            )
            st.caption(caption_text)
            st.plotly_chart(fig, use_container_width=True, key=chart_key)

        with c4_1:
            st.markdown("#### Employee Experience vs Procurement Breaches — by Service")
            _c4_dual_chart(
                _c4_ee_scores,
                "Avg EE Score (0 = Strongly Disagree → 4 = Strongly Agree)",
                "c4_1_bar",
                "Blue = average EE score (left axis, higher = more positive). Red = procurement value (right axis). Note: 'Strategic Asset Management' excluded — no survey respondents found.",
            )

        with c4_2:
            st.markdown("#### Ways of Working vs Procurement Breaches — by Service")
            for _wi, _wt in enumerate(WOW_THEMES):
                _wc = WOW_PLACE_COLS[_wi]
                _ws = WOW_PLACE_STATEMENTS.get(_wt, "")
                st.markdown(f"#### {_wt}")
                if _ws:
                    st.caption(f'"{_ws}"')
                _wv = _c4_wow_p[_wc]
                _wsort = sorted(range(len(_c4_labels)), key=lambda i: _wv[i])
                _c4_dual_chart(
                    [_wv[i] for i in _wsort],
                    f"{_wt} (Place) — Avg Score",
                    f"c4_2_bar_{_wi}",
                    f"Blue = average score for '{_wt}' (Place). Red = procurement breaches.",
                    x_labels=[_c4_labels[i] for i in _wsort],
                    comp_override=[_c4_comp[i] for i in _wsort],
                )
    with c5:
        c5_1, c5_2 = st.tabs(["C5.1: Employee Experience", "C5.2: Ways of Working"])

        _c5_labels, _c5_ee_scores, _c5_wow_p, _c5_wow_i = _precompute_c5(filtered)
        _c5_comp = [C5_COMPLETION.get(lbl) for lbl in _c5_labels]

        def _c5_dual_chart(y_vals, y_title, chart_key, caption_text, x_labels=None, comp_override=None):
            _x = x_labels if x_labels is not None else _c5_labels
            _comp = comp_override if comp_override is not None else _c5_comp
            fig = go.Figure()
            fig.add_trace(go.Bar(
                name=y_title, x=_x, y=y_vals,
                text=[f"{v:.2f}" if v is not None else "" for v in y_vals],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color=PRIMARY, yaxis="y1", offsetgroup=0,
                hovertemplate="<b>%{x}</b><br>" + y_title + ": %{y:.2f}<extra></extra>",
            ))
            fig.add_trace(go.Bar(
                name="Survey Completion %", x=_x, y=_comp,
                text=[f"{v:.1f}%" if v is not None else "" for v in _comp],
                textposition="outside", textfont=dict(color="#1A2B3C", size=8),
                marker_color="#E07070", yaxis="y2", offsetgroup=1,
                hovertemplate="<b>%{x}</b><br>Survey completion: %{y:.1f}%<extra></extra>",
            ))
            fig.update_layout(
                barmode="group",
                font=dict(family="Inter", color="#1A2B3C"),
                paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                margin=dict(l=10, r=60, t=40, b=200),
                xaxis=dict(tickangle=-45, tickfont=dict(size=9, color="#1A2B3C")),
                yaxis=dict(
                    title=dict(text=y_title, font=dict(color=PRIMARY, size=10)),
                    range=[-0.2, 4.5], tickfont=dict(size=10, color=PRIMARY), gridcolor="#E8EEF2",
                ),
                yaxis2=dict(
                    title=dict(text="Survey Completion (%)", font=dict(color="#C0392B", size=10)),
                    overlaying="y", side="right", range=[0, 170],
                    tickfont=dict(size=10, color="#C0392B"), showgrid=False, ticksuffix="%",
                ),
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
                            font=dict(color="#1A2B3C", size=11)),
                height=600,
            )
            st.caption(caption_text)
            st.plotly_chart(fig, use_container_width=True, key=chart_key)

        with c5_1:
            st.markdown("#### Employee Experience vs Survey Completion — by Service")
            _c5_dual_chart(
                _c5_ee_scores,
                "Avg EE Score (0 = Strongly Disagree → 4 = Strongly Agree)",
                "c5_1_bar",
                "Blue = average EE score (left axis, higher = more positive). Red = survey completion rate (right axis). Note: 'Commercial and Investment' excluded — no survey respondents. CEO Office & Directors shows >100% as headcount denominator may be understated.",
            )

        with c5_2:
            st.markdown("#### Ways of Working vs Survey Completion — by Service")
            for _wi, _wt in enumerate(WOW_THEMES):
                _wc = WOW_PLACE_COLS[_wi]
                _ws = WOW_PLACE_STATEMENTS.get(_wt, "")
                st.markdown(f"#### {_wt}")
                if _ws:
                    st.caption(f'"{_ws}"')
                _wv = _c5_wow_p[_wc]
                _wsort = sorted(range(len(_c5_labels)), key=lambda i: _wv[i])
                _c5_dual_chart(
                    [_wv[i] for i in _wsort],
                    f"{_wt} (Place) — Avg Score",
                    f"c5_2_bar_{_wi}",
                    f"Blue = average score for '{_wt}' (Place). Red = survey completion rate.",
                    x_labels=[_c5_labels[i] for i in _wsort],
                    comp_override=[_c5_comp[i] for i in _wsort],
                )


# ═══════════════════════════════════════════════════════════════════════════════════
# SECTION D
# ═══════════════════════════════════════════════════════════════════════════════════
with _timed("Section D render"), sec_d:
    d1, = st.tabs(["D1: Descriptive Analysis"])

    with d1:
        if not st.session_state.get('d1_loaded'):
            st.info(
                'EDI views are not loaded yet — click below to load them. '
                'This keeps the rest of the app fast.'
            )
            if st.button('▶ Load EDI Views', type='primary', key='load_d1_btn'):
                st.session_state['d1_loaded'] = True
                st.rerun()
        else:
            d1_1, d1_2 = st.tabs([
                "D1.1: Ways of Working",
                "D1.2: Sentiment Outcomes",
            ])

            # ── Shared chart builder ──────────────────────────────────────────
            def _edi_bar_chart(x_labels, overall_vals, filtered_vals, n_overall, n_filtered, y_title, chart_key):
                """Returns a Plotly figure for EDI comparison."""
                fig = go.Figure()
                fig.add_trace(go.Bar(
                    name=f"Council overall (n={n_overall:,})",
                    x=x_labels, y=overall_vals,
                    marker_color="#8DC0D4",
                    hovertemplate="<b>%{x}</b><br>Council overall: %{y:.2f}<extra></extra>",
                ))
                fig.add_trace(go.Bar(
                    name=f"Selected EDI group (n={n_filtered:,})",
                    x=x_labels, y=filtered_vals,
                    marker_color=PRIMARY,
                    hovertemplate="<b>%{x}</b><br>Selected group: %{y:.2f}<extra></extra>",
                ))
                fig.update_layout(
                    barmode="group",
                    font=dict(family="Inter", color="#1A2B3C"),
                    paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                    margin=dict(l=10, r=10, t=40, b=180),
                    xaxis=dict(tickangle=-45, tickfont=dict(size=10, color="#1A2B3C")),
                    yaxis=dict(
                        title=dict(text=y_title, font=dict(color="#1A2B3C", size=11)),
                        range=[-0.2, 4.5], tickfont=dict(size=11, color="#1A2B3C"), gridcolor="#E8EEF2",
                    ),
                    legend=dict(orientation="h", yanchor="bottom", y=1.02,
                                xanchor="right", x=1, font=dict(color="#1A2B3C", size=12)),
                    height=520,
                )
                return fig

            # Pre-build all figures (needed for both display and report)
            _n_all   = len(df.dropna(subset=WOW_PLACE_COLS, how="all"))
            _n_filt  = len(filtered.dropna(subset=WOW_PLACE_COLS, how="all"))
            _n_out_all  = len(df.dropna(subset=OUTCOME_COLS, how="all"))
            _n_out_filt = len(filtered.dropna(subset=OUTCOME_COLS, how="all"))
            _y_title = "Average Score (0 = Strongly Disagree, 4 = Strongly Agree)"

            _fig_place = _edi_bar_chart(
                WOW_THEMES,
                [df[c].mean() for c in WOW_PLACE_COLS],
                [filtered[c].mean() for c in WOW_PLACE_COLS],
                _n_all, _n_filt, _y_title, "d1_1_place_fig"
            )
            _fig_ind = _edi_bar_chart(
                WOW_THEMES,
                [df[c].mean() for c in WOW_IND_COLS],
                [filtered[c].mean() for c in WOW_IND_COLS],
                _n_all, _n_filt, _y_title, "d1_1_ind_fig"
            )
            _fig_outcomes = _edi_bar_chart(
                OUTCOME_LABELS,
                [df[c].mean() for c in OUTCOME_COLS],
                [filtered[c].mean() for c in OUTCOME_COLS],
                _n_out_all, _n_out_filt, _y_title, "d1_2_fig"
            )

            # ── D1.1: Ways of Working ─────────────────────────────────────────
            with d1_1:
                st.caption(
                    "For each Way of Working theme, shows the average score for the council overall "
                    "and for every group within the selected EDI filter. Scroll down to see all 22 themes."
                )
                st.markdown(
                    '<p style="font-size:11px;color:#5A7080;margin-bottom:8px">'
                    '<strong>Scale:</strong> 0 = Strongly disagree &nbsp;·&nbsp; 1 = Disagree &nbsp;·&nbsp; '
                    '2 = Neither agree nor disagree &nbsp;·&nbsp; 3 = Agree &nbsp;·&nbsp; 4 = Strongly agree</p>',
                    unsafe_allow_html=True,
                )

                _d1_edi_tabs = st.tabs([
                    "Length of Service", "Age", "Gender",
                    "Ethnic Group", "Sexual Orientation", "Disabled", "Carer",
                ])
                _d1_edi_configs = [
                    (["Q8"],   "Length of Service"),
                    (["Q72"],  "Age"),
                    (["Q73"],  "Gender"),
                    (["Q74"],  "Ethnic Group"),
                    (["Q75"],  "Sexual Orientation"),
                    (["Q76"],  "Disabled"),
                    (["Q77"],  "Carer"),
                ]

                def _make_wow_docx(edi_cols, edi_display):
                    """Build and return WoW Word report bytes for the given EDI group."""
                    import io as _io, copy as _copy
                    from docx import Document as _Document
                    from docx.shared import Inches as _Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
                    import plotly.io as _pio
                    _x = ["Council Overall"]
                    _p_ov = [df[c].mean() for c in WOW_PLACE_COLS]
                    _i_ov = [df[c].mean() for c in WOW_IND_COLS]
                    _gp, _gi, _grps = [], [], []
                    for eq in edi_cols:
                        _vals = sorted(
                            (v for v in df[eq].dropna().unique()
                             if str(v).strip() not in ("", "Not Answered")),
                            key=lambda v: LOS_ORDER.index(v) if eq == "Q8" and v in LOS_ORDER else v,
                        )
                        _pfx = (EDI_FILTERS.get(eq, eq) + ": ") if len(edi_cols) > 1 else ""
                        for val in _vals:
                            _sub = df[df[eq] == val]
                            _x.append(f"{_pfx}{val}")
                            _gp.append([_sub[c].mean() for c in WOW_PLACE_COLS])
                            _gi.append([_sub[c].mean() for c in WOW_IND_COLS])
                            _grps.append((f"{_pfx}{val}", eq, val))
                    _tot = len(df)
                    _parts = [f"Council Overall: {_tot:,} respondents (100%)"]
                    for _bl, eq, val in _grps:
                        _n = int((df[eq] == val).sum())
                        _parts.append(f"{_bl}: {_n:,} ({_n / _tot:.0%})")
                    _nb = len(_x)
                    _pc = ["#8DC0D4"] + [PRIMARY] * (_nb - 1)
                    _ic = ["#F0B0B0"] + ["#E07070"] * (_nb - 1)
                    _doc = _Document()
                    _doc.core_properties.title = f"Ways of Working — {edi_display}"
                    _h = _doc.add_heading(f"Ways of Working — {edi_display}", 0)
                    _h.alignment = _ALIGN.CENTER
                    _sh = _doc.add_paragraph("Somerset Council Cultural Diagnostic — D1.1")
                    _sh.alignment = _ALIGN.CENTER
                    _doc.add_paragraph()
                    _doc.add_paragraph(
                        "Scale: 0 = Strongly disagree · 1 = Disagree · "
                        "2 = Neither agree nor disagree · 3 = Agree · 4 = Strongly agree"
                    )
                    _doc.add_paragraph()
                    _doc.add_heading("Sample Breakdown", 2)
                    for _sp in _parts:
                        _doc.add_paragraph(_sp, style="List Bullet")
                    _doc.add_page_break()
                    for _wi, _theme in enumerate(WOW_THEMES):
                        _py = [_p_ov[_wi]] + [g[_wi] for g in _gp]
                        _iy = [_i_ov[_wi]] + [g[_wi] for g in _gi]
                        _doc.add_heading(_theme, 2)
                        _ps = WOW_PLACE_STATEMENTS.get(_theme, "")
                        _is = WOW_IND_STATEMENTS.get(_theme, "")
                        if _ps:
                            _doc.add_paragraph().add_run(f"Place: \"{_ps}\"").italic = True
                        if _is:
                            _doc.add_paragraph().add_run(f"Individual: \"{_is}\"").italic = True
                        _f = go.Figure()
                        _f.add_trace(go.Bar(name="Place (P)", x=_x, y=_py, marker_color=_pc,
                                            text=[f"{v:.2f}" for v in _py], textposition="outside",
                                            textfont=dict(size=12, color="#1A2B3C"),
                                            hovertemplate="<b>%{x}</b><br>Place: %{y:.2f}<extra></extra>"))
                        _f.add_trace(go.Bar(name="Individual (I)", x=_x, y=_iy, marker_color=_ic,
                                            text=[f"{v:.2f}" for v in _iy], textposition="outside",
                                            textfont=dict(size=12, color="#1A2B3C"),
                                            hovertemplate="<b>%{x}</b><br>Individual: %{y:.2f}<extra></extra>"))
                        _f.update_layout(
                            barmode="group",
                            title=dict(text=_theme, font=dict(color="#0F4C6B", size=18)),
                            font=dict(family="Inter", color="#1A2B3C", size=14),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=100, r=30, t=80, b=160),
                            xaxis=dict(tickangle=-35, tickfont=dict(size=13, color="#1A2B3C")),
                            yaxis=dict(range=[-0.2, 4.5],
                                       title=dict(text="Avg Score", font=dict(color="#1A2B3C", size=13)),
                                       tickfont=dict(size=13, color="#1A2B3C"), gridcolor="#E8EEF2"),
                            legend=dict(orientation="h", yanchor="bottom", y=1.02,
                                        xanchor="right", x=1, font=dict(color="#1A2B3C", size=13)),
                            height=520,
                        )
                        _img = _pio.to_image(_f, format="png", width=1200, height=520, scale=2)
                        _doc.add_picture(_io.BytesIO(_img), width=_Inches(6.5))
                        _doc.add_paragraph()
                    _buf = _io.BytesIO()
                    _doc.save(_buf)
                    return _buf.getvalue()

                def _wow_edi_charts(edi_cols, edi_display, key_prefix):
                    """Render 22 bar charts (one per WoW theme) with per-tab report download."""
                    _x_labels = ["Council Overall"]
                    _p_overall = [df[c].mean() for c in WOW_PLACE_COLS]
                    _i_overall = [df[c].mean() for c in WOW_IND_COLS]
                    _group_p, _group_i, _groups = [], [], []
                    for eq in edi_cols:
                        _eq_vals = sorted(
                            (v for v in df[eq].dropna().unique()
                             if str(v).strip() not in ("", "Not Answered")),
                            key=lambda v: LOS_ORDER.index(v) if eq == "Q8" and v in LOS_ORDER else v,
                        )
                        _prefix = (EDI_FILTERS.get(eq, eq) + ": ") if len(edi_cols) > 1 else ""
                        for val in _eq_vals:
                            _sub = df[df[eq] == val]
                            _x_labels.append(f"{_prefix}{val}")
                            _group_p.append([_sub[c].mean() for c in WOW_PLACE_COLS])
                            _group_i.append([_sub[c].mean() for c in WOW_IND_COLS])
                            _groups.append((f"{_prefix}{val}", eq, val))
                    _total_resp = len(df)
                    _summary_parts = [f"Council Overall: {_total_resp:,} respondents (100%)"]
                    for _bar_lbl, eq, val in _groups:
                        _n = int((df[eq] == val).sum())
                        _pct = _n / _total_resp
                        _summary_parts.append(f"{_bar_lbl}: {_n:,} ({_pct:.0%})")
                    _n_bars = len(_x_labels)
                    _p_colors = ["#8DC0D4"] + [PRIMARY] * (_n_bars - 1)
                    _i_colors = ["#F0B0B0"] + ["#E07070"] * (_n_bars - 1)
                    _theme_figs = []
                    for _wi, _theme in enumerate(WOW_THEMES):
                        _p_y = [_p_overall[_wi]] + [gp[_wi] for gp in _group_p]
                        _i_y = [_i_overall[_wi]] + [gi[_wi] for gi in _group_i]
                        _fig_w = go.Figure()
                        _fig_w.add_trace(go.Bar(
                            name="Place (P)", x=_x_labels, y=_p_y,
                            marker_color=_p_colors,
                            text=[f"{v:.2f}" for v in _p_y], textposition="outside",
                            textfont=dict(size=8, color="#1A2B3C"),
                            hovertemplate="<b>%{x}</b><br>Place: %{y:.2f}<extra></extra>",
                        ))
                        _fig_w.add_trace(go.Bar(
                            name="Individual (I)", x=_x_labels, y=_i_y,
                            marker_color=_i_colors,
                            text=[f"{v:.2f}" for v in _i_y], textposition="outside",
                            textfont=dict(size=8, color="#1A2B3C"),
                            hovertemplate="<b>%{x}</b><br>Individual: %{y:.2f}<extra></extra>",
                        ))
                        _fig_w.update_layout(
                            barmode="group",
                            title=dict(text=_theme, font=dict(color="#0F4C6B", size=14)),
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=40, b=120),
                            xaxis=dict(tickangle=-35, tickfont=dict(size=10, color="#1A2B3C")),
                            yaxis=dict(range=[-0.2, 4.5],
                                       title=dict(text="Avg Score", font=dict(color="#1A2B3C", size=10)),
                                       tickfont=dict(size=10, color="#1A2B3C"), gridcolor="#E8EEF2"),
                            legend=dict(orientation="h", yanchor="bottom", y=1.02,
                                        xanchor="right", x=1, font=dict(color="#1A2B3C", size=11)),
                            height=420,
                        )
                        _theme_figs.append((_theme, _fig_w))
                    # ── Report download button ───────────────────────────────
                    _fname = f"WoW_{edi_display.replace(' ', '')}.docx"
                    _rpt_key = f"{key_prefix}_rpt"
                    _mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    if _rpt_key not in st.session_state:
                        with st.spinner("Preparing report…"):
                            st.session_state[_rpt_key] = _make_wow_docx(edi_cols, edi_display)
                    st.download_button(
                        "📄 Download Report (.docx)",
                        data=st.session_state[_rpt_key],
                        file_name=_fname,
                        mime=_mime,
                        key=f"{key_prefix}_dl",
                    )
                    # ── Summary + charts ─────────────────────────────────────
                    st.markdown(
                        '<p style="font-size:12px;color:#5A7080;margin-bottom:12px">' +
                        " &nbsp;|&nbsp; ".join(_summary_parts) + "</p>",
                        unsafe_allow_html=True
                    )
                    for _wi, (_theme, _fig_w) in enumerate(_theme_figs):
                        _p_stmt = WOW_PLACE_STATEMENTS.get(_theme, "")
                        _i_stmt = WOW_IND_STATEMENTS.get(_theme, "")
                        if _p_stmt:
                            st.markdown(
                                f'<p style="font-size:11px;color:#5A7080;font-style:italic;margin-bottom:1px">'
                                f'Place: "{_p_stmt}"</p>'
                                f'<p style="font-size:11px;color:#5A7080;font-style:italic;margin-bottom:4px">'
                                f'Individual: "{_i_stmt}"</p>',
                                unsafe_allow_html=True
                            )
                        st.plotly_chart(_fig_w, use_container_width=True, key=f"{key_prefix}_{_wi}")
                        st.markdown("---")

                for (_edi_tab, (_edi_cols, _edi_name)) in zip(_d1_edi_tabs, _d1_edi_configs):
                    with _edi_tab:
                        st.markdown(f"#### Ways of Working — by {_edi_name}")
                        _wow_edi_charts(_edi_cols, _edi_name,
                                        f"d11_{_edi_name.replace(' ','_')}")

            # ── D1.2: Sentiment Outcomes ──────────────────────────────────────
            with d1_2:
                st.caption(
                    "For each employee experience theme, shows the average score for the council overall "
                    "and for every group within the selected EDI filter. Scroll down to see all 15 themes."
                )
                st.markdown(
                    '<p style="font-size:11px;color:#5A7080;margin-bottom:8px">'
                    '<strong>Scale:</strong> 0 = Strongly disagree &nbsp;·&nbsp; 1 = Disagree &nbsp;·&nbsp; '
                    '2 = Neither agree nor disagree &nbsp;·&nbsp; 3 = Agree &nbsp;·&nbsp; 4 = Strongly agree</p>',
                    unsafe_allow_html=True,
                )

                _d12_edi_tabs = st.tabs([
                    "Length of Service", "Age", "Gender",
                    "Ethnic Group", "Sexual Orientation", "Disabled", "Carer",
                ])
                _d12_edi_configs = [
                    (["Q8"],   "Length of Service"),
                    (["Q72"],  "Age"),
                    (["Q73"],  "Gender"),
                    (["Q74"],  "Ethnic Group"),
                    (["Q75"],  "Sexual Orientation"),
                    (["Q76"],  "Disabled"),
                    (["Q77"],  "Carer"),
                ]

                def _make_empex_docx(edi_cols, edi_display):
                    """Build and return Employee Experience Word report bytes for the given EDI group."""
                    import io as _io, copy as _copy
                    from docx import Document as _Document
                    from docx.shared import Inches as _Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
                    import plotly.io as _pio
                    _x = ["Council Overall"]
                    _o_ov = [df[c].mean() for c in OUTCOME_COLS]
                    _go, _grps = [], []
                    for eq in edi_cols:
                        _vals = sorted(
                            (v for v in df[eq].dropna().unique()
                             if str(v).strip() not in ("", "Not Answered")),
                            key=lambda v: LOS_ORDER.index(v) if eq == "Q8" and v in LOS_ORDER else v,
                        )
                        _pfx = (EDI_FILTERS.get(eq, eq) + ": ") if len(edi_cols) > 1 else ""
                        for val in _vals:
                            _sub = df[df[eq] == val]
                            _x.append(f"{_pfx}{val}")
                            _go.append([_sub[c].mean() for c in OUTCOME_COLS])
                            _grps.append((f"{_pfx}{val}", eq, val))
                    _tot = len(df)
                    _parts = [f"Council Overall: {_tot:,} respondents (100%)"]
                    for _bl, eq, val in _grps:
                        _n = int((df[eq] == val).sum())
                        _parts.append(f"{_bl}: {_n:,} ({_n / _tot:.0%})")
                    _oc = ["#8DC0D4"] + [PRIMARY] * (len(_x) - 1)
                    _doc = _Document()
                    _doc.core_properties.title = f"Employee Experience — {edi_display}"
                    _h = _doc.add_heading(f"Employee Experience — {edi_display}", 0)
                    _h.alignment = _ALIGN.CENTER
                    _sh = _doc.add_paragraph("Somerset Council Cultural Diagnostic — D1.2")
                    _sh.alignment = _ALIGN.CENTER
                    _doc.add_paragraph()
                    _doc.add_paragraph(
                        "Scale: 0 = Strongly disagree · 1 = Disagree · "
                        "2 = Neither agree nor disagree · 3 = Agree · 4 = Strongly agree"
                    )
                    _doc.add_paragraph()
                    _doc.add_heading("Sample Breakdown", 2)
                    for _sp in _parts:
                        _doc.add_paragraph(_sp, style="List Bullet")
                    _doc.add_page_break()
                    for _oi, _lbl in enumerate(OUTCOME_LABELS):
                        _is_rating = (_lbl == "Employer rating")
                        if _is_rating:
                            _y = [df["Q69_raw"].mean()] + [
                                df[df[eq] == val]["Q69_raw"].mean()
                                for _, eq, val in _grps
                            ]
                            _y_range = [0, 11]
                            _y_title = "Avg Rating (1–10)"
                        else:
                            _y = [_o_ov[_oi]] + [g[_oi] for g in _go]
                            _y_range = [-0.2, 4.5]
                            _y_title = "Avg Score"
                        _doc.add_heading(_lbl, 2)
                        _os = OUTCOME_STATEMENTS.get(_lbl, "")
                        if _os:
                            _doc.add_paragraph().add_run(f"\"{_os}\"").italic = True
                        _f = go.Figure()
                        _f.add_trace(go.Bar(name="Score", x=_x, y=_y, marker_color=_oc,
                                            text=[f"{v:.2f}" for v in _y], textposition="outside",
                                            textfont=dict(size=12, color="#1A2B3C"),
                                            hovertemplate="<b>%{x}</b><br>Score: %{y:.2f}<extra></extra>"))
                        _f.update_layout(
                            title=dict(text=_lbl, font=dict(color="#0F4C6B", size=18)),
                            font=dict(family="Inter", color="#1A2B3C", size=14),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=100, r=30, t=80, b=160),
                            xaxis=dict(tickangle=-35, tickfont=dict(size=13, color="#1A2B3C")),
                            yaxis=dict(range=_y_range,
                                       title=dict(text=_y_title, font=dict(color="#1A2B3C", size=13)),
                                       tickfont=dict(size=13, color="#1A2B3C"), gridcolor="#E8EEF2"),
                            showlegend=False, height=520,
                        )
                        _img = _pio.to_image(_f, format="png", width=1200, height=520, scale=2)
                        _doc.add_picture(_io.BytesIO(_img), width=_Inches(6.5))
                        _doc.add_paragraph()
                    _buf = _io.BytesIO()
                    _doc.save(_buf)
                    return _buf.getvalue()

                def _outcome_edi_charts(edi_cols, edi_display, key_prefix):
                    """Render 15 bar charts (one per outcome theme) with per-tab report download."""
                    _x_labels = ["Council Overall"]
                    _o_overall = [df[c].mean() for c in OUTCOME_COLS]
                    _group_o, _groups = [], []
                    for eq in edi_cols:
                        _eq_vals = sorted(
                            (v for v in df[eq].dropna().unique()
                             if str(v).strip() not in ("", "Not Answered")),
                            key=lambda v: LOS_ORDER.index(v) if eq == "Q8" and v in LOS_ORDER else v,
                        )
                        _prefix = (EDI_FILTERS.get(eq, eq) + ": ") if len(edi_cols) > 1 else ""
                        for val in _eq_vals:
                            _sub = df[df[eq] == val]
                            _x_labels.append(f"{_prefix}{val}")
                            _group_o.append([_sub[c].mean() for c in OUTCOME_COLS])
                            _groups.append((f"{_prefix}{val}", eq, val))
                    _total_resp = len(df)
                    _summary_parts = [f"Council Overall: {_total_resp:,} respondents (100%)"]
                    for _bar_lbl, eq, val in _groups:
                        _n = int((df[eq] == val).sum())
                        _pct = _n / _total_resp
                        _summary_parts.append(f"{_bar_lbl}: {_n:,} ({_pct:.0%})")
                    _o_colors = ["#8DC0D4"] + [PRIMARY] * (len(_x_labels) - 1)
                    _outcome_figs = []
                    for _oi, _lbl in enumerate(OUTCOME_LABELS):
                        _is_rating = (_lbl == "Employer rating")
                        if _is_rating:
                            _y = [df["Q69_raw"].mean()] + [
                                df[df[eq] == val]["Q69_raw"].mean()
                                for _, eq, val in _groups
                            ]
                            _y_range = [0, 11]
                            _y_title = "Avg Rating (1–10)"
                        else:
                            _y = [_o_overall[_oi]] + [_gdata[_oi] for _gdata in _group_o]
                            _y_range = [-0.2, 4.5]
                            _y_title = "Avg Score"
                        _fig_o = go.Figure()
                        _fig_o.add_trace(go.Bar(
                            name="Score", x=_x_labels, y=_y,
                            marker_color=_o_colors,
                            text=[f"{v:.2f}" for v in _y], textposition="outside",
                            textfont=dict(size=8, color="#1A2B3C"),
                            hovertemplate="<b>%{x}</b><br>Score: %{y:.2f}<extra></extra>",
                        ))
                        _fig_o.update_layout(
                            title=dict(text=_lbl, font=dict(color="#0F4C6B", size=14)),
                            font=dict(family="Inter", color="#1A2B3C"),
                            paper_bgcolor="#F7F9FC", plot_bgcolor="#F7F9FC",
                            margin=dict(l=10, r=10, t=40, b=120),
                            xaxis=dict(tickangle=-35, tickfont=dict(size=10, color="#1A2B3C")),
                            yaxis=dict(range=_y_range,
                                       title=dict(text=_y_title, font=dict(color="#1A2B3C", size=10)),
                                       tickfont=dict(size=10, color="#1A2B3C"), gridcolor="#E8EEF2"),
                            showlegend=False, height=420,
                        )
                        _outcome_figs.append((_lbl, _fig_o))
                    # ── Report download button ───────────────────────────────
                    _fname = f"EmpEx_{edi_display.replace(' ', '')}.docx"
                    _rpt_key = f"{key_prefix}_rpt"
                    _mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    if _rpt_key not in st.session_state:
                        with st.spinner("Preparing report…"):
                            st.session_state[_rpt_key] = _make_empex_docx(edi_cols, edi_display)
                    st.download_button(
                        "📄 Download Report (.docx)",
                        data=st.session_state[_rpt_key],
                        file_name=_fname,
                        mime=_mime,
                        key=f"{key_prefix}_dl",
                    )
                    # ── Summary + charts ─────────────────────────────────────
                    st.markdown(
                        '<p style="font-size:12px;color:#5A7080;margin-bottom:12px">' +
                        " &nbsp;|&nbsp; ".join(_summary_parts) + "</p>",
                        unsafe_allow_html=True
                    )
                    for _oi, (_lbl, _fig_o) in enumerate(_outcome_figs):
                        _stmt = OUTCOME_STATEMENTS.get(_lbl, "")
                        if _stmt:
                            st.markdown(
                                f'<p style="font-size:11px;color:#5A7080;font-style:italic;margin-bottom:4px">'
                                f'"{_stmt}"</p>',
                                unsafe_allow_html=True
                            )
                        st.plotly_chart(_fig_o, use_container_width=True, key=f"{key_prefix}_{_oi}")
                        st.markdown("---")

                for (_edi_tab, (_edi_cols, _edi_name)) in zip(_d12_edi_tabs, _d12_edi_configs):
                    with _edi_tab:
                        st.markdown(f"#### Employee Experience — by {_edi_name}")
                        _outcome_edi_charts(_edi_cols, _edi_name,
                                            f"d12_{_edi_name.replace(' ', '_')}")

            # ── Download all EDI reports as zip ───────────────────────────────
            st.markdown("---")
            st.markdown("#### Download All EDI Reports")
            st.caption(
                "Downloads all Ways of Working (D1.1) and Employee Experience (D1.2) "
                "reports as a single zip file, with one report per EDI group."
            )
            _zip_key = "d1_all_reports_zip"
            _edi_zip_configs = [
                ("Length_of_Service", ["Q8"],   "Length of Service"),
                ("Age",               ["Q72"],  "Age"),
                ("Gender",            ["Q73"],  "Gender"),
                ("Ethnic_Group",      ["Q74"],  "Ethnic Group"),
                ("Sexual_Orientation",["Q75"],  "Sexual Orientation"),
                ("Disabled",          ["Q76"],  "Disabled"),
                ("Carer",             ["Q77"],  "Carer"),
            ]
            if _zip_key not in st.session_state:
                with st.spinner("Assembling all reports into zip…"):
                    import zipfile as _zipfile
                    import io as _zip_io
                    _zip_buf = _zip_io.BytesIO()
                    with _zipfile.ZipFile(_zip_buf, "w", _zipfile.ZIP_DEFLATED) as _zf:
                        for _kname, _ecols, _ename in _edi_zip_configs:
                            _wow_bytes = (
                                st.session_state.get(f"d11_{_kname}_rpt")
                                or _make_wow_docx(_ecols, _ename)
                            )
                            _zf.writestr(
                                f"D1.1_Ways_of_Working/WoW_{_ename.replace(' ', '')}.docx",
                                _wow_bytes,
                            )
                            _ex_bytes = (
                                st.session_state.get(f"d12_{_kname}_rpt")
                                or _make_empex_docx(_ecols, _ename)
                            )
                            _zf.writestr(
                                f"D1.2_Employee_Experience/EmpEx_{_ename.replace(' ', '')}.docx",
                                _ex_bytes,
                            )
                    st.session_state[_zip_key] = _zip_buf.getvalue()
            st.download_button(
                "⬇ Download All EDI Reports (.zip)",
                data=st.session_state[_zip_key],
                file_name="EDI_Reports.zip",
                mime="application/zip",
                key="d1_zip_dl",
            )

    # ── Timing display ────────────────────────────────────────────────────────────────
    if st.session_state.get("_show_timing") and _timings:
        with st.sidebar:
            st.markdown("---")
            st.markdown(
                '<p style="color:#FFFFFF;font-weight:700;margin-bottom:4px">⏱ Render timing (this run)</p>',
                unsafe_allow_html=True,
            )
            for _lbl, _elapsed in sorted(_timings, key=lambda x: -x[1]):
                if _elapsed >= 5:
                    _dot = "🔴"
                elif _elapsed >= 1:
                    _dot = "🟠"
                else:
                    _dot = "🟢"
                st.markdown(
                    f'<p style="color:#FFFFFF;font-size:12px;margin:2px 0">'
                    f'{_dot} {_lbl}: <b>{_elapsed:.2f}s</b></p>',
                    unsafe_allow_html=True,
                )
